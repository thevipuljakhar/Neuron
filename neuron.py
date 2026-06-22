"""
Neuron v4 - Indian RE Industry Intelligence Monitor
by Vipul Jakhar
v4: Daily Intelligence Brief, Signal Intelligence Score (CRITICAL/HIGH/MED/LOW),
    7-day sparklines, NDC Progress Tracker, IEX Power Market, ICED data integration
v3: SAATVIKGL, SQLite alert persistence, source health, SECI scraper,
    correlation matrix, forex/USDINR, configurable watchlist, alert history
"""
import json, time, re, os, io, sqlite3, math
from datetime import datetime, timedelta
from flask import Flask, jsonify, render_template, request
import yfinance as yf
import requests
import feedparser
import numpy as np
import pandas as pd
from concurrent.futures import ThreadPoolExecutor, as_completed
from bs4 import BeautifulSoup
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
import sources as v11_sources
import swot_engine
import brain_centers.metacognitive as _metacog
import brain_centers.backup_center  as _backup_center
import brain_centers.security_center as _security_center
import curiosity_engine as _curiosity

app = Flask(__name__)
cache = {}

# ── Security headers (applied to every response) ─────────────────────────────
@app.after_request
def _add_security_headers(response):
    # Prevent clickjacking
    response.headers.setdefault("X-Frame-Options", "SAMEORIGIN")
    # Prevent MIME-type sniffing
    response.headers.setdefault("X-Content-Type-Options", "nosniff")
    # Basic XSS filter for older browsers
    response.headers.setdefault("X-XSS-Protection", "1; mode=block")
    # Don't leak Referer to external sites
    response.headers.setdefault("Referrer-Policy", "same-origin")
    # Restrict powerful features
    response.headers.setdefault("Permissions-Policy", "geolocation=(), camera=(), microphone=()")
    return response

# ── Telegram Notifications (optional — set env vars or .env to enable) ────────
# Secrets live OUTSIDE source: real env vars, or a local .env file next to
# neuron.py (KEY=VALUE lines, never committed). No defaults in code — the old
# committed token is burned and must be rotated via @BotFather.
def _load_dotenv():
    try:
        p = os.path.join(os.path.dirname(__file__), ".env")
        if os.path.exists(p):
            for line in open(p, encoding="utf-8"):
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    k, v = line.split("=", 1)
                    os.environ.setdefault(k.strip(), v.strip())
    except Exception:
        pass
_load_dotenv()

TELEGRAM_BOT_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN", "")
TELEGRAM_CHAT_ID   = os.environ.get("TELEGRAM_CHAT_ID", "")
# P18 — OpenWeatherMap (optional, keyless-degrading). Enriches the RE-hub weather
# with current cloud cover/temp/wind (cloud cover ⇒ near-term solar yield).
OPENWEATHER_API_KEY = os.environ.get("OPENWEATHER_API_KEY", "")
# P20 — optional macro/trade keys. IMF DataMapper is keyless (always on); EIA and
# TradingEconomics light up only when their keys are present (degrade otherwise).
EIA_API_KEY = os.environ.get("EIA_API_KEY", "")
TRADINGECONOMICS_KEY = os.environ.get("TRADINGECONOMICS_KEY", "")

def _get_chat_id():
    """Return TELEGRAM_CHAT_ID — env var takes precedence, then SQLite-stored value."""
    if TELEGRAM_CHAT_ID:
        return TELEGRAM_CHAT_ID
    try:
        con = sqlite3.connect(DB_PATH)
        row = con.execute("SELECT value FROM kv_store WHERE key='tg_chat_id' LIMIT 1").fetchone()
        con.close()
        return row[0] if row else ""
    except Exception:
        return ""

def _save_chat_id(chat_id: str):
    """Persist discovered chat_id to SQLite kv_store."""
    try:
        con = sqlite3.connect(DB_PATH)
        con.execute("CREATE TABLE IF NOT EXISTS kv_store (key TEXT PRIMARY KEY, value TEXT)")
        con.execute("INSERT OR REPLACE INTO kv_store(key,value) VALUES('tg_chat_id',?)", (chat_id,))
        con.commit(); con.close()
    except Exception:
        pass

def send_telegram(message, parse_mode="HTML"):
    """Send message via Telegram Bot API. Falls back to plain text if HTML parse fails."""
    chat_id = _get_chat_id()
    if not TELEGRAM_BOT_TOKEN or not chat_id:
        return False
    url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage"
    try:
        # Try with requested parse_mode first
        r = requests.post(url, json={"chat_id": chat_id, "text": message,
                                      "parse_mode": parse_mode}, timeout=10)
        if r.status_code == 200:
            return True
        # If HTML parse failed, retry as plain text (strip tags)
        if parse_mode == "HTML":
            import re as _re
            plain = _re.sub(r'<[^>]+>', '', message)
            r2 = requests.post(url, json={"chat_id": chat_id, "text": plain}, timeout=10)
            return r2.status_code == 200
        return False
    except Exception:
        return False

# Telegram alert cooldowns — prevent duplicate sends within a session
_tg_last_intel_alert = 0   # epoch seconds of last intel score alert
_tg_last_seci_count  = 0   # last SECI tender count alerted

CACHE_TTL = 300

# ── NaN-safe JSON provider — browser rejects Python's NaN/Infinity ───────────
def _fix_nan(obj):
    """Recursively replace float NaN/Inf with None so JSON is browser-safe."""
    if isinstance(obj, float):
        return None if (math.isnan(obj) or math.isinf(obj)) else obj
    if isinstance(obj, dict):  return {k: _fix_nan(v) for k, v in obj.items()}
    if isinstance(obj, list):  return [_fix_nan(i) for i in obj]
    if isinstance(obj, tuple): return tuple(_fix_nan(i) for i in obj)
    return obj

_orig_jsonify = jsonify
def jsonify(*args, **kwargs):
    if args and not kwargs:
        args = (_fix_nan(args[0]),)
    else:
        kwargs = {k: _fix_nan(v) for k, v in kwargs.items()}
    return _orig_jsonify(*args, **kwargs)

# ── Stock Universe ───────────────────────────────────────────────────────────
RE_STOCKS_DEFAULT = {
    "ADANIGREEN.NS":"Adani Green","NHPC.NS":"NHPC","NTPC.NS":"NTPC",
    "SUZLON.NS":"Suzlon","SWSOLAR.NS":"Sterling Wilson","IREDA.NS":"IREDA",
    "SJVN.NS":"SJVN","WAAREEENER.NS":"Waaree","PREMIERENE.NS":"Premier Energy",
    "TATAPOWER.NS":"Tata Power","TORNTPOWER.NS":"Torrent Power",
    "INOXWIND.NS":"Inox Wind","BORORENEW.NS":"Borosil Renewables",
    "SAATVIKGL.NS":"Saatvik Green",
}

WATCHLIST_PATH = os.path.join(os.path.dirname(__file__), "watchlist.json")

def load_watchlist():
    wl = dict(RE_STOCKS_DEFAULT)
    try:
        if os.path.exists(WATCHLIST_PATH):
            with open(WATCHLIST_PATH) as f:
                custom = json.load(f)
            wl.update(custom)
    except Exception:
        pass
    return wl

RE_STOCKS = load_watchlist()

# ── Market Universe ──────────────────────────────────────────────────────────
RSS_FEEDS = [
    # ── India RE-specialist sources (high-trust, RE-only content) ──────────────
    ("Mercom India",    "https://mercomindia.com/feed/"),
    ("SAU Energy",      "https://www.saurenergy.com/feed"),
    ("Solar Quarter",   "https://solarquarter.com/feed/"),
    ("REGlobal",        "https://reglobal.co/feed/"),
    ("EQ Mag",          "https://www.eqmagpro.com/feed/"),
    # ── Indian business/energy (RE section — filtered by score) ──────────────
    ("ET Energy",       "https://economictimes.indiatimes.com/industry/energy/rssfeeds/13357270.cms"),
    ("Mint Energy",     "https://www.livemint.com/rss/energy"),   # energy-specific, not general industry
    # ── Global RE/clean-energy specialist ────────────────────────────────────
    ("PV Tech",         "https://www.pv-tech.org/feed/"),
    ("CleanTechnica",   "https://cleantechnica.com/feed/"),
    ("Energy Monitor",  "https://www.energymonitor.ai/feed/"),
    # ── Global wire — filtered strictly to RE/energy keywords before display ──
    ("Reuters Energy",  "https://feeds.reuters.com/reuters/businessNews"),
    ("BBC World",       "http://feeds.bbci.co.uk/news/world/rss.xml"),
]

GLOBAL_NEWS_SOURCES = {"Reuters Energy", "BBC World"}

# Articles from global wire sources MUST match at least one of these to be included
_GLOBAL_RE_REQUIRED = [
    "solar","wind","renewable","clean energy","green hydrogen","battery storage","bess",
    "electric vehicle","ev charg","carbon","emissions","climate","net zero","energy transition",
    "power grid","coal plant","nuclear plant","lng","polysilicon","panel",
    "mnre","seci","ntpc","adani","waaree","vikram","irena","iea energy",
]

COMMODITIES = {
    "GC=F":"Gold","CL=F":"Crude Oil","SI=F":"Silver",
    "NG=F":"Natural Gas","ALI=F":"Aluminum",
}
FOREX = {
    "USDINR=X":"USD/INR","EURINR=X":"EUR/INR",
}
GLOBAL_RE = {
    "ICLN":"iShares Global CE","QCLN":"First Trust NASDAQ CE",
    "NEE":"NextEra Energy","ENPH":"Enphase Energy","SEDG":"SolarEdge",
    "FSLR":"First Solar","BEP":"Brookfield RE","TAN":"Invesco Solar ETF",
    "RNW":"ReNew Energy","CSIQ":"Canadian Solar",
    # World panel tickers
    "^GSPC":"S&P 500","MCHI":"iShares MSCI China","EWJ":"iShares MSCI Japan",
}

# Tickers that must NOT have .NS appended (they are US/global, not NSE-listed)
_NON_NSE_TICKERS = (
    set(GLOBAL_RE.keys())
    | set(COMMODITIES.keys())
    | set(FOREX.keys())
    | {"^GSPC", "^NSEI", "^CNXENERGY", "^CNX100", "^CNXAUTO"}
)

# India macro indices — fetched separately, not in GLOBAL_RE (avoid polluting that feed)
INDIA_INDICES = {
    "^NSEI":      "Nifty 50",
    "^CNXENERGY": "Nifty Energy",
    "^CNX100":    "Nifty 100",
}

def fetch_india_indices():
    cached = get_cache("india_indices", 900)
    if cached: return cached
    results = {}
    with ThreadPoolExecutor(max_workers=4) as ex:
        futs = {ex.submit(fetch_quote_generic, s, n): s for s, n in INDIA_INDICES.items()}
        for f in as_completed(futs):
            d = f.result()
            if d: results[futs[f]] = d
    set_cache("india_indices", results)
    return results

@app.route("/api/india_indices")
def api_india_indices(): return jsonify(fetch_india_indices())

ALERT_KEYWORDS = {
    "SUPPLY CHAIN":["polysilicon","supply chain","cell shortage","module price","wafer","backsheet","glass shortage"],
    "TRADE":       ["tariff","anti-dumping","ALMM","BCD","AD/CVD","import duty","safeguard","trade war","customs"],
    "POLICY":      ["PLI","VGF","SECI tender","MNRE notification","carbon credit","RPO","REC","budget allocation"],
    "BESS/GH2":    ["battery storage","BESS","green hydrogen","electrolyzer","GH2","hydrogen mission","NGHM","pumped hydro"],
    "MARKET":      ["L1 tariff","auction result","bid","capacity addition","MW commissioned","GW installed","solar auction"],
    "COMPANIES":   ["Waaree","Premier Energy","Adani Green","IREDA","NHPC","Suzlon","Tata Power","Saatvik","Vikram Solar","SAATVIKGL"],
    "GLOBAL":      ["China solar","US IRA","EU solar","IRENA report","BloombergNEF","polysilicon price","panel oversupply"],
}

# ── Intel Engine Constants (v6) ──────────────────────────────────────────────
INTEL_CATEGORIES = {
    "PROJECT_WIN":         ["wins order","bags order","awarded","l1 bidder","letter of award"," loa ","emerges l1","lowest bidder","bid win","wins bid","project awarded","secures order","gets order","wins contract"],
    "COMMISSIONING":       ["commissions","commissioned","inaugurates","operationalises","goes live","begins operations","first unit","capacity operationalised","coc received","formally inaugurated"],
    "EXPANSION":           ["expansion","expands capacity","new plant","greenfield","brownfield","sets up","plans to add","new facility","to build","capacity addition","to install","to develop","scaling up"],
    "TENDER_ISSUED":       ["issues tender","floats tender","invites bids","rfp issued","eoi invited","new tender","fresh tender","calls for bids","invites expression","issues rfp","floats rfp"],
    "POLICY_NOTIFICATION": ["mnre notification","cerc order","serc order","ministry circular","gazette notification","policy amendment","new regulation","rpo target","mnre order","mnre circular","bcd notification","almm order"],
    "FUNDING":             ["raises funds","ipo","fpo","qip","secures debt","loan sanctioned","ncd","bonds issued","equity raise","fundraise","raises ₹","raises rs","secured funding"],
    "SUPPLY_CHAIN":        ["polysilicon","module price","cell price","wafer price","supply chain","import duty","anti-dumping","ad/cvd","glass shortage","backsheet","module shortage","panel price","input cost"],
    "TARIFF_SIGNAL":       ["₹/kwh","rs/kwh","tariff discovered","tariff trend","auction tariff","l1 tariff","bid tariff","power tariff","lowest tariff","tariff falls","tariff rises"],
    "GLOBAL_MACRO":        ["china solar","us ira","eu solar","polysilicon china","panel oversupply","irena","bloombergnef","global capacity","china module","us tariff on","wto ruling"],
    "REGULATORY_RISK":     ["penalty","show cause","notice issued","default","cancellation","terminated","curtailment","discoms default","land acquisition issue","regulatory hurdle","forced shutdown"],
    "M_AND_A":             ["acquires","merger","acquisition","stake buy","joint venture","mou signed","strategic partnership","takeover","buyout","stake acquisition"],
    "EARNINGS_SIGNAL":     ["quarterly results","q1 fy","q2 fy","q3 fy","q4 fy"," pat ","revenue growth","order book","ebitda","guidance","margin expansion","margin pressure","results beat","results miss"],
}
INTEL_DIRECTION = {
    "PROJECT_WIN":"POSITIVE","COMMISSIONING":"POSITIVE","EXPANSION":"POSITIVE",
    "FUNDING":"POSITIVE","TARIFF_SIGNAL":"POSITIVE","M_AND_A":"POSITIVE",
    "TENDER_ISSUED":"POSITIVE","POLICY_NOTIFICATION":"NEUTRAL",
    "SUPPLY_CHAIN":"NEGATIVE","REGULATORY_RISK":"NEGATIVE",
    "GLOBAL_MACRO":"NEUTRAL","EARNINGS_SIGNAL":"NEUTRAL",
}
INTEL_CATEGORY_SCORE = {
    "PROJECT_WIN":90,"COMMISSIONING":80,"REGULATORY_RISK":75,"EXPANSION":70,
    "FUNDING":65,"SUPPLY_CHAIN":65,"TARIFF_SIGNAL":60,"EARNINGS_SIGNAL":55,
    "TENDER_ISSUED":50,"M_AND_A":50,"POLICY_NOTIFICATION":45,"GLOBAL_MACRO":35,
}
COMPANY_ENTITIES = {
    "saatvik green":   "SAATVIKGL.NS",  "saatvik":         "SAATVIKGL.NS",
    "adani green":     "ADANIGREEN.NS", "agel":            "ADANIGREEN.NS",
    "waaree":          "WAAREEENER.NS",
    "premier energy":  "PREMIERENE.NS",
    "sterling wilson": "SWSOLAR.NS",    "swsolar":         "SWSOLAR.NS",
    "nhpc":            "NHPC.NS",
    "ntpc":            "NTPC.NS",
    "sjvn":            "SJVN.NS",
    "ireda":           "IREDA.NS",
    "suzlon":          "SUZLON.NS",
    "inox wind":       "INOXWIND.NS",
    "tata power":      "TATAPOWER.NS",
    "torrent power":   "TORNTPOWER.NS",
    "borosil":         "BORORENEW.NS",
}

# v12: each channel carries a known-good 24/7 fallback videoId (worldmonitor
# pattern) so a failed live-resolution degrades to a working stream, never a
# black frame. "fallback" IDs are long-running live streams.
LIVE_CHANNELS = [
    {"name":"WION",       "id":"UCsB-sMFo8gznkLsNt5NU0mg","color":"#b58900","fallback":""},
    {"name":"ET Now",     "id":"UCJim7HNvOmCJRLJnQJp0czw","color":"#268bd2","fallback":""},
    {"name":"CNBC TV18",  "id":"UC7HExiGZiGPJqOfkfANczQA","color":"#2aa198","fallback":""},
    {"name":"India TV",   "id":"UCE_Uy-xEpFiRLmpTEkxc-LA","color":"#859900","fallback":""},
    {"name":"Bloomberg",  "id":"UCIALMKvObZNtJ6AmdCLP7Lg","color":"#dc322f","fallback":"iEpJwprxDdk"},
    {"name":"CNBC",       "id":"UCvJJ_dzjViJCoLf5uKUTwoA","color":"#e8a030","fallback":"9NyxcX3rhQs"},
    {"name":"Sky News",   "id":"UCoMdktPbSTixAyNGwb-UYkQ","color":"#4a9eff","fallback":"uvviIF4725I"},
    {"name":"DW News",    "id":"UCknLrEdhRCp1aegoMqRaCZg","color":"#8b5cf6","fallback":"LuKwFajn37U"},
    {"name":"DD News",    "id":"UCF2MmTp-Q6HlCJDjdHWF9Rw","color":"#6c71c4","fallback":""},
    {"name":"NewsX",      "id":"UCFZIZGhSwRSGF6Mn__cTjIA","color":"#d33682","fallback":""},
    {"name":"Al Jazeera", "id":"UCNye-wNBqNL5ZzHSJj3l8Bg","color":"#cb4b16","fallback":"gCNeDWCI0vo"},
]

# ── SQLite Persistence ───────────────────────────────────────────────────────
DB_PATH = os.path.join(os.path.dirname(__file__), "neuron.db")

def init_db():
    con = sqlite3.connect(DB_PATH)
    # WAL removes write contention between the daily-brief daemon and request threads
    con.execute("PRAGMA journal_mode=WAL")
    con.execute("""CREATE TABLE IF NOT EXISTS alerts_seen
                   (uid TEXT PRIMARY KEY, ts REAL)""")
    con.execute("""CREATE TABLE IF NOT EXISTS alerts_log
                   (uid TEXT PRIMARY KEY, title TEXT, source TEXT,
                    category TEXT, keywords TEXT, link TEXT, ts REAL)""")
    con.execute("""CREATE TABLE IF NOT EXISTS almm_meta
                   (id INTEGER PRIMARY KEY, pdf_url TEXT, pub_date TEXT,
                    parsed_at REAL, record_count INTEGER)""")
    con.execute("""CREATE TABLE IF NOT EXISTS almm_list
                   (id INTEGER PRIMARY KEY AUTOINCREMENT,
                    mfr TEXT, model TEXT, capacity_wp REAL, efficiency REAL,
                    technology TEXT, validity_date TEXT, parent_company TEXT,
                    meta_id INTEGER)""")
    con.execute("""CREATE TABLE IF NOT EXISTS news_archive
                   (uid TEXT PRIMARY KEY, source TEXT, title TEXT, link TEXT,
                    summary TEXT, published_dt TEXT, days_old REAL, ts REAL)""")
    con.execute("""CREATE TABLE IF NOT EXISTS almm_modules_meta
                   (id INTEGER PRIMARY KEY, pdf_url TEXT, pub_date TEXT,
                    parsed_at REAL, record_count INTEGER)""")
    con.execute("""CREATE TABLE IF NOT EXISTS almm_modules
                   (id INTEGER PRIMARY KEY AUTOINCREMENT,
                    mfr TEXT, model TEXT, capacity_mw_yr REAL, efficiency REAL,
                    module_type TEXT, module_wp REAL, validity_from TEXT,
                    validity_to TEXT, parent_company TEXT, meta_id INTEGER)""")
    con.execute("""CREATE TABLE IF NOT EXISTS pulse_history
                   (id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts REAL, date TEXT, pulse INTEGER, label TEXT,
                    hot_topics TEXT, articles_processed INTEGER)""")
    con.execute("""CREATE TABLE IF NOT EXISTS cea_statewise_snap
                   (id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts REAL, date TEXT, region TEXT, state TEXT,
                    cumulative_mw REAL, monthly_mw REAL)""")
    # National RE capacity snapshots — for historical growth chart (P6.3)
    con.execute("""CREATE TABLE IF NOT EXISTS cea_national_snap
                   (id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts REAL, snap_date TEXT,
                    re_total_mw REAL, solar_mw REAL, wind_mw REAL,
                    hydro_mw REAL, grand_total_mw REAL)""")
    # Key-value store — persists Telegram chat_id and other runtime config
    con.execute("""CREATE TABLE IF NOT EXISTS kv_store
                   (key TEXT PRIMARY KEY, value TEXT)""")
    # P7: Sector breadth history — daily % of RE stocks above SMA
    con.execute("""CREATE TABLE IF NOT EXISTS breadth_history
                   (id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts REAL, snap_date TEXT,
                    breadth_20 REAL, breadth_50 REAL, breadth_200 REAL)""")
    # P7: Fear & Greed history — daily composite RE investment climate score
    con.execute("""CREATE TABLE IF NOT EXISTS fear_greed_history
                   (id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts REAL, snap_date TEXT, score REAL, label TEXT,
                    breadth REAL, pulse_dir REAL, rsi_nifty REAL, brent_dir REAL)""")
    # P7: Add sentiment columns to pulse_history (safe re-run)
    for col in ("bullish_count INTEGER DEFAULT 0",
                "bearish_count INTEGER DEFAULT 0",
                "neutral_count INTEGER DEFAULT 0"):
        try:
            con.execute(f"ALTER TABLE pulse_history ADD COLUMN {col}")
        except Exception:
            pass
    # P21 — Tender tracker (v21_tenders): entity/project/type/capacity/dates/sector
    con.execute("""CREATE TABLE IF NOT EXISTS v21_tenders
                   (id INTEGER PRIMARY KEY AUTOINCREMENT,
                    entity TEXT, project_name TEXT, tender_type TEXT,
                    capacity_mw REAL, sector TEXT, state TEXT,
                    announced_date TEXT, bid_deadline TEXT, ppa_signed_date TEXT,
                    status TEXT DEFAULT 'OPEN', source_url TEXT, source_title TEXT,
                    ts REAL DEFAULT (strftime('%s','now')))""")
    con.commit(); con.close()
    _init_user_data()

def _init_user_data():
    """Create user_data/ folder and template Excel files if they don't exist."""
    try:
        import openpyxl
        ud = os.path.join(os.path.dirname(__file__), "user_data")
        os.makedirs(ud, exist_ok=True)
        # PM Kusum template
        kusum_path = os.path.join(ud, "pm_kusum.xlsx")
        if not os.path.exists(kusum_path):
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "PM_KUSUM"
            ws.append(["State_Name", "Total_Sanction_MW", "Total_Installed_MW"])
            sample = [
                ("Assam",2,0),("Chhattisgarh",366,9),("Goa",24,4),
                ("Gujarat",175,0),("Haryana",158,60.21),("Himachal Pradesh",100,100),
                ("Madhya Pradesh",1401,148.57),("Maharashtra",260,21),("Odisha",90,0),
                ("Rajasthan",5618,1040.99),("Tamil Nadu",3,3),("Telangana",1797,50),
                ("Tripura",5,0),("Uttar Pradesh",1,1),("Total",10000,1437.77),
            ]
            for row in sample: ws.append(row)
            ws["A1"].font = openpyxl.styles.Font(bold=True)
            wb.save(kusum_path)
        # PM Surya Ghar template
        sg_path = os.path.join(ud, "pm_surya_ghar.xlsx")
        if not os.path.exists(sg_path):
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "PM_Surya_Ghar"
            ws.append(["State / UT", "Applications (No.)", "Installations (No.)",
                       "Households Covered (No.)", "Installation Capacity (MW)",
                       "Subsidy Released (Cr)"])
            states_sg = [
                "Andhra Pradesh","Bihar","Goa","Gujarat","Haryana",
                "Himachal Pradesh","Karnataka","Kerala","Madhya Pradesh","Maharashtra",
                "Odisha","Punjab","Rajasthan","Tamil Nadu","Telangana",
                "Uttar Pradesh","Uttarakhand","West Bengal",
            ]
            for s in states_sg: ws.append([s, 0, 0, 0, 0, 0])
            ws["A1"].font = openpyxl.styles.Font(bold=True)
            wb.save(sg_path)
        # README
        readme = os.path.join(ud, "README.txt")
        if not os.path.exists(readme):
            with open(readme, "w") as f:
                f.write("NEURON user_data/ — manually maintained data files\n\n"
                        "pm_kusum.xlsx     : PM-KUSUM Component-A state-wise data\n"
                        "                    Columns: State_Name | Total_Sanction_MW | Total_Installed_MW\n"
                        "                    Update this file with latest MNRE data as needed.\n\n"
                        "pm_surya_ghar.xlsx: PM Surya Ghar state-wise data\n"
                        "                    Columns: State_Name | Applications_Registered | Sanctioned | Installed\n"
                        "                    Update this file with latest MNRE data as needed.\n\n"
                        "Neuron reads these files on every request (no cache bypass needed).\n"
                        "Just save the file and refresh the dashboard.\n")
    except Exception as _e:
        pass  # openpyxl not available or write error — non-fatal

def load_seen_uids():
    try:
        con = sqlite3.connect(DB_PATH)
        rows = con.execute("SELECT uid FROM alerts_seen").fetchall()
        con.close()
        return set(r[0] for r in rows)
    except Exception:
        return set()

def persist_alert_db(uid, alert):
    try:
        con = sqlite3.connect(DB_PATH)
        con.execute("INSERT OR IGNORE INTO alerts_seen VALUES (?,?)", (uid, time.time()))
        con.execute("INSERT OR IGNORE INTO alerts_log VALUES (?,?,?,?,?,?,?)",
            (uid, alert.get("title",""), alert.get("source",""),
             alert.get("category",""), ",".join(alert.get("keywords",[])),
             alert.get("link",""), time.time()))
        con.commit(); con.close()
    except Exception:
        pass

def get_alert_history(limit=100):
    try:
        con = sqlite3.connect(DB_PATH)
        rows = con.execute(
            "SELECT title,source,category,keywords,link,ts FROM alerts_log ORDER BY ts DESC LIMIT ?",
            (limit,)
        ).fetchall()
        con.close()
        return [{"title":r[0],"source":r[1],"category":r[2],
                 "keywords":r[3].split(","),"link":r[4],
                 "date":datetime.fromtimestamp(r[5]).strftime("%Y-%m-%d %H:%M")} for r in rows]
    except Exception:
        return []

init_db()
SEEN_ALERTS = load_seen_uids()

# ── P7: Term Spike Detector — in-memory sliding window ──────────────────────
_TERM_FREQ     = {}   # {term: [(ts, source), ...]}  — 2-hour sliding window
_TERM_BASELINE = {}   # {term: daily_avg (EMA)}
_TERM_COOLDOWN = {}   # {term: last_alert_ts}  — 30-min per-term cooldown

SPIKE_TERMS = [
    "waaree","adani","vikram","saatvik","premier","torrent","ireda","suzlon","nhpc","ntpc",
    "hybrid","bess","battery storage","green hydrogen","electrolyzer",
    "almm","bcd","rpo","vgf","mnre","seci","tariff","anti-dumping",
    "polysilicon","curtailment","commissioning","tender","mw","gw",
]

def _update_term_freq(articles):
    now    = time.time()
    cutoff = now - 7200
    for a in articles:
        text = (a.get("title","") + " " + a.get("summary","")).lower()
        src  = a.get("source","")
        for term in SPIKE_TERMS:
            if term in text:
                if term not in _TERM_FREQ:
                    _TERM_FREQ[term] = []
                _TERM_FREQ[term].append((now, src))
    for term in list(_TERM_FREQ.keys()):
        _TERM_FREQ[term] = [(ts, s) for ts, s in _TERM_FREQ[term] if ts > cutoff]
        if not _TERM_FREQ[term]:
            del _TERM_FREQ[term]

def _compute_spikes():
    spikes = []
    for term, events in _TERM_FREQ.items():
        sources  = set(s for _, s in events)
        count    = len(events)
        baseline = _TERM_BASELINE.get(term, 1.0)
        if count >= baseline * 2.5 and len(sources) >= 2:
            spikes.append({
                "term": term, "count": count,
                "baseline": round(baseline, 1),
                "ratio":    round(count / max(baseline, 1), 1),
                "sources":  list(sources)[:4],
            })
    spikes.sort(key=lambda x: x["ratio"], reverse=True)
    return spikes[:6]

def _update_baseline(articles):
    counts = {}
    for a in articles:
        text = (a.get("title","") + " " + a.get("summary","")).lower()
        for term in SPIKE_TERMS:
            if term in text:
                counts[term] = counts.get(term, 0) + 1
    alpha = 0.15
    for term, cnt in counts.items():
        daily_equiv = cnt / 6
        old = _TERM_BASELINE.get(term, daily_equiv)
        _TERM_BASELINE[term] = round(old * (1 - alpha) + daily_equiv * alpha, 2)

# ── Source Health ────────────────────────────────────────────────────────────
SOURCE_HEALTH = {}

def mark_health(source, ok, msg=""):
    SOURCE_HEALTH[source] = {
        "status": "ok" if ok else "error",
        "last_checked": datetime.now().strftime("%H:%M:%S"),
        "msg": msg[:120] if msg else "",
    }

# ── Gov-site HTTP ────────────────────────────────────────────────────────────
GOV_UA = {"User-Agent": "Mozilla/5.0 Chrome/124 Safari/537.36"}

def gov_get(url, timeout=20, headers=None, method="get", **kw):
    """Single chokepoint for government-site HTTP. verify=False is a deliberate
    tradeoff (NIC/gov TLS chains are broken) accepted ONCE, here — never inline."""
    fn = requests.head if method == "head" else requests.get
    return fn(url, headers=headers or GOV_UA, timeout=timeout, verify=False, **kw)

# ── Cache ────────────────────────────────────────────────────────────────────
def get_cache(key, ttl=None):
    e = cache.get(key)
    if e and time.time()-e["ts"] < (ttl or CACHE_TTL): return e["data"]
    return None

def set_cache(key, data):
    cache[key] = {"data":data, "ts":time.time()}

def serialized(fn):
    """Stampede guard: serialize concurrent calls to one expensive fetcher.
    Two cold hits on /api/almm/modules used to trigger two 280-page PDF parses;
    with this, the second caller blocks then returns the freshly warmed cache.
    Cache-hit calls pay only a microsecond lock acquire."""
    import threading as _th
    lock = _th.Lock()
    def wrap(*a, **k):
        with lock:
            return fn(*a, **k)
    wrap.__name__ = getattr(fn, "__name__", "serialized")
    return wrap

# ── Stock Fetchers ───────────────────────────────────────────────────────────
def fetch_quote(symbol):
    t = yf.Ticker(symbol)
    hist = None
    # Try 5d first; fall back to 1mo if empty (some tickers throttled on short periods)
    for period in ("5d", "1mo"):
        try:
            h = t.history(period=period)
            if not h.empty:
                hist = h
                break
            time.sleep(1)
        except Exception:
            time.sleep(1)
    if hist is None or hist.empty:
        mark_health(symbol, False, "no data from yfinance (5d+1mo)")
        return None
    try:
        close = float(hist["Close"].iloc[-1])
        prev  = float(hist["Close"].iloc[-2]) if len(hist)>1 else close
        chg   = close-prev; chgp = (chg/prev*100) if prev else 0
        mark_health(symbol, True)
        # 7-day sparkline for stock card mini-chart
        sparkline = [round(float(v), 2) for v in hist["Close"].values[-7:]]
        return {"symbol":symbol,"name":RE_STOCKS.get(symbol,symbol),
                "price":round(close,2),"change":round(chg,2),"change_pct":round(chgp,2),
                "volume":int(hist["Volume"].iloc[-1]),
                "high":round(float(hist["High"].iloc[-1]),2),
                "low":round(float(hist["Low"].iloc[-1]),2),
                "mktcap":getattr(t.fast_info,"market_cap",None),
                "sparkline": sparkline}
    except Exception as e:
        mark_health(symbol, False, str(e))
        return None

def fetch_history(symbol, period="1y"):
    try:
        t = yf.Ticker(symbol); hist = t.history(period=period)
        return [{"date":str(d.date()),"open":round(float(r.Open),2),"high":round(float(r.High),2),
                 "low":round(float(r.Low),2),"close":round(float(r.Close),2),"volume":int(r.Volume)}
                for d,r in hist.iterrows()]
    except: return []

# High-liquidity stocks safe for parallel fetch; others go sequential
_LIQUID = {
    "ADANIGREEN.NS","NHPC.NS","NTPC.NS","SUZLON.NS","TATAPOWER.NS",
    "TORNTPOWER.NS","INOXWIND.NS","IREDA.NS","SJVN.NS","WAAREEENER.NS",
    "SWSOLAR.NS","SAATVIKGL.NS",
}

def fetch_all_quotes():
    cached = get_cache("all_quotes")
    if cached: return cached
    results = {}
    liquid    = {s:n for s,n in RE_STOCKS.items() if s in _LIQUID}
    illiquid  = {s:n for s,n in RE_STOCKS.items() if s not in _LIQUID}

    # Parallel fetch for liquid large-caps
    with ThreadPoolExecutor(max_workers=6) as ex:
        futs = {ex.submit(fetch_quote,s):s for s in liquid}
        for f in as_completed(futs):
            d = f.result()
            if d: results[futs[f]] = d

    # Sequential fetch for smaller tickers — 3s gap avoids Yahoo rate-limit after burst
    if illiquid:
        time.sleep(3)
        for sym in illiquid:
            d = fetch_quote(sym)
            if d: results[sym] = d

    # Final retry for any still missing
    failed = [s for s in RE_STOCKS if s not in results]
    if failed:
        time.sleep(2)
        for sym in failed:
            d = fetch_quote(sym)
            if d: results[sym] = d

    set_cache("all_quotes", results); return results

def fetch_quote_generic(symbol, name, health_group=None):
    """Generic quote with the same resilience shape as fetch_quote: period
    fallback (5d→1mo) + one delayed retry. Thin tickers (TAN, CSIQ, RNW) get
    throttled by Yahoo and a bare period="2d"/no-retry fetch dropped them
    silently. When health_group is given, per-symbol failures are recorded in
    SOURCE_HEALTH (e.g. "Global RE: TAN") so a dropped ticker is visible in
    /api/health instead of vanishing behind the aggregate flag."""
    def _hkey(): return f"{health_group}: {symbol}" if health_group else None
    hist = None
    for period in ("5d", "1mo"):
        try:
            h = yf.Ticker(symbol).history(period=period)
            if not h.empty:
                hist = h
                break
            time.sleep(1)
        except Exception:
            time.sleep(1)
    if hist is None or hist.empty:
        if health_group: mark_health(_hkey(), False, "no data from yfinance (5d+1mo)")
        return None
    try:
        close = float(hist["Close"].iloc[-1])
        prev  = float(hist["Close"].iloc[-2]) if len(hist)>1 else close
        chg   = close-prev
        if health_group: mark_health(_hkey(), True)
        return {"symbol":symbol,"name":name,"price":round(close,2),"change":round(chg,2),
                "change_pct":round((chg/prev*100) if prev else 0, 2)}
    except Exception as e:
        if health_group: mark_health(_hkey(), False, str(e))
        return None

def fetch_commodities():
    cached = get_cache("commodities")
    if cached: return cached
    results = {}
    with ThreadPoolExecutor(max_workers=6) as ex:
        futs = {ex.submit(fetch_quote_generic,s,n):s for s,n in {**COMMODITIES,**FOREX}.items()}
        for f in as_completed(futs):
            d = f.result()
            if d: results[futs[f]] = d
    mark_health("Commodities/Forex", bool(results))
    set_cache("commodities", results); return results

def fetch_global_re():
    cached = get_cache("global_re")
    if cached: return cached
    results = {}
    # Sequential, not 10-way parallel: global RE ETFs/ADRs (TAN, CSIQ, RNW...)
    # are thin and Yahoo throttles them under a parallel burst, dropping tickers
    # silently. Mirror fetch_all_quotes()'s throttled pattern instead.
    for sym, name in GLOBAL_RE.items():
        d = fetch_quote_generic(sym, name, health_group="Global RE")
        if d: results[sym] = d
        time.sleep(0.5)
    # One retry pass for any still missing (matches fetch_all_quotes resilience).
    missing = [s for s in GLOBAL_RE if s not in results]
    if missing:
        time.sleep(2)
        for sym in missing:
            d = fetch_quote_generic(sym, GLOBAL_RE[sym], health_group="Global RE")
            if d: results[sym] = d
            time.sleep(0.5)
    still_missing = [s for s in GLOBAL_RE if s not in results]
    mark_health("Global RE", bool(results),
                "" if not still_missing else f"{len(still_missing)} missing: {','.join(still_missing)}")
    set_cache("global_re", results); return results

# ── Global Installed RE Capacity (IRENA PxWeb — public, no key) ───────────────
# P14 Item 2: real nameplate installed-capacity-by-country (MW), not stock
# prices. Source is IRENA's open PxWeb statistics API (English). GEM requires a
# registration form and IEA/Ember are Cloudflare-blocked, so IRENA — the actual
# authority on installed capacity — is the honest key-less choice. Updates
# ~half-yearly, so cached 24h and persisted to kv_store for stale-fallback.
IRENA_CAP_URL = ("https://pxweb.irena.org/api/v1/en/IRENASTAT/"
                 "Power Capacity and Generation/Country_ELECCAP_2026_H1_v-PX 1.px")
# Technology codes (from table metadata): 0=Total RE, 2=Solar PV, 4=Wind.
_IRENA_TECH = {"0": "total_re_mw", "2": "solar_pv_mw", "4": "wind_mw"}

def _kv_read(key):
    try:
        con = sqlite3.connect(DB_PATH, timeout=15)
        con.execute("CREATE TABLE IF NOT EXISTS kv_store (key TEXT PRIMARY KEY, value TEXT)")
        row = con.execute("SELECT value FROM kv_store WHERE key=?", (key,)).fetchone()
        con.close()
        return json.loads(row[0]) if row else None
    except Exception:
        return None

def _kv_write(key, obj):
    try:
        con = sqlite3.connect(DB_PATH, timeout=15)
        con.execute("CREATE TABLE IF NOT EXISTS kv_store (key TEXT PRIMARY KEY, value TEXT)")
        con.execute("INSERT OR REPLACE INTO kv_store(key,value) VALUES(?,?)",
                    (key, json.dumps(obj)))
        con.commit(); con.close()
    except Exception:
        pass

def fetch_global_installed_capacity():
    cached = get_cache("global_capacity", ttl=86400)
    if cached: return cached
    query = {"query": [
        {"code": "Technology", "selection": {"filter": "item", "values": ["0", "2", "4"]}},
        {"code": "Grid connection", "selection": {"filter": "item", "values": ["0"]}},
        {"code": "Year", "selection": {"filter": "item", "values": ["25"]}},
    ], "response": {"format": "json-stat2"}}
    try:
        r = requests.post(IRENA_CAP_URL, headers={**GOV_UA, "Content-Type": "application/json"},
                          data=json.dumps(query), timeout=30)
        r.raise_for_status()
        js = r.json()
        dims = js["id"]                      # dimension order, e.g. [Country, Technology, Grid, Year]
        size = js["size"]
        ci, ti = dims.index("Country/area"), dims.index("Technology")
        cdim = js["dimension"]["Country/area"]
        clabels = cdim["category"]["label"]
        cindex  = cdim["category"]["index"]  # code -> position
        code_by_pos = {pos: code for code, pos in cindex.items()}
        tdim = js["dimension"]["Technology"]
        tindex = tdim["category"]["index"]
        values = js["value"]
        # strides for row-major json-stat2 value array
        stride = [1] * len(size)
        for i in range(len(size) - 2, -1, -1):
            stride[i] = stride[i + 1] * size[i + 1]
        year = js["dimension"]["Year"]["category"]["label"]
        as_of = list(year.values())[0] if year else "?"

        # IRENA's Country/area dimension includes one regional aggregate (REA =
        # Eurasia); drop it so the ranking is genuinely by country.
        IRENA_AGGREGATES = {"REA"}
        countries = {}
        for cpos in range(size[ci]):
            code = code_by_pos.get(cpos)
            if code in IRENA_AGGREGATES:
                continue
            row = {"country": clabels.get(code, code), "code": code}
            ok = False
            for tcode, field in _IRENA_TECH.items():
                tpos = tindex.get(tcode)
                if tpos is None: continue
                idx = cpos * stride[ci] + tpos * stride[ti]
                v = values[idx] if idx < len(values) else None
                row[field] = round(v, 1) if isinstance(v, (int, float)) else None
                if row[field]: ok = True
            if ok: countries[code] = row
        ranked = sorted(countries.values(), key=lambda x: x.get("total_re_mw") or 0, reverse=True)
        india = next((c for c in ranked if c["code"] == "IND"), None)
        out = {"as_of": as_of, "unit": "MW", "source": "IRENA Renewable Capacity Statistics",
               "top": ranked[:25], "india": india, "country_count": len(ranked),
               "fetched_at": datetime.now().strftime("%Y-%m-%d %H:%M"), "stale": False}
        _kv_write("global_capacity_lastgood", out)
        mark_health("IRENA Capacity", True)
        set_cache("global_capacity", out); return out
    except Exception as e:
        mark_health("IRENA Capacity", False, str(e))
        last = _kv_read("global_capacity_lastgood")
        if last:
            last = {**last, "stale": True}
            set_cache("global_capacity", last); return last
        return {"as_of": None, "unit": "MW", "source": "IRENA Renewable Capacity Statistics",
                "top": [], "india": None, "country_count": 0, "stale": True,
                "error": "IRENA fetch failed and no cached data yet"}

def fetch_solar_capacity_history(country="IND", project_to=2030):
    """P14 Item 5: real India solar-PV installed-capacity time series (GW) from
    IRENA's yearly figures, replacing the old hardcoded JS arrays — plus a
    simple CAGR projection forward (no ML, matching project style). Cached 24h,
    persisted for stale-fallback like the capacity fetcher."""
    ck = f"solar_cap_hist_{country}"
    cached = get_cache(ck, ttl=86400)
    if cached: return cached
    query = {"query": [
        {"code": "Country/area", "selection": {"filter": "item", "values": [country]}},
        {"code": "Technology", "selection": {"filter": "item", "values": ["2"]}},
        {"code": "Grid connection", "selection": {"filter": "item", "values": ["0"]}},
    ], "response": {"format": "json-stat2"}}
    try:
        r = requests.post(IRENA_CAP_URL, headers={**GOV_UA, "Content-Type": "application/json"},
                          data=json.dumps(query), timeout=30)
        r.raise_for_status()
        js = r.json()
        ydim = js["dimension"]["Year"]
        ylabels = ydim["category"]["label"]            # code -> "2000"
        yindex  = ydim["category"]["index"]            # code -> position
        values  = js["value"]
        # Only Year varies here (country/tech/grid each fixed to 1) → position == value index.
        series = []
        for code, pos in sorted(yindex.items(), key=lambda kv: kv[1]):
            v = values[pos] if pos < len(values) else None
            if isinstance(v, (int, float)):
                series.append({"year": int(ylabels[code]), "gw": round(v / 1000.0, 2)})
        series.sort(key=lambda x: x["year"])
        # CAGR over the last 5 available years → forward projection to project_to.
        proj = []
        if len(series) >= 2:
            last = series[-1]
            base = series[-6] if len(series) >= 6 else series[0]
            n = last["year"] - base["year"]
            cagr = ((last["gw"] / base["gw"]) ** (1.0 / n) - 1.0) if (n and base["gw"] > 0) else 0.0
            cur = last["gw"]
            for yr in range(last["year"] + 1, project_to + 1):
                cur = cur * (1.0 + cagr)
                proj.append({"year": yr, "gw": round(cur, 1)})
        out = {"country": country, "as_of": series[-1]["year"] if series else None,
               "source": "IRENA Renewable Capacity Statistics", "unit": "GW",
               "history": series, "projection": proj,
               "cagr_pct": round(cagr * 100, 1) if len(series) >= 2 else None,
               "fetched_at": datetime.now().strftime("%Y-%m-%d %H:%M"), "stale": False}
        _kv_write(f"{ck}_lastgood", out)
        mark_health("IRENA Solar History", True)
        set_cache(ck, out); return out
    except Exception as e:
        mark_health("IRENA Solar History", False, str(e))
        last = _kv_read(f"{ck}_lastgood")
        if last:
            last = {**last, "stale": True}; set_cache(ck, last); return last
        return {"country": country, "as_of": None, "unit": "GW", "history": [],
                "projection": [], "stale": True, "error": "IRENA fetch failed, no cache"}

def fetch_wind_tech_mix(country="IND"):
    """P14 Item 6: real wind technology mix (Onshore vs Offshore installed MW)
    from IRENA. This is the honest, live wind-technology breakdown — drivetrain
    class (DFIG/PMSG) is not published in any clean key-less list, so onshore/
    offshore is the technology dimension we can show truthfully. Sits next to
    the ALMM solar-module-tech chart."""
    ck = f"wind_tech_mix_{country}"
    cached = get_cache(ck, ttl=86400)
    if cached: return cached
    query = {"query": [
        {"code": "Country/area", "selection": {"filter": "item", "values": [country]}},
        {"code": "Technology", "selection": {"filter": "item", "values": ["5", "6"]}},
        {"code": "Grid connection", "selection": {"filter": "item", "values": ["0"]}},
        {"code": "Year", "selection": {"filter": "item", "values": ["25"]}},
    ], "response": {"format": "json-stat2"}}
    try:
        r = requests.post(IRENA_CAP_URL, headers={**GOV_UA, "Content-Type": "application/json"},
                          data=json.dumps(query), timeout=30)
        r.raise_for_status()
        js = r.json()
        tdim = js["dimension"]["Technology"]
        tindex = tdim["category"]["index"]; tlabels = tdim["category"]["label"]
        values = js["value"]
        mix = []
        for code, pos in sorted(tindex.items(), key=lambda kv: kv[1]):
            v = values[pos] if pos < len(values) else None
            mix.append({"tech": tlabels[code], "mw": round(v, 1) if isinstance(v, (int, float)) else 0})
        as_of = list(js["dimension"]["Year"]["category"]["label"].values())[0]
        out = {"country": country, "as_of": as_of, "unit": "MW", "mix": mix,
               "note": "Drivetrain class (DFIG/PMSG) not publicly listed; onshore/offshore per IRENA.",
               "source": "IRENA Renewable Capacity Statistics", "stale": False,
               "fetched_at": datetime.now().strftime("%Y-%m-%d %H:%M")}
        _kv_write(f"{ck}_lastgood", out)
        mark_health("IRENA Wind Tech", True)
        set_cache(ck, out); return out
    except Exception as e:
        mark_health("IRENA Wind Tech", False, str(e))
        last = _kv_read(f"{ck}_lastgood")
        if last:
            last = {**last, "stale": True}; set_cache(ck, last); return last
        return {"country": country, "as_of": None, "unit": "MW", "mix": [],
                "stale": True, "error": "IRENA fetch failed, no cache"}

# ── Energy Market Prices ─────────────────────────────────────────────────────
ENERGY_TICKERS = {
    "NG=F":    {"name": "Henry Hub Gas",  "unit": "$/MMBtu", "region": "us"},
    "TTF=F":   {"name": "TTF Natural Gas","unit": "€/MWh",   "region": "europe"},
    "CL=F":    {"name": "Brent Crude",    "unit": "$/bbl",   "region": "global"},
    "MTF=F":   {"name": "Thermal Coal",   "unit": "$/t",     "region": "global"},
}

def fetch_energy_prices():
    cached = get_cache("energy_prices", 1800)
    if cached: return cached
    results = []
    for sym, meta in ENERGY_TICKERS.items():
        try:
            t = yf.Ticker(sym)
            h = t.history(period="2d", interval="1d")
            if h.empty: continue
            row   = h.iloc[-1]
            prev  = h.iloc[-2] if len(h) > 1 else row
            price = round(float(row["Close"]), 3)
            chg   = round(float(row["Close"] - prev["Close"]), 3)
            chg_p = round(chg / float(prev["Close"]) * 100, 2) if prev["Close"] else 0
            results.append({
                "symbol":  sym,
                "name":    meta["name"],
                "unit":    meta["unit"],
                "region":  meta["region"],
                "price":   price,
                "change":  chg,
                "change_pct": chg_p,
            })
        except Exception:
            pass
    # EU ETS carbon — yfinance has no clean ticker; use static reference
    results.append({
        "symbol":  "EUA",
        "name":    "EU Carbon (ETS)",
        "unit":    "€/t",
        "region":  "europe",
        "price":   65.0,
        "change":  0,
        "change_pct": 0,
        "static":  True,
    })
    set_cache("energy_prices", results)
    return results

# ── MNRE Live ────────────────────────────────────────────────────────────────
@serialized
def fetch_mnre_live():
    cached = get_cache("mnre_live", 3600)
    if cached: return cached
    try:
        hdr = {"User-Agent":"Mozilla/5.0 Chrome/124 Safari/537.36"}
        r   = gov_get("https://mnre.gov.in/en/physical-progress/", timeout=20)
        soup = BeautifulSoup(r.text,"html.parser")
        table = soup.find("table")
        if not table: return {}
        rows = table.find_all("tr")
        header = [c.get_text(strip=True) for c in rows[0].find_all(["th","td"])]
        # Footnote markers attached to source names — strip them and record their meaning
        # * = includes off-grid solar (small systems not grid-metered)
        # ^ = includes allocated shares in joint/central sector utilities (CEA basis)
        # # = includes nuclear capacity additions under MNRE tracking scope
        FOOTNOTES = {
            "*": "incl. off-grid solar",
            "^": "incl. allocated shares in joint/central utilities",
            "#": "incl. nuclear capacity tracked under non-fossil target",
        }
        def _strip_footnote(raw):
            notes = []
            clean = raw.strip()
            for marker, desc in FOOTNOTES.items():
                if clean.endswith(marker):
                    clean = clean[:-len(marker)].strip()
                    notes.append(desc)
            return clean, notes

        data = {}
        for row in rows[2:]:
            cells = [c.get_text(strip=True) for c in row.find_all(["td","th"])]
            if len(cells)>=4 and cells[0] and cells[3]:
                try:
                    clean_name, notes = _strip_footnote(cells[0])
                    # Aggregate rows (Sub Total / Total RE / Total Non-Fossil) must never
                    # be summed together with technology rows — that quadruples the total.
                    is_agg = "total" in clean_name.lower()
                    data[clean_name] = {
                        "monthly_mw":   float(cells[1].replace(",","")) if cells[1] else 0,
                        "cumulative_mw":float(cells[3].replace(",","")) if cells[3] else 0,
                        "footnotes":    notes,
                        "is_aggregate": is_agg,
                        "is_re":        not (is_agg or "nuclear" in clean_name.lower()),
                    }
                except: pass
        # Canonical totals read straight from MNRE's own aggregate rows
        _tre = next((v["cumulative_mw"] for k,v in data.items() if k.lower()=="total re"), 0)
        if not _tre:  # fall back to summing technology (non-aggregate, RE) rows only
            _tre = sum(v["cumulative_mw"] for v in data.values() if v.get("is_re"))
        result = {"data":data,"as_on":header[3] if len(header)>3 else "",
                  "total_re_mw": round(_tre, 2),
                  "fetched_at":datetime.now().strftime("%Y-%m-%d %H:%M")}
        set_cache("mnre_live", result)
        mark_health("MNRE", True)
        return result
    except Exception as e:
        mark_health("MNRE", False, str(e)); return {"error":str(e)}

# ── CEA Excel ────────────────────────────────────────────────────────────────
def _cea_url():
    hdr = {"User-Agent":"Mozilla/5.0 Chrome/124 Safari/537.36"}
    # Try scraping the page for an explicit xlsx link first
    try:
        r   = gov_get("https://cea.nic.in/installed-capacity-report/?lang=en", timeout=15)
        soup = BeautifulSoup(r.text,"html.parser")
        for a in soup.find_all("a",href=True):
            if ".xlsx" in a["href"].lower() and "installed" in a["href"].lower():
                href = a["href"]
                if not href.startswith("http"): href = "https://cea.nic.in" + href
                return href
    except: pass
    # CEA data typically lags 1-2 months — try up to 4 months back
    n = datetime.now()
    for delta in range(4):
        month = n.month - delta
        year  = n.year
        if month <= 0:
            month += 12
            year  -= 1
        url = f"https://cea.nic.in/wp-content/uploads/installed/{year}/{month:02d}/Website.xlsx"
        try:
            chk = gov_get(url, timeout=8, method="head")
            if chk.status_code == 200:
                return url
        except: pass
    # Hard fallback — last known good file (April 2026)
    return "https://cea.nic.in/wp-content/uploads/installed/2026/04/Website.xlsx"

def fetch_cea_capacity():
    cached = get_cache("cea_cap", 86400)
    if cached: return cached
    try:
        hdr = {"User-Agent":"Mozilla/5.0 Chrome/124 Safari/537.36"}
        url = _cea_url()
        r   = gov_get(url, timeout=20)
        if r.status_code != 200: return {}
        df  = pd.read_excel(io.BytesIO(r.content),sheet_name="Summary",header=None,engine='openpyxl')
        result = {"source_url":url,"fetched_at":datetime.now().strftime("%Y-%m-%d")}
        for _,row in df.iterrows():
            non_null = [v for v in row.values if pd.notna(v)]
            if len(non_null)>=2:
                try:
                    strs = [str(v).strip() for v in non_null if isinstance(v,str) and str(v).strip()]
                    nums = [float(v) for v in non_null if isinstance(v,(int,float)) and not isinstance(v,bool)]
                    if strs and len(nums)>=2:
                        label=strs[-1]; mw=nums[0]; pct=nums[1]
                        if mw>100: result[label]={"mw":round(mw,2),"pct":round(pct*100,2)}
                except: pass
        set_cache("cea_cap", result)
        mark_health("CEA", True)
        return result
    except Exception as e:
        mark_health("CEA", False, str(e)); return {"error":str(e)}

@serialized
def fetch_statewise():
    cached = get_cache("statewise", 86400)
    if cached: return cached
    try:
        hdr = {"User-Agent":"Mozilla/5.0 Chrome/124 Safari/537.36"}
        url = _cea_url()
        r   = gov_get(url, timeout=30)
        r.raise_for_status()
        raw = r.content
        df  = pd.read_excel(io.BytesIO(raw), sheet_name="IC", header=None, engine='openpyxl')
        # IC sheet columns (0-indexed):
        # 1=Region  2=Ownership  3=Coal  4=Lignite  5=Gas  6=Diesel
        # 7=Thermal Total  8=Nuclear  9=Hydro  10=RES*(MNRE)  11=RE Total(Hydro+RES)  12=Grand Total

        # Parse Summary sheet for national solar/wind breakdown
        national = {"solar_mw": 0, "wind_mw": 0, "re_total_mw": 0, "hydro_mw": 0}
        try:
            ds = pd.read_excel(io.BytesIO(raw), sheet_name="Summary", header=None, engine='openpyxl')
            for _, sr in ds.iterrows():
                vals = [v for v in sr.values if pd.notna(v)]
                strs = [str(v).strip() for v in vals if isinstance(v, str)]
                nums = [float(v) for v in vals if isinstance(v, (int, float)) and not isinstance(v, bool)]
                if strs and nums:
                    lbl = " ".join(strs).lower()
                    if "solar" in lbl and "wind" not in lbl and nums:
                        national["solar_mw"] = round(nums[0], 2)
                    elif "wind" in lbl and "solar" not in lbl and nums:
                        national["wind_mw"] = round(nums[0], 2)
                    elif "hydro (including" in lbl and nums:
                        national["hydro_mw"] = round(nums[0], 2)
                    elif "res (including hydro)" in lbl and nums:
                        national["re_total_mw"] = round(nums[0], 2)
        except Exception:
            pass
        REGIONS = {"Northern Region","Western Region","Southern Region",
                   "Eastern Region","North Eastern Region","Islands"}
        STATE_REGION = {
            "Delhi":"Northern","Haryana":"Northern","Himachal Pradesh":"Northern",
            "Jammu & Kashmir and \nLadakh":"Northern","Jammu & Kashmir and Ladakh":"Northern",
            "Punjab":"Northern","Rajasthan":"Northern","Uttar Pradesh":"Northern","Uttarakhand":"Northern",
            "Chhattisgarh":"Western","Goa":"Western","Gujarat":"Western",
            "Madhya Pradesh":"Western","Maharashtra":"Western",
            "Andhra Pradesh":"Southern","Karnataka":"Southern","Kerala":"Southern",
            "Tamil Nadu":"Southern","Telangana":"Southern","Puducherry":"Southern",
            "Bihar":"Eastern","Jharkhand":"Eastern","Odisha":"Eastern",
            "Sikkim":"Eastern","West Bengal":"Eastern",
            "Arunachal Pradesh":"North Eastern","Assam":"North Eastern",
            "Manipur":"North Eastern","Meghalaya":"North Eastern","Mizoram":"North Eastern",
            "Nagaland":"North Eastern","Tripura":"North Eastern",
        }
        sd = {}             # regions
        state_data = {}     # individual states
        current_region = None
        current_state  = None
        SKIP_HEADERS   = {"region", "state ", "state", "ownership/ sector",
                           "all india", "nan", ""}
        for _, row in df.iterrows():
            cell1 = re.sub(r'\s+', ' ', str(row.iloc[1])).strip() if pd.notna(row.iloc[1]) else ""
            cell2 = str(row.iloc[2]).strip() if pd.notna(row.iloc[2]) else ""
            if cell1 in REGIONS:
                current_region = cell1; current_state = None; continue
            # Region subtotal (no hyphen)
            # col 11 = RE Total (Large Hydro + RES/MNRE) — the correct 279.3 GW figure
            # col 10 = RES*(MNRE) only = 227.6 GW (excludes large hydro, wrong for RE Total)
            if cell2 == "Sub Total" and current_region:
                try:
                    res = float(row.iloc[11]) if pd.notna(row.iloc[11]) else 0
                    tot = float(row.iloc[12]) if pd.notna(row.iloc[12]) else 0
                    if res > 0 or tot > 0:
                        sd[current_region] = {"res_mw": round(res,2), "total_mw": round(tot,2)}
                except: pass
                continue
            # State name row (cell2 == "State" and cell1 is a real state)
            if cell2 == "State" and cell1 and cell1.lower() not in SKIP_HEADERS:
                current_state = cell1.strip()
                continue
            # State subtotal (hyphen)
            if cell2 == "Sub-Total" and current_state:
                try:
                    res = float(row.iloc[11]) if pd.notna(row.iloc[11]) else 0
                    tot = float(row.iloc[12]) if pd.notna(row.iloc[12]) else 0
                    if res > 0 or tot > 0:
                        rgn = STATE_REGION.get(current_state, current_region or "")
                        state_data[current_state] = {
                            "res_mw":   round(res, 2),
                            "total_mw": round(tot, 2),
                            "region":   rgn,
                        }
                except: pass
                current_state = None
        srt = sorted(sd.items(), key=lambda x: x[1]["res_mw"], reverse=True)
        top = {k: v for k, v in srt}
        top_states = dict(sorted(state_data.items(),
                                 key=lambda x: x[1]["res_mw"], reverse=True)[:15])
        result = {"states": top, "top_states": top_states,
                  "national": national,
                  "as_on": datetime.now().strftime("%B %Y"),
                  "source_url": url, "granularity": "region"}
        set_cache("statewise", result)
        # Store national snapshot for historical growth chart
        if national.get("re_total_mw", 0) > 0:
            try:
                snap_date = datetime.now().strftime("%Y-%m")
                _con = sqlite3.connect(DB_PATH)
                exists = _con.execute("SELECT 1 FROM cea_national_snap WHERE snap_date=?", (snap_date,)).fetchone()
                if not exists:
                    _con.execute(
                        "INSERT INTO cea_national_snap(ts,snap_date,re_total_mw,solar_mw,wind_mw,hydro_mw,grand_total_mw) VALUES(?,?,?,?,?,?,?)",
                        (time.time(), snap_date, national["re_total_mw"], national.get("solar_mw",0),
                         national.get("wind_mw",0), national.get("hydro_mw",0),
                         sum(v["total_mw"] for v in top.values())))
                    _con.commit()
                _con.close()
            except Exception:
                pass
        mark_health("CEA Statewise", True, f"{len(top)} regions · {url.split('/')[-3]}/{url.split('/')[-2]}")
        return result
    except Exception as e:
        mark_health("CEA Statewise", False, str(e)); return {"error": str(e), "states": {}}

# ── SECI Tenders (requests + BS4, no JS needed, 4hr cache) ──────────────────
def fetch_seci_tenders():
    cached = get_cache("seci_tenders", 14400)
    if cached: return cached
    try:
        hdr = {"User-Agent": "Mozilla/5.0 Chrome/124 Safari/537.36"}
        r   = gov_get("https://seci.co.in/tenders/", timeout=20)
        soup = BeautifulSoup(r.text, "html.parser")
        tenders = []
        table = soup.find("table")
        if table:
            for row in table.find_all("tr")[1:]:
                cells   = [c.get_text(strip=True) for c in row.find_all(["td","th"])]
                link_el = row.find("a", href=True)
                link    = ("https://seci.co.in" + link_el["href"]) if link_el else ""
                # Row format: [blank, TenderID, ETS_ref, TenderRef, Title, PubDate, BidDate, ViewDetails]
                if len(cells) >= 5 and cells[4]:
                    title = cells[4]
                    # Extract MW/GW capacity from title
                    mw = 0.0
                    # (?!h) rejects MWh/GWh (BESS energy ratings are not capacity)
                    mw_m = re.search(r'([\d,]+(?:\.\d+)?)\s*(GW|MW|MWp|MWdc)(?!h)', title, re.IGNORECASE)
                    if mw_m:
                        val = float(mw_m.group(1).replace(",", ""))
                        mw  = val * 1000 if mw_m.group(2).upper() == "GW" else val
                    # Detect technology type
                    tech = "Solar"
                    tl = title.lower()
                    if "wind" in tl and "solar" not in tl:    tech = "Wind"
                    elif "hybrid" in tl or ("solar" in tl and "wind" in tl): tech = "Hybrid"
                    elif "bess" in tl or "storage" in tl or "battery" in tl: tech = "BESS"
                    elif "hydrogen" in tl or "green h" in tl: tech = "GH2"
                    tenders.append({
                        "tender_id": cells[1],
                        "ets_ref":   cells[2],
                        "ref":       cells[3],
                        "title":     title,
                        "pub_date":  cells[5] if len(cells) > 5 else "",
                        "deadline":  cells[6] if len(cells) > 6 else "",
                        "link":      link,
                        "status":    "Active",
                        "mw":        mw,
                        "tech":      tech,
                    })
        total_mw  = sum(t["mw"] for t in tenders)
        tech_dist = {}
        for t in tenders:
            tech_dist[t["tech"]] = tech_dist.get(t["tech"], 0) + 1
        result = {
            "tenders":    tenders[:20],
            "fetched_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "source":     "seci.co.in",
            "count":      len(tenders),
            "total_mw":   round(total_mw, 0),
            "tech_dist":  tech_dist,
        }
        # Telegram alert — fire only if tender count is higher than last alerted
        global _tg_last_seci_count
        if tenders and len(tenders) > _tg_last_seci_count:
            new_n = len(tenders) - _tg_last_seci_count
            titles = "\n".join(f"• {t['title'][:60]}" for t in tenders[:3])
            send_telegram(
                f"📋 <b>SECI Tenders Update</b>\n"
                f"{new_n} new tenders found ({len(tenders)} total)\n\n"
                f"{titles}"
            )
            _tg_last_seci_count = len(tenders)
        set_cache("seci_tenders", result)
        mark_health("SECI", True)
        return result
    except Exception as e:
        mark_health("SECI", False, str(e))
        return {"error": str(e), "tenders": []}

# ── PM Schemes (direct Playwright — scrapling's patchright dep is not installed) ──
def _playwright_fetch(url, wait_selector=None):
    """Shared Playwright helper — returns (html_content, plain_text) or raises."""
    from playwright.sync_api import sync_playwright
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page    = browser.new_page()
        page.goto(url, timeout=30000)
        try:
            page.wait_for_load_state("networkidle", timeout=12000)
        except Exception:
            pass
        if wait_selector:
            try:
                page.wait_for_selector(wait_selector, timeout=8000)
            except Exception:
                pass
        html = page.content()
        browser.close()
    return html

# ── MNRE State-wise RE Installed Capacity PDF (P9) ───────────────────────────
MNRE_STATE_CAP_FALLBACK = "https://cdnbbsr.s3waas.gov.in/s3716e1b8c6cd17b771da77391355749f3/uploads/2026/06/20260608821718495.pdf"

def _mnre_state_cap_url():
    hdr = {"User-Agent": "Mozilla/5.0 Chrome/124 Safari/537.36"}
    for page_url in [
        "https://mnre.gov.in/physical-progress-achievements/",
        "https://mnre.gov.in/",
    ]:
        try:
            r = gov_get(page_url, timeout=12)
            soup = BeautifulSoup(r.text, "html.parser")
            for a in soup.find_all("a", href=True):
                href = a["href"]
                txt  = (a.get_text(strip=True) + " " + href).lower()
                if href.lower().endswith(".pdf") and any(k in txt for k in ("state","statewise","installed","capacity","location")):
                    if not href.startswith("http"): href = "https://mnre.gov.in" + href
                    return href
        except Exception:
            pass
    return MNRE_STATE_CAP_FALLBACK

def fetch_mnre_state_capacity():
    """Parse MNRE state-wise RE installed capacity PDF. Columns:
    0=SNo, 1=LGD, 2=State, 3=Solar Ground, 4=RTS(PM-Surya), 5=Hybrid,
    6=KUSUM Comp-B, 7=Solar Total, 8=Wind, 9=Bio Bagasse, 10=Bio NonBagasse,
    11=WtE, 12=WtE Offgrid, 13=Bio Total, 14=SHP, 15=Large Hydro, 16=Total RE
    """
    cached = get_cache("mnre_state_cap", 86400)
    if cached: return cached
    try:
        import pdfplumber, tempfile
        hdr = {"User-Agent": "Mozilla/5.0 Chrome/124 Safari/537.36"}
        url = _mnre_state_cap_url()
        r = gov_get(url, timeout=45)
        r.raise_for_status()
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(r.content); tmp_path = tmp.name

        states = {}
        national = {}

        _SKIP_CELLS = {"state", "states", "uts", "s.no", "lgd", "lgdcode", "solar",
                       "wind", "bio", "power", "energy", "hydro", "capacity",
                       "ministry", "mw", "mega", "biomass", "cogen", "waste",
                       "ground", "mounted", "hybrid", "yojana", "inclusive"}

        def _mw(cells, idx):
            try:
                if idx >= len(cells): return 0.0
                v = cells[idx].replace(",", "").strip()
                return float(v) if v and v not in ("", "-") else 0.0
            except Exception:
                return 0.0

        with pdfplumber.open(tmp_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    if not table: continue
                    for row in table:
                        if not row or len(row) < 10: continue
                        cells = [re.sub(r'\s+', ' ', str(c or '')).strip() for c in row]
                        sno  = cells[0]
                        name = cells[2] if len(cells) > 2 else ""
                        if not name: continue
                        name_lo = name.lower().replace(" ", "")
                        # Data row: S.No is a plain integer
                        if re.match(r'^\d+$', sno.strip()):
                            if not re.search(r'[A-Za-z]{3,}', name): continue
                            if any(k in name_lo for k in _SKIP_CELLS): continue
                            states[name] = {
                                "solar_ground":  _mw(cells, 3),
                                "solar_rts":     _mw(cells, 4),
                                "solar_hybrid":  _mw(cells, 5),
                                "solar_kusum":   _mw(cells, 6),
                                "solar_total":   _mw(cells, 7),
                                "wind":          _mw(cells, 8),
                                "bio_bagasse":   _mw(cells, 9),
                                "bio_noncogen":  _mw(cells, 10),
                                "waste_energy":  _mw(cells, 11),
                                "bio_total":     _mw(cells, 13),
                                "small_hydro":   _mw(cells, 14),
                                "large_hydro":   _mw(cells, 15),
                                "total_re":      _mw(cells, 16),
                            }
                        # National total row
                        elif "total" in name_lo and len(cells) >= 16:
                            national = {
                                "solar":      _mw(cells, 7),
                                "wind":       _mw(cells, 8),
                                "bio":        _mw(cells, 13),
                                "small_hydro":_mw(cells, 14),
                                "large_hydro":_mw(cells, 15),
                                "total_re":   _mw(cells, 16),
                            }

        os.unlink(tmp_path)
        result = {
            "states":     states,
            "national":   national,
            "source_url": url,
            "fetched_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        }
        set_cache("mnre_state_cap", result)
        mark_health("MNRE State Cap", True, f"{len(states)} states · {url.split('/')[-1]}")
        return result
    except Exception as e:
        mark_health("MNRE State Cap", False, str(e))
        return {"error": str(e), "states": {}, "national": {}}

# ── user_data xlsx helpers (P18) — total-row-safe, JSON-safe ──────────────────
_TOTAL_LABELS = {"total", "grand total", "all india", "all-india", "india",
                 "sub total", "subtotal", "g.total", "grand-total"}

def _clean_num(v):
    """Coerce a cell to a JSON-safe number (None for NaN/blank), else passthrough."""
    try:
        if v is None:
            return None
        f = float(v)
        return None if (math.isnan(f) or math.isinf(f)) else round(f, 2)
    except (TypeError, ValueError):
        return None

def _split_total_rows(df, name_col):
    """Separate genuine state/UT rows from any 'Total' summary row, so national
    figures are summed from states only (never double-counted by a Total row)."""
    states, totals = [], []
    for r in df.to_dict(orient="records"):
        nm = str(r.get(name_col, "")).strip().lower()
        if not nm or nm == "nan":
            continue
        (totals if nm in _TOTAL_LABELS else states).append(r)
    return states, totals

@serialized
def fetch_pm_surya_ghar():
    # Cache key includes the xlsx mtime so user edits show up on the next request
    # (README promises live reads — a plain 6-hour TTL broke that promise)
    _ud = os.path.join(os.path.dirname(__file__), "user_data", "pm_surya_ghar.xlsx")
    _mt = int(os.path.getmtime(_ud)) if os.path.exists(_ud) else 0
    _ck = f"surya_ghar:{_mt}"
    cached = get_cache(_ck, 21600)
    if cached: return cached
    # Read from user-maintained Excel file if it exists
    if os.path.exists(_ud):
        try:
            df = pd.read_excel(_ud, engine="openpyxl")
            df.columns = [str(c).strip() for c in df.columns]
            # P18 schema: State / UT | Applications (No.) | Installations (No.) |
            # Households Covered (No.) | Installation Capacity (MW) | Subsidy Released (Cr)
            name_col = next((c for c in df.columns if c.lower().startswith("state")),
                            df.columns[0])
            colmap = [("Applications (No.)", "applications"),
                      ("Installations (No.)", "installations"),
                      ("Households Covered (No.)", "households"),
                      ("Installation Capacity (MW)", "capacity_mw"),
                      ("Subsidy Released (Cr)", "subsidy_cr")]
            present = [(s, d) for s, d in colmap if s in df.columns]
            states, _tot = _split_total_rows(df, name_col)
            def _remap(r):
                out = {"state": str(r.get(name_col, "")).strip()}
                for s, d in present:
                    out[d] = _clean_num(r.get(s))
                return out
            state_data = [_remap(r) for r in states]
            national = {d: round(sum((row.get(d) or 0) for row in state_data), 2)
                        for _s, d in present}
            result = {"national": national, "state_data": state_data,
                      "top_states": state_data, "columns": [d for _s, d in present],
                      "source": "user_data/pm_surya_ghar.xlsx",
                      "fetched_at": datetime.now().strftime("%Y-%m-%d %H:%M")}
            set_cache(_ck, result)
            mark_health("PM Surya Ghar", True, "user_data file (6-col schema)")
            return result
        except Exception as _e:
            mark_health("PM Surya Ghar", False, f"user_data read error: {_e}")
    try:
        # Use hash-route page — needs JS execution for state table
        html  = _playwright_fetch("https://pmsuryaghar.gov.in/", wait_selector=None)
        soup  = BeautifulSoup(html, "html.parser")
        text  = soup.get_text(" ", strip=True)

        # Extract national totals — look for large counter numbers
        national = {}
        for label, patterns in [
            ("applications", [r"([\d,]+)\s*(?:lakh\s*)?applications?\s*registered",
                              r"registered[^\d]*([\d,]+)"]),
            ("sanctioned",   [r"([\d,]+)\s*(?:lakh\s*)?sanctioned",
                              r"applications?\s+sanctioned[^\d]*([\d,]+)"]),
            ("installed",    [r"([\d,]+)\s*(?:lakh\s*)?installed",
                              r"installations?\s*completed[^\d]*([\d,]+)"]),
        ]:
            for pat in patterns:
                m = re.search(pat, text, re.I)
                if m:
                    national[label] = m.group(1).replace(",","")
                    break

        # Counter widgets (often Angular/React counters)
        counters = {}
        for el in soup.select(".counter,.count,[class*=number],[class*=stat],[class*=count],[class*=total]"):
            t2 = el.get_text(strip=True)
            if re.search(r"\d{4,}", t2): counters[len(counters)] = t2[:80]

        # State-wise table — find table with state names + numbers
        top_states = []
        for table in soup.find_all("table"):
            rows = table.find_all("tr")
            if len(rows) > 5:
                for row in rows[1:]:
                    cells = [c.get_text(strip=True) for c in row.find_all(["td","th"])]
                    if len(cells) >= 2 and re.search(r"[A-Za-z]", cells[0]) and re.search(r"\d", cells[-1]):
                        top_states.append({"state": cells[0][:30], "value": cells[-1][:20]})
                if top_states: break
        top_states = top_states[:10]

        result = {
            "national": national,
            "counters": counters,
            "top_states": top_states,
            "fetched_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        }
        set_cache(_ck, result)
        mark_health("PM Surya Ghar", True)
        return result
    except Exception as e:
        mark_health("PM Surya Ghar", False, str(e))
        return {"error": str(e), "note": "Playwright required — run: playwright install chromium"}

@serialized
def fetch_pm_kusum():
    # Cache key includes the xlsx mtime so user edits show up on the next request
    _ud = os.path.join(os.path.dirname(__file__), "user_data", "pm_kusum.xlsx")
    _mt = int(os.path.getmtime(_ud)) if os.path.exists(_ud) else 0
    _ck = f"pm_kusum:{_mt}"
    cached = get_cache(_ck, 21600)
    if cached: return cached
    # Read from user-maintained Excel file if it exists
    if os.path.exists(_ud):
        try:
            xf = pd.ExcelFile(_ud, engine="openpyxl")
            def _read_comp(sheet, val_cols):
                df = pd.read_excel(xf, sheet_name=sheet)
                df.columns = [str(c).strip() for c in df.columns]
                name_col = next((c for c in df.columns if c.lower().startswith("state")),
                                df.columns[0])
                # P18 fix: exclude any 'Total' row from BOTH state_data and the
                # national sum (it was double-counting national before).
                states, _tot = _split_total_rows(df, name_col)
                rows = [{k: (_clean_num(v) if k in val_cols else
                             (None if (hasattr(v, '__class__') and v.__class__.__name__ == 'NAType')
                              else v)) for k, v in r.items()} for r in states]
                nat = {col: round(sum((_clean_num(r.get(col)) or 0) for r in states), 2)
                       for col in val_cols}
                return {"state_data": rows, "national": nat}
            sheets = xf.sheet_names
            comp_a = _read_comp("PM_KUSUM_A" if "PM_KUSUM_A" in sheets else sheets[0],
                                ["Total_Sanction_MW",   "Total_Installed_MW"])
            comp_b = _read_comp("PM_KUSUM_B", ["Total_Sanction_Nos.", "Total_Installed_Nos."]) if "PM_KUSUM_B" in sheets else None
            comp_c = _read_comp("PM_KUSUM_C", ["Total_Sanction_Nos.", "Total_Installed_Nos."]) if "PM_KUSUM_C" in sheets else None
            result = {
                "state_data": comp_a["state_data"], "national": comp_a["national"],
                "comp_a": comp_a, "comp_b": comp_b, "comp_c": comp_c,
                "source": "user_data/pm_kusum.xlsx",
                "fetched_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
            }
            set_cache(_ck, result)
            mark_health("PM KUSUM", True, "user_data file — A/B/C loaded")
            return result
        except Exception as _e:
            mark_health("PM KUSUM", False, f"user_data read error: {_e}")
    try:
        html  = _playwright_fetch("https://pmkusum.mnre.gov.in/landing.html")
        soup  = BeautifulSoup(html, "html.parser")
        text  = soup.get_text(" ", strip=True)

        # Extract Component A / B / C structured numbers
        def _extract(patterns):
            for pat in patterns:
                m = re.search(pat, text, re.I)
                if m: return m.group(1).strip()
            return None

        comp_a_gw      = _extract([r"component[\s-]*a[^\d]*([\d,\.]+)\s*(?:GW|MW)", r"([\d,\.]+)\s*GW.*?ground.?mounted"])
        comp_b_pumps   = _extract([r"component[\s-]*b[^\d]*([\d,\.]+)\s*(?:lakh|million)?\s*pumps?",
                                   r"([\d,\.]+)\s*(?:lakh|M)\s*pumps?.*?standalone"])
        comp_c_pumps   = _extract([r"component[\s-]*c[^\d]*([\d,\.]+)\s*(?:lakh|million)?\s*pumps?",
                                   r"grid.?connected[^\d]*([\d,\.]+)\s*(?:lakh|M)?\s*pumps?"])
        budget_cr      = _extract([r"([\d,]+)\s*(?:crore|cr)",
                                   r"₹\s*([\d,]+)"])
        total_target   = _extract([r"total\s*target[^\d]*([\d,\.]+)\s*(GW|MW|MW\b)",
                                   r"([\d,\.]+)\s*GW\s*target"])

        # State-wise table if available
        state_data = []
        for table in soup.find_all("table"):
            rows = table.find_all("tr")
            if len(rows) > 5:
                for row in rows[1:]:
                    cells = [c.get_text(strip=True) for c in row.find_all(["td","th"])]
                    if len(cells) >= 3 and re.search(r"[A-Za-z]{3,}", cells[0]):
                        state_data.append({
                            "state": cells[0][:30],
                            "sanctioned": cells[1][:20] if len(cells) > 1 else "",
                            "installed":  cells[2][:20] if len(cells) > 2 else "",
                        })
                if state_data: break
        state_data = state_data[:20]

        # Fallback raw numbers
        raw_nums = re.findall(r"([\d,\.]+)\s*(GW|MW|pumps?|Lakh|crore|beneficiar\w*)", text, re.I)

        result = {
            "comp_a_gw":    comp_a_gw,
            "comp_b_pumps": comp_b_pumps,
            "comp_c_pumps": comp_c_pumps,
            "budget_cr":    budget_cr,
            "total_target": total_target,
            "state_data":   state_data,
            "raw_nums":     raw_nums[:15],
            "fetched_at":   datetime.now().strftime("%Y-%m-%d %H:%M"),
        }
        set_cache(_ck, result)
        mark_health("PM KUSUM", True)
        return result
    except Exception as e:
        mark_health("PM KUSUM", False, str(e))
        return {"error": str(e), "note": "Playwright required — run: playwright install chromium"}

# ── News Intelligence (P1.5) ─────────────────────────────────────────────────
_USEFULNESS_HIGH = [
    "tender","auction","bid","tariff","almm","bcd","anti-dumping","mnre","ntpc","seci",
    "nhpc","sjvn","mw","gw","commissioned","sanctioned","capacity","pli","vgf",
    "ireda","loan","finance","saatvik","waaree","adani","vikram solar","premier",
    "battery","bess","green hydrogen","electrolyzer","curtailment","grid","rpm","rpo",
    "solar panel","module","inverter","wind turbine","power purchase","ppa",
    "renewable energy","clean energy","energy storage","rooftop solar",
]
_USEFULNESS_LOW = [
    "opinion","sponsored","advertis","quiz","listicle","click here","subscribe",
    "newsletter","prize","giveaway","promoted","explainer","what is","recipes",
    "lifestyle","fashion","travel","sports","cricket","ipl","bollywood","film",
    "movie","celebrity","actor","actress","wedding","party","stock tip","horoscope",
    "health tip","beauty","diet","fitness","gaming","entertainment","music album",
]

# India-focused RE sources get a bonus — they are almost always relevant
_INDIA_RE_SOURCES = {"Mercom India","ET Energy","Solar Quarter","EQ Mag","REGlobal","SAU Energy"}

def _usefulness_score(title, summary, source):
    text = (title + " " + summary).lower()
    # Global wire sources: hard-gate — must contain at least one RE keyword
    if source in GLOBAL_NEWS_SOURCES:
        if not any(kw in text for kw in _GLOBAL_RE_REQUIRED):
            return 0  # discard immediately — not RE-related
        score = 35  # starts lower, wire sources need explicit RE content
    else:
        score = 30  # base for specialist sources
    for kw in _USEFULNESS_HIGH:
        if kw in text: score += 8
    for kw in _USEFULNESS_LOW:
        if kw in text: score -= 15  # stronger penalty
    # Bonus for India RE specialist sources
    if source in _INDIA_RE_SOURCES: score += 12
    return min(100, max(0, score))

def _parse_pub_dt(entry):
    """Return (datetime_obj, days_old_float). Falls back to now if unparseable."""
    tp = entry.get("published_parsed")
    if tp:
        try:
            import calendar, datetime as _dtmod
            _utc = _dtmod.timezone.utc
            dt_aware = _dtmod.datetime.fromtimestamp(calendar.timegm(tp), _utc)
            days_old = (_dtmod.datetime.now(_utc) - dt_aware).total_seconds() / 86400.0
            return dt_aware.replace(tzinfo=None), max(0.0, days_old)
        except Exception:
            pass
    # Undated article: penalize (3 days) instead of treating as fresh, so
    # recency-weighted scoring stops promoting undated junk to the top
    return datetime.now(), 3.0

def _archive_old_articles(articles):
    """Move articles older than 5 days to news_archive SQLite table (keep up to 15 days)."""
    con = sqlite3.connect(DB_PATH)
    try:
        for a in articles:
            if a["days_old"] > 5 and a["days_old"] <= 15:
                uid = a["link"] + a["source"]
                con.execute("""INSERT OR IGNORE INTO news_archive
                               (uid,source,title,link,summary,published_dt,days_old,ts)
                               VALUES(?,?,?,?,?,?,?,?)""",
                            (uid, a["source"], a["title"], a["link"],
                             a["summary"], a.get("published_dt",""), a["days_old"], time.time()))
        # Purge archive entries older than 15 days
        cutoff = time.time() - 15 * 86400
        con.execute("DELETE FROM news_archive WHERE ts < ?", (cutoff,))
        con.commit()
    finally:
        con.close()

def fetch_news():
    cached = get_cache("news", 1800)
    if cached: return cached
    articles = []
    ok_count = 0
    for source, url in RSS_FEEDS:
        try:
            feed = feedparser.parse(url)
            for e in feed.entries[:8]:
                dt, days_old = _parse_pub_dt(e)
                title   = e.get("title","")
                summary = (e.get("summary","") or "")[:300]
                score   = _usefulness_score(title, summary, source)
                if score < 15: continue  # discard pure noise
                articles.append({
                    "source":       source,
                    "source_type":  "global" if source in GLOBAL_NEWS_SOURCES else "india",
                    "title":        title,
                    "link":         e.get("link",""),
                    "date":         e.get("published",""),
                    "published_dt": dt.strftime("%Y-%m-%d %H:%M"),
                    "days_old":     round(days_old, 1),
                    "summary":      summary,
                    "score":        score,
                })
            ok_count += 1
        except Exception:
            pass
    articles.sort(key=lambda x: x["published_dt"], reverse=True)
    _archive_old_articles(articles)
    active = [a for a in articles if a["days_old"] <= 5]
    mark_health("RSS Feeds", ok_count >= 3, f"{ok_count}/{len(RSS_FEEDS)} feeds live")
    set_cache("news", active)
    return active

def fetch_news_archive():
    """Return 5-15 day old articles from SQLite archive."""
    con = sqlite3.connect(DB_PATH)
    try:
        rows = con.execute("""SELECT source,title,link,summary,published_dt,days_old
                              FROM news_archive ORDER BY published_dt DESC LIMIT 100""").fetchall()
        return [{"source":r[0],"title":r[1],"link":r[2],"summary":r[3],
                 "published_dt":r[4],"days_old":round(r[5],1)} for r in rows]
    finally:
        con.close()

def get_alerts():
    articles = fetch_news(); alerts = []
    for a in articles:
        text = (a["title"]+" "+a["summary"]).lower()
        for cat,kws in ALERT_KEYWORDS.items():
            matched = [kw for kw in kws if kw.lower() in text]
            if matched:
                uid      = a["link"]+cat
                is_new   = uid not in SEEN_ALERTS
                if is_new:
                    SEEN_ALERTS.add(uid)
                    persist_alert_db(uid, {**a,"category":cat,"keywords":matched})
                alerts.append({**a,"category":cat,"keywords":matched,"is_new":is_new})
                break
    return alerts

# ── Correlation Matrix ───────────────────────────────────────────────────────
def fetch_correlation():
    cached = get_cache("correlation", 3600)
    if cached: return cached
    try:
        # Use top RE stocks + macro proxies
        syms = [
            "ADANIGREEN.NS","NHPC.NS","SUZLON.NS","WAAREEENER.NS",
            "SAATVIKGL.NS","IREDA.NS","INOXWIND.NS","TATAPOWER.NS",
            "USDINR=X","CL=F","TAN","ICLN"
        ]
        name_map = {
            "ADANIGREEN.NS":"Adani Green","NHPC.NS":"NHPC","SUZLON.NS":"Suzlon",
            "WAAREEENER.NS":"Waaree","SAATVIKGL.NS":"Saatvik","IREDA.NS":"IREDA",
            "INOXWIND.NS":"Inox Wind","TATAPOWER.NS":"Tata Power",
            "USDINR=X":"USD/INR","CL=F":"Crude","TAN":"TAN ETF","ICLN":"ICLN ETF"
        }
        series = {}
        for sym in syms:
            try:
                h = yf.Ticker(sym).history(period="3mo")
                if not h.empty:
                    s = h["Close"].copy()
                    # Strip tz info so NSE (IST) and US (EST) dates align by date
                    s.index = pd.to_datetime([d.date() for d in s.index])
                    series[name_map.get(sym,sym)] = s
            except: pass
        if len(series)<3:
            return {"error":"insufficient data","matrix":[],"labels":[]}
        df   = pd.DataFrame(series).ffill().dropna()
        corr = df.corr().round(3).fillna(0)
        result = {
            "matrix": corr.values.tolist(),
            "labels": list(corr.columns),
            "fetched_at": datetime.now().strftime("%Y-%m-%d %H:%M")
        }
        set_cache("correlation", result)
        mark_health("Correlation", True)
        return result
    except Exception as e:
        mark_health("Correlation", False, str(e))
        return {"error":str(e),"matrix":[],"labels":[]}

# ── World Bank ───────────────────────────────────────────────────────────────
def india_energy_worldbank():
    cached = get_cache("wb_energy", 86400)
    if cached: return cached
    try:
        r = requests.get("https://api.worldbank.org/v2/country/IN/indicator/EG.ELC.RNEW.ZS?format=json&mrv=12",timeout=15)
        data = r.json()
        series = [{"year":d["date"],"value":d["value"]} for d in data[1] if d["value"]]
        set_cache("wb_energy", series)
        mark_health("World Bank", True)
        return series
    except Exception as e:
        mark_health("World Bank", False, str(e)); return []

# ── YouTube ───────────────────────────────────────────────────────────────────
def get_youtube_live(channel_id):
    """Resolve the channel's CURRENT live stream (v12 — worldmonitor pattern).
    The old version returned the latest RSS upload, i.e. a random recent clip.
    Now: scrape /channel/<id>/live for the live videoId; RSS-latest is only the
    last-resort fallback and is flagged is_live: false."""
    cached = get_cache(f"yt_{channel_id}", 300)
    if cached: return cached
    result = {"video_id": "", "title": "", "is_live": False}
    try:
        r = requests.get(f"https://www.youtube.com/channel/{channel_id}/live",
                         timeout=12, cookies={"CONSENT": "YES+1"},
                         headers={"User-Agent": "Mozilla/5.0 Chrome/124 Safari/537.36",
                                  "Accept-Language": "en"})
        html = r.text
        if '"isLive":true' in html or '"isLiveNow":true' in html:
            m = re.search(r'"videoId":"([\w-]{11})"', html)
            t = re.search(r'<title>([^<]+)</title>', html)
            if m:
                result = {"video_id": m.group(1),
                          "title": (t.group(1).replace(" - YouTube", "") if t else ""),
                          "is_live": True}
    except Exception:
        pass
    if not result["video_id"]:
        try:
            import xml.etree.ElementTree as ET
            r = requests.get(f"https://www.youtube.com/feeds/videos.xml?channel_id={channel_id}", timeout=10)
            ns = {"yt": "http://www.youtube.com/xml/schemas/2015", "atom": "http://www.w3.org/2005/Atom"}
            entries = ET.fromstring(r.text).findall("atom:entry", ns)
            if entries:
                vid = entries[0].find("yt:videoId", ns)
                title = entries[0].find("atom:title", ns)
                result = {"video_id": vid.text if vid is not None else "",
                          "title": title.text if title is not None else "", "is_live": False}
        except Exception:
            pass
    set_cache(f"yt_{channel_id}", result)
    return result

# ── Technicals ───────────────────────────────────────────────────────────────
def compute_projection(prices, days=90):
    if len(prices)<30: return []
    y      = np.array([p["close"] for p in prices[-90:]])
    x      = np.arange(len(y))
    coeffs = np.polyfit(x,y,2)
    fx     = np.arange(len(y), len(y)+days)
    fy     = np.polyval(coeffs, fx)
    last   = datetime.strptime(prices[-1]["date"],"%Y-%m-%d")
    return [{"date":(last+timedelta(days=i+1)).strftime("%Y-%m-%d"),"projected":round(float(v),2)} for i,v in enumerate(fy)]

def compute_technicals(prices):
    if len(prices)<26: return {}
    closes = np.array([p["close"] for p in prices])
    delta  = np.diff(closes)
    gain   = np.where(delta>0,delta,0); loss=np.where(delta<0,-delta,0)
    avg_g  = np.mean(gain[-14:]); avg_l=np.mean(loss[-14:])
    rsi    = 100-(100/(1+avg_g/avg_l)) if avg_l else 100
    ema12  = float(np.mean(closes[-12:])); ema26=float(np.mean(closes[-26:]))
    macd   = ema12-ema26; sig=float(np.mean(closes[-9:])-np.mean(closes[-18:]))
    sma20  = float(np.mean(closes[-20:])); std20=float(np.std(closes[-20:]))
    return {
        "rsi":round(float(rsi),2),"macd":round(macd,2),"macd_signal":round(sig,2),
        "bb_upper":round(sma20+2*std20,2),"bb_lower":round(sma20-2*std20,2),"bb_mid":round(sma20,2),
        "sma20":round(sma20,2),"sma50":round(float(np.mean(closes[-50:])),2) if len(closes)>=50 else None
    }

# ── Signal Intelligence (v4) ────────────────────────────────────────────────
# Severity weight per category — reflects impact on Indian RE sector
SIGNAL_SEVERITY = {
    "TRADE": 90, "SUPPLY CHAIN": 85, "POLICY": 75,
    "BESS/GH2": 65, "MARKET": 60, "COMPANIES": 50, "GLOBAL": 45,
}

def compute_signal_score():
    alerts = get_alerts()
    if not alerts:
        return {"score": 0, "level": "LOW", "active": 0, "composite": False,
                "categories": [], "signals": []}
    cats    = list(set(a["category"] for a in alerts))
    scores  = [min(SIGNAL_SEVERITY.get(a["category"], 40) + (15 if a.get("is_new") else 0), 100)
               for a in alerts]
    max_s   = max(scores)
    composite = len(cats) >= 3
    if   composite and max_s >= 75: level, score = "CRITICAL", min(100, max_s + 10)
    elif max_s >= 85 or (composite and max_s >= 65): level, score = "HIGH",     max_s
    elif max_s >= 65 or len(alerts) >= 5:            level, score = "MEDIUM",   max_s
    else:                                             level, score = "LOW",      max_s
    return {
        "score": round(score), "level": level,
        "active": len(alerts), "composite": composite, "categories": cats,
        "signals": [{"category": a["category"], "title": a["title"][:80],
                     "is_new": a.get("is_new", False)} for a in alerts[:5]],
    }

# ── Daily Intelligence Brief (v4) ────────────────────────────────────────────
def fetch_daily_brief():
    cached = get_cache("daily_brief", 300)
    if cached: return cached
    quotes  = fetch_all_quotes()
    news    = fetch_news()
    mnre    = fetch_mnre_live()
    tenders = fetch_seci_tenders()
    sig     = compute_signal_score()
    # Sector verdict from stock universe
    up    = sum(1 for q in quotes.values() if q.get("change_pct", 0) > 0)
    total = len(quotes)
    bull  = up / total * 100 if total else 0
    if   bull >= 65: verdict, vcol = "BULLISH",  "green"
    elif bull <= 35: verdict, vcol = "BEARISH",  "red"
    else:            verdict, vcol = "NEUTRAL",  "yellow"
    # Top mover by absolute % change
    top_mover = None
    if quotes:
        q = max(quotes.values(), key=lambda x: abs(x.get("change_pct", 0)))
        top_mover = {"name": q["name"], "pct": q["change_pct"], "price": q["price"]}
    # India total RE installed
    total_gw = 0
    if mnre and mnre.get("data"):
        total_gw = round(mnre["data"].get("Total RE", {}).get("cumulative_mw", 0) / 1000, 1)
    # Top tender
    top_tender = None
    if tenders and tenders.get("tenders"):
        t = tenders["tenders"][0]
        top_tender = {"title": t["title"][:70], "deadline": t.get("deadline", ""), "link": t.get("link", "")}
    result = {
        "date":          datetime.now().strftime("%d %b %Y"),
        "time":          datetime.now().strftime("%H:%M IST"),
        "verdict":       verdict,  "verdict_color": vcol,
        "bulls": up,     "bears":  total - up, "total_stocks": total,
        "top_mover":     top_mover,
        "total_re_gw":   total_gw,
        "mnre_as_on":    mnre.get("as_on", "") if mnre else "",
        "top_news":      {"title": news[0]["title"][:100], "source": news[0]["source"], "link": news[0]["link"]} if news else None,
        "top_tender":    top_tender,
        "signal_level":  sig["level"], "signal_score": sig["score"],
        "active_signals": sig["active"], "composite": sig["composite"],
        "signal_categories": sig["categories"],
    }
    set_cache("daily_brief", result)
    return result

# ── Power Market Widget (v6 — IEX has no public API, replaced with PRAAPTI) ──
def fetch_iex_power():
    cached = get_cache("iex_power", 1800)
    if cached: return cached
    comm = fetch_commodities()
    usdinr = comm.get("USDINR=X", {}).get("price")
    result = {
        "source": "PRAAPTI",
        "usdinr": usdinr,
        "praapti_url": "https://praapti.in/",
        "iex_url": "https://www.iexindia.com/",
        "fetched_at": datetime.now().strftime("%H:%M"),
    }
    set_cache("iex_power", result)
    return result

# ── Intel Helper Functions (v6) ──────────────────────────────────────────────
def classify_article(text):
    tl = text.lower()
    for cat, keywords in INTEL_CATEGORIES.items():
        for kw in keywords:
            if kw in tl:
                return cat, INTEL_DIRECTION.get(cat, "NEUTRAL")
    return None, "NEUTRAL"

def extract_entities(text):
    tl = text.lower()
    company = None
    for name, sym in COMPANY_ENTITIES.items():
        if name in tl:
            company = sym; break
    # (?!h) rejects energy units (MWh/GWh) — capacity only; GW converts to MW
    mw_m     = re.search(r"([\d,]+(?:\.\d+)?)\s*(mwp?|gw)(?!h)", tl)
    tariff_m = re.search(r"(?:₹|rs\.?)\s*([\d]+\.[\d]+)\s*/\s*kwh", tl)
    crore_m  = re.search(r"(?:₹|rs\.?)\s*([\d,]+)\s*(?:crore|cr\.?\b)", tl)
    _mw = None
    if mw_m:
        _mw = float(mw_m.group(1).replace(",",""))
        if mw_m.group(2) == "gw": _mw *= 1000
    return {
        "company": company,
        "mw":      _mw,
        "tariff":  float(tariff_m.group(1))                 if tariff_m else None,
        "crore":   float(crore_m.group(1).replace(",",""))  if crore_m  else None,
    }

def compute_pulse_score(intel_items):
    CWEIGHT = {
        "PROJECT_WIN":12,"COMMISSIONING":10,"EXPANSION":8,"FUNDING":7,
        "TENDER_ISSUED":6,"POLICY_NOTIFICATION":5,"M_AND_A":5,"EARNINGS_SIGNAL":4,
        "TARIFF_SIGNAL":4,"GLOBAL_MACRO":3,"SUPPLY_CHAIN":-8,"REGULATORY_RISK":-12,
    }
    DMULT = {"POSITIVE":1.0,"NEGATIVE":-1.0,"NEUTRAL":0.5}
    score = 40
    for item in intel_items:
        w = CWEIGHT.get(item.get("category",""), 2)
        d = DMULT.get(item.get("direction","NEUTRAL"), 0.5)
        b = 5 if item.get("entities",{}).get("company") else 0
        score += w * d + b
    return min(100, max(0, round(score)))

# ── Government Portal Scrapers (v6) ──────────────────────────────────────────
def _scrape_gov_tenders(url, source_name, cache_key, title_col=1, date_col=None):
    cached = get_cache(cache_key, 14400)
    if cached: return cached
    # Check 1hr failure cache — don't hammer unreachable hosts
    fail_cached = get_cache(cache_key + "_fail", 3600)
    if fail_cached: return fail_cached
    hdr   = {"User-Agent": "Mozilla/5.0 Chrome/124 Safari/537.36"}
    items = []
    try:
        r    = gov_get(url, timeout=8)
        soup = BeautifulSoup(r.text, "html.parser")
        table = soup.find("table")
        if not table:
            try:
                html  = _playwright_fetch(url)
                soup  = BeautifulSoup(html, "html.parser")
                table = soup.find("table")
            except Exception: pass
        if table:
            for row in table.find_all("tr")[1:16]:
                cells   = [c.get_text(strip=True) for c in row.find_all(["td","th"])]
                link_el = row.find("a", href=True)
                if len(cells) > title_col and cells[title_col]:
                    href = link_el["href"] if link_el else url
                    if href.startswith("/"): href = url.split("/")[0]+"//"+url.split("/")[2]+href
                    items.append({
                        "title":  cells[title_col][:120],
                        "date":   cells[date_col] if date_col is not None and len(cells)>date_col else "",
                        "link":   href,
                        "source": source_name,
                    })
    except requests.exceptions.ConnectionError as e:
        # DNS/network unreachable — cache failure for 1hr, return portal link
        fail_result = {"items":[],"error":"portal_unreachable","portal_url":url,
                       "source":source_name,"fetched_at":datetime.now().strftime("%Y-%m-%d %H:%M")}
        set_cache(cache_key + "_fail", fail_result)
        mark_health(source_name, False, f"unreachable: {str(e)[:80]}")
        return fail_result
    except Exception as e:
        mark_health(source_name, False, str(e))
    result = {"items":items[:12],"fetched_at":datetime.now().strftime("%Y-%m-%d %H:%M"),
              "source":source_name,"portal_url":url}
    if items:
        set_cache(cache_key, result)
        mark_health(source_name, True)
    return result

def fetch_mnre_notifications():
    cached = get_cache("mnre_notifs", 3600)
    if cached: return cached
    try:
        hdr  = {"User-Agent": "Mozilla/5.0 Chrome/124 Safari/537.36"}
        r    = gov_get("https://mnre.gov.in/notifications/", timeout=20)
        soup = BeautifulSoup(r.text, "html.parser")
        items = []
        rows  = soup.select("table tr") or soup.select("li") or []
        for row in rows[1:21]:
            cells   = [c.get_text(strip=True) for c in row.find_all(["td","th","p","span"])]
            link_el = row.find("a", href=True)
            if len(cells) >= 2 and cells[1]:
                href = link_el["href"] if link_el else ""
                if href and href.startswith("/"): href = "https://mnre.gov.in" + href
                items.append({"title":cells[1][:120],"date":cells[0],"link":href,"source":"MNRE"})
        result = {"items":items[:15],"fetched_at":datetime.now().strftime("%Y-%m-%d %H:%M")}
        set_cache("mnre_notifs", result)
        mark_health("MNRE Notifications", bool(items))
        return result
    except Exception as e:
        mark_health("MNRE Notifications", False, str(e))
        return {"error":str(e),"items":[]}

def fetch_ntpc_tenders():
    return _scrape_gov_tenders("https://www.ntpc.co.in/en/tenders","NTPC","ntpc_tenders",title_col=1,date_col=2)

def fetch_nhpc_tenders():
    # NHPC uses nhpcindia.com (new domain) — no dedicated tender page;
    # tenders issued via SECI/MNRE/GeM. Return portal info immediately.
    mark_health("NHPC", True, "no tender portal — tenders via SECI/MNRE")
    return {"items": [], "source": "NHPC", "note": "NHPC tenders issued via SECI/MNRE/GeM",
            "portal_url": "https://www.nhpcindia.com", "fetched_at": datetime.now().strftime("%Y-%m-%d %H:%M")}

def fetch_sjvn_tenders():
    return _scrape_gov_tenders("https://www.sjvn.nic.in/tenders-1","SJVN","sjvn_tenders")

# ── Company News via yfinance (v6) ────────────────────────────────────────────
def fetch_company_news():
    cached = get_cache("company_news", 1800)
    if cached: return cached
    result = {}
    for sym in list(RE_STOCKS.keys()):
        try:
            raw = yf.Ticker(sym).news or []
            result[sym] = [
                {
                    "title":     n.get("title",""),
                    "link":      n.get("link",""),
                    "publisher": n.get("publisher",""),
                    "published": datetime.fromtimestamp(n["providerPublishTime"]).strftime("%Y-%m-%d %H:%M")
                                 if n.get("providerPublishTime") else "",
                }
                for n in raw[:5]
            ]
        except Exception:
            result[sym] = []
    set_cache("company_news", result)
    mark_health("Company News", True)
    return result

# ── NEURON Intel Engine (v6) ──────────────────────────────────────────────────
def fetch_intel_engine():
    cached = get_cache("intel_engine", 1800)
    if cached: return cached

    # 1. Aggregate: RSS feeds + yfinance company news
    articles = []
    for a in fetch_news():
        articles.append({**a, "text": (a.get("title","")+" "+a.get("summary","")).lower()})
    for sym, items in fetch_company_news().items():
        for n in items:
            articles.append({
                "source":  RE_STOCKS.get(sym, sym),
                "title":   n.get("title",""),
                "link":    n.get("link",""),
                "date":    n.get("published",""),
                "summary": "",
                "text":    n.get("title","").lower(),
            })

    # 2. Classify + extract entities
    intel_items = []
    seen_links  = set()
    for a in articles:
        lnk = a.get("link","") or a.get("title","")
        if lnk in seen_links: continue
        seen_links.add(lnk)
        cat, direction = classify_article(a.get("text",""))
        if not cat: continue
        entities = extract_entities(a.get("text",""))
        score    = INTEL_CATEGORY_SCORE.get(cat, 40)
        if entities.get("mw"):      score += 10
        if entities.get("tariff"):  score += 8
        if entities.get("company"): score += 10
        intel_items.append({
            "category":  cat, "direction": direction,
            "title":     a.get("title","")[:100],
            "source":    a.get("source",""),
            "link":      a.get("link",""),
            "date":      a.get("date",""),
            "entities":  entities,
            "score":     score,
            "is_new":    lnk not in SEEN_ALERTS,
        })
    intel_items.sort(key=lambda x: x["score"], reverse=True)

    # P7 — Update term spike tracking + baseline
    _update_term_freq(articles)
    _update_baseline(articles)
    term_spikes = _compute_spikes()
    # Telegram alert on high-ratio spikes
    for spike in term_spikes[:2]:
        if spike["ratio"] >= 3.0:
            cd = _TERM_COOLDOWN.get(spike["term"], 0)
            if time.time() - cd > 1800:
                send_telegram(
                    f"🔍 <b>TERM SPIKE: {spike['term'].upper()}</b>\n"
                    f"{spike['count']}× in 2hr (baseline {spike['baseline']}× · ratio {spike['ratio']}×)\n"
                    f"Sources: {', '.join(spike['sources'][:3])}"
                )
                _TERM_COOLDOWN[spike["term"]] = time.time()

    # P7 — Sentiment counts
    bullish_count = sum(1 for i in intel_items if i.get("direction") == "POSITIVE")
    bearish_count = sum(1 for i in intel_items if i.get("direction") == "NEGATIVE")
    neutral_count = sum(1 for i in intel_items if i.get("direction") == "NEUTRAL")

    # 3. Per-company signals
    company_signals = {}
    for sym in RE_STOCKS:
        hits = [i for i in intel_items if i.get("entities",{}).get("company")==sym]
        if hits:
            top = hits[0]
            company_signals[sym] = {
                "news_count":        len(hits),
                "dominant_category": top["category"],
                "direction":         top["direction"],
                "top_headline":      top["title"],
                "score":             top["score"],
            }

    # 4. Action flags
    action_flags = []
    for item in intel_items[:10]:
        if item["score"] < 50: continue
        company = item.get("entities",{}).get("company","")
        cname   = RE_STOCKS.get(company,"") if company else ""
        mw_str  = f" — {int(item['entities']['mw'])}MW" if item.get("entities",{}).get("mw") else ""
        action_flags.append({
            "urgency":     "HIGH" if item["score"]>=75 else "MEDIUM",
            "flag":        f"{cname+': ' if cname else ''}{item['title'][:70]}{mw_str}",
            "category":    item["category"],
            "company_sym": company,
            "link":        item.get("link",""),
        })

    # 5. Tender pipeline (SECI + gov portals)
    tender_pipeline = []
    seci = fetch_seci_tenders()
    for t in (seci.get("tenders") or [])[:5]:
        tender_pipeline.append({"issuer":"SECI","title":t.get("title","")[:80],"deadline":t.get("deadline",""),"link":t.get("link","")})
    for fn, nm in [(fetch_ntpc_tenders,"NTPC"),(fetch_nhpc_tenders,"NHPC"),(fetch_sjvn_tenders,"SJVN")]:
        try:
            d = fn()
            for t in (d.get("items") or [])[:3]:
                tender_pipeline.append({"issuer":nm,"title":t.get("title","")[:80],"deadline":t.get("date",""),"link":t.get("link","")})
        except Exception: pass

    # 6. Hot topics (top 3 by article count)
    cat_counts = {}
    for item in intel_items:
        cat_counts[item["category"]] = cat_counts.get(item["category"],0)+1
    hot_topics = [c.replace("_"," ").title()
                  for c,_ in sorted(cat_counts.items(),key=lambda x:x[1],reverse=True)[:3]]

    pulse = compute_pulse_score(intel_items)

    # ── Cross-Source Signals (P6.2 — WorldMonitor pattern) ──────────────────
    # Detects composite events when multiple categories fire together
    cross_signals = []
    def _cs(sig_type, severity, summary, cats, score, companies=None):
        cross_signals.append({
            "type": sig_type, "severity": severity, "summary": summary,
            "contributing": cats, "score": score, "companies": companies or [],
        })
    cat_set = set(i["category"] for i in intel_items[:15])
    # COMPOSITE: tariff + capacity addition + tender all active
    if len({"TARIFF_SIGNAL","CAPACITY_ADDITION","TENDER_ISSUED"} & cat_set) >= 3:
        _cs("COMPOSITE_ESCALATION", "HIGH",
            "Simultaneous tariff pressure, capacity additions, and new tenders — full-cycle activity",
            list({"TARIFF_SIGNAL","CAPACITY_ADDITION","TENDER_ISSUED"} & cat_set), 82)
    # POLICY_SHIFT: 3+ policy items
    pol_count = sum(1 for i in intel_items if i["category"]=="POLICY_NOTIFICATION")
    if pol_count >= 3:
        _cs("POLICY_SHIFT", "HIGH",
            f"{pol_count} policy notifications in latest cycle — regulatory environment shifting",
            ["POLICY_NOTIFICATION"], 75 + pol_count*2)
    # FUNDING_SURGE: multiple funding/loan items
    fund_count = sum(1 for i in intel_items if i["category"]=="FUNDING")
    if fund_count >= 2:
        cos = list(set(i["entities"].get("company","") for i in intel_items if i["category"]=="FUNDING" and i["entities"].get("company")))[:4]
        _cs("FUNDING_SURGE", "MEDIUM",
            f"₹ capital deployment surge — {fund_count} financing events detected",
            ["FUNDING"], 65, cos)
    # MARKET_MOVE: high-score commissioning or capacity events
    high_cap = [i for i in intel_items if i["category"] in ("COMMISSIONING","CAPACITY_ADDITION") and i["score"]>=70]
    if len(high_cap) >= 2:
        _cs("CAPACITY_SURGE", "MEDIUM",
            f"{len(high_cap)} high-signal capacity events — sector commissioning acceleration",
            ["COMMISSIONING","CAPACITY_ADDITION"], 70)
    # TARIFF_ALERT: any high-score tariff
    tariff_high = [i for i in intel_items if i["category"]=="TARIFF_SIGNAL" and i["score"]>=75]
    if tariff_high:
        _cs("TARIFF_ALERT", "HIGH" if tariff_high[0]["score"]>=80 else "MEDIUM",
            tariff_high[0]["title"][:90],
            ["TARIFF_SIGNAL"], tariff_high[0]["score"])
    cross_signals.sort(key=lambda x: x["score"], reverse=True)

    result = {
        "industry_pulse":     pulse,
        "pulse_label":        "BULLISH" if pulse>=60 else "BEARISH" if pulse<=40 else "NEUTRAL",
        "hot_topics":         hot_topics,
        "action_flags":       action_flags[:5],
        "cross_signals":      cross_signals[:6],
        "company_signals":    company_signals,
        "intel_stream":       intel_items[:20],
        "tender_pipeline":    tender_pipeline[:10],
        "sources_live":       len(RSS_FEEDS),
        "articles_processed": len(articles),
        "intel_items":        len(intel_items),
        "refreshed_at":       datetime.now().strftime("%H:%M IST"),
        # P7 additions
        "term_spikes":        term_spikes,
        "bullish_count":      bullish_count,
        "bearish_count":      bearish_count,
        "neutral_count":      neutral_count,
    }
    # Store pulse in history (once per run, dedup by date)
    try:
        _con = sqlite3.connect(DB_PATH)
        today = datetime.now().strftime("%Y-%m-%d")
        exists = _con.execute("SELECT 1 FROM pulse_history WHERE date=?", (today,)).fetchone()
        if not exists:
            _con.execute("""INSERT INTO pulse_history
                            (ts,date,pulse,label,hot_topics,articles_processed,
                             bullish_count,bearish_count,neutral_count)
                            VALUES(?,?,?,?,?,?,?,?,?)""",
                         (time.time(), today, pulse, result["pulse_label"],
                          ",".join(hot_topics), len(articles),
                          bullish_count, bearish_count, neutral_count))
            _con.commit()
        _con.close()
    except Exception:
        pass

    # Telegram alert — fire only once per 4 hours if pulse crosses 80
    global _tg_last_intel_alert
    import time as _time
    if pulse >= 80 and (_time.time() - _tg_last_intel_alert) > 14400:
        top_flag = action_flags[0]["flag"] if action_flags else "No specific flag"
        send_telegram(
            f"⚡ <b>NEURON Alert</b>\n"
            f"Intel Pulse: <b>{pulse}</b> — {result['pulse_label']}\n"
            f"Top: {top_flag}\n"
            f"Hot: {', '.join(hot_topics)}"
        )
        _tg_last_intel_alert = _time.time()

    set_cache("intel_engine", result)
    return result

# ── India Macro Pulse — World Bank (v6) ───────────────────────────────────────
def fetch_india_macro():
    cached = get_cache("india_macro", 86400)
    if cached: return cached
    try:
        base = "https://api.worldbank.org/v2/country/IN/indicator/"
        out  = {}
        for key, ind in [("gdp_growth","NY.GDP.MKTP.KD.ZG"),("cpi","FP.CPI.TOTL.ZG")]:
            r    = requests.get(f"{base}{ind}?format=json&mrv=6", timeout=15)
            data = r.json()
            out[key] = [{"year":d["date"],"value":round(float(d["value"]),2)} for d in data[1] if d["value"]]
        out["fetched_at"] = datetime.now().strftime("%Y-%m-%d")
        set_cache("india_macro", out)
        mark_health("India Macro", True)
        return out
    except Exception as e:
        mark_health("India Macro", False, str(e)); return {"error":str(e)}

# ── Sector Performance Overlay (v6) ──────────────────────────────────────────
SECTOR_OVERLAY_STOCKS = [
    ("SAATVIKGL.NS","Saatvik ★"),("ADANIGREEN.NS","Adani Green"),
    ("NHPC.NS","NHPC"),("SUZLON.NS","Suzlon"),
    ("WAAREEENER.NS","Waaree"),("NTPC.NS","NTPC"),
]

def fetch_sector_history(period="1mo"):
    key = f"sector_hist_{period}"
    cached = get_cache(key, 3600)
    if cached: return cached
    result = {}
    for sym, name in SECTOR_OVERLAY_STOCKS:
        try:
            h = yf.Ticker(sym).history(period=period)
            if h.empty: continue
            closes = [round(float(v),2) for v in h["Close"].values]
            base   = closes[0] or 1
            result[sym] = {
                "name":       name,
                "dates":      [str(d.date()) for d in h.index],
                "normalized": [round(v/base*100,2) for v in closes],
            }
        except Exception: pass
    out = {"stocks":result,"period":period,"fetched_at":datetime.now().strftime("%Y-%m-%d %H:%M")}
    set_cache(key, out)
    mark_health("Sector Overlay", bool(result))
    return out

# ── Saatvik Focus Stock Deep-Dive (v6) ────────────────────────────────────────
def fetch_focus_stock_deep():
    cached = get_cache("focus_deep", 3600)
    if cached: return cached
    try:
        t      = yf.Ticker("SAATVIKGL.NS")
        fi     = t.fast_info
        w52h   = getattr(fi,"year_high",None)
        w52l   = getattr(fi,"year_low",None)
        price  = getattr(fi,"last_price",None)
        mktcap = getattr(fi,"market_cap",None)
        pe, promoter_pct = None, None
        try:
            info = t.info
            pe   = info.get("trailingPE") or info.get("forwardPE")
        except Exception: pass
        try:
            mh = t.major_holders
            if mh is not None and not mh.empty:
                for _, row in mh.iterrows():
                    desc = str(row.iloc[1]).lower()
                    if "insider" in desc or "promoter" in desc:
                        try: promoter_pct = round(float(str(row.iloc[0]).replace("%","").strip()),1)
                        except Exception: pass
        except Exception: pass
        pos_pct = None
        if w52h and w52l and price:
            rng = w52h - w52l
            if rng > 0: pos_pct = round((price-w52l)/rng*100, 1)
        result = {
            "symbol":"SAATVIKGL.NS","name":"Saatvik Green Energy",
            "price":        round(price,2)      if price   else None,
            "mktcap_cr":    round(mktcap/1e7,0) if mktcap  else None,
            "week52_high":  round(w52h,2)        if w52h    else None,
            "week52_low":   round(w52l,2)        if w52l    else None,
            "pos_in_range_pct": pos_pct,
            "pe":           round(pe,1)           if pe      else None,
            "promoter_pct": promoter_pct,
            "fetched_at":   datetime.now().strftime("%H:%M"),
        }
        set_cache("focus_deep", result)
        mark_health("Focus Stock", True)
        return result
    except Exception as e:
        mark_health("Focus Stock", False, str(e)); return {"error":str(e)}

# ── Supply Chain Stress — PV Module Prices (v6) ───────────────────────────────
def fetch_pv_prices():
    cached = get_cache("pv_prices", 86400)
    if cached: return cached
    try:
        hdr  = {"User-Agent":"Mozilla/5.0 Chrome/124 Safari/537.36"}
        r    = gov_get("https://pvinsights.com/", timeout=20)
        soup = BeautifulSoup(r.text, "html.parser")
        prices = {}
        for row in soup.select("tr"):
            cells = [c.get_text(strip=True) for c in row.find_all(["td","th"])]
            if len(cells) < 2: continue
            label = cells[0].lower()
            for prod in ["polysilicon","wafer","cell","module"]:
                if prod in label:
                    for cell in cells[1:]:
                        m = re.search(r"(\d+\.?\d*)", cell.replace(",",""))
                        if m: prices[prod] = {"price":float(m.group(1)),"raw":cell}; break
        mod_p  = prices.get("module",{}).get("price",0)
        stress = "HIGH" if mod_p>0.28 else "MEDIUM" if mod_p>0.22 else "LOW" if mod_p>0 else "UNKNOWN"
        result = {"prices":prices,"stress":stress,"module_price":mod_p,
                  "fetched_at":datetime.now().strftime("%Y-%m-%d"),"source":"pvinsights.com"}
        set_cache("pv_prices", result)
        mark_health("PV Prices", bool(prices))
        return result
    except Exception as e:
        mark_health("PV Prices", False, str(e))
        return {"error":str(e),"prices":{},"stress":"UNKNOWN"}

# ── SECI Auction Results (v6) ─────────────────────────────────────────────────
def fetch_seci_results():
    cached = get_cache("seci_results", 14400)
    if cached: return cached
    SECI_COS = ["Adani","ReNew","NTPC","NHPC","SJVN","Torrent","Tata","ACME",
                "Greenko","Waaree","Saatvik","Sterling","Amp","Avaada","Hero"]
    try:
        hdr = {"User-Agent":"Mozilla/5.0 Chrome/124 Safari/537.36"}
        r   = gov_get("https://seci.co.in/loa/", timeout=20)
        soup = BeautifulSoup(r.text,"html.parser")
        results = []
        for row in soup.select("tr"):
            cells   = [c.get_text(" ",strip=True) for c in row.find_all(["td","th"])]
            if not cells: continue
            text    = " ".join(cells)
            t_m     = re.search(r"(?:₹|rs\.?)\s*([\d]+\.[\d]+)\s*/\s*kwh", text, re.I)
            mw_m    = re.search(r"([\d,]+)\s*(?:mwp?|gw)", text, re.I)
            link_el = row.find("a", href=True)
            if t_m or mw_m:
                results.append({
                    "text":      text[:120],
                    "tariff":    float(t_m.group(1))                   if t_m  else None,
                    "mw":        int(mw_m.group(1).replace(",",""))    if mw_m else None,
                    "companies": [c for c in SECI_COS if c.lower() in text.lower()],
                    "link":      link_el["href"] if link_el else "",
                })
        out = {"results":results[:15],"fetched_at":datetime.now().strftime("%Y-%m-%d %H:%M"),
               "source":"seci.co.in/loa","count":len(results)}
        set_cache("seci_results", out)
        mark_health("SECI Results", True)
        return out
    except Exception as e:
        mark_health("SECI Results", False, str(e))
        return {"error":str(e),"results":[]}

# ── CEA Daily RE Generation (v6) ─────────────────────────────────────────────
def fetch_cea_generation():
    cached = get_cache("cea_gen", 86400)
    if cached: return cached
    try:
        hdr  = {"User-Agent":"Mozilla/5.0 Chrome/124 Safari/537.36"}
        r    = gov_get("https://gen-re.cea.gov.in/", timeout=20)
        soup = BeautifulSoup(r.text, "html.parser")
        text = soup.get_text()
        data = {}
        for table in soup.find_all("table"):
            for row in table.find_all("tr"):
                cells = [c.get_text(strip=True) for c in row.find_all(["td","th"])]
                if len(cells) < 2: continue
                lbl = cells[0].lower()
                for tech in ["solar","wind","hydro","total"]:
                    if tech in lbl:
                        nums = []
                        for c in cells[1:]:
                            m = re.search(r"([\d,\.]+)", c)
                            if m:
                                try: nums.append(float(m.group(1).replace(",","")))
                                except: pass
                        if nums: data[tech] = {"generation_mu": nums[0]}
        for tech, pat in [("solar",r"[Ss]olar[^\d]{0,30}([\d,\.]+)\s*(?:MU|mu|MWh)"),
                          ("wind", r"[Ww]ind[^\d]{0,30}([\d,\.]+)\s*(?:MU|mu|MWh)")]:
            if tech not in data:
                m = re.search(pat, text)
                if m:
                    try: data[tech] = {"generation_mu": float(m.group(1).replace(",",""))}
                    except: pass
        result = {"data":data,"fetched_at":datetime.now().strftime("%Y-%m-%d %H:%M"),"source":"gen-re.cea.gov.in"}
        set_cache("cea_gen", result)
        mark_health("CEA Generation", bool(data))
        return result
    except Exception as e:
        mark_health("CEA Generation", False, str(e))
        return {"error":str(e),"data":{}}

# ── Watchlist API ────────────────────────────────────────────────────────────
def save_watchlist(wl):
    with open(WATCHLIST_PATH,"w") as f:
        json.dump(wl, f, indent=2)

# ── Flask Routes ─────────────────────────────────────────────────────────────
@app.route("/")
@app.route("/v20")
def index_v20():
    """v20 — the rebuilt premium responsive cockpit."""
    return render_template("index20.html")

@app.route("/v19")
@app.route("/cockpit")
def index_v19():
    """P19 — the rebuilt 'Cockpit of a Mind' is now the default at /. The full
    legacy app is preserved at /legacy (every endpoint, panel and feature intact)."""
    return render_template("index19.html")

@app.route("/legacy")
@app.route("/v18")
def index_legacy():
    """The complete pre-P19 cockpit — preserved, nothing removed."""
    return render_template("index.html")

@app.route("/api/quotes")
def api_quotes(): return jsonify(fetch_all_quotes())

@app.route("/api/news")
def api_news(): return jsonify(fetch_news())

@app.route("/api/commodities")
def api_commodities(): return jsonify(fetch_commodities())

@app.route("/api/energy_prices")
def api_energy_prices(): return jsonify(fetch_energy_prices())

@app.route("/api/global_re")
def api_global_re(): return jsonify(fetch_global_re())

@app.route("/api/global_capacity")
def api_global_capacity(): return jsonify(fetch_global_installed_capacity())

@app.route("/api/solar_capacity_history")
def api_solar_capacity_history(): return jsonify(fetch_solar_capacity_history())

@app.route("/api/wind_tech_mix")
def api_wind_tech_mix(): return jsonify(fetch_wind_tech_mix())

@app.route("/api/mnre_live")
def api_mnre_live(): return jsonify(fetch_mnre_live())

# ── v11 Observatory: source registry & regional article store ────────────────
@app.route("/api/sources/stats")
def api_sources_stats(): return jsonify(v11_sources.source_stats())

@app.route("/api/news/region/<region>")
def api_news_region(region):
    if region not in v11_sources.REGIONS:
        return jsonify({"error": f"unknown region; one of {v11_sources.REGIONS}"}), 404
    return jsonify(v11_sources.recent_articles(region=region, hours=48, limit=60))

@app.route("/api/region_velocity")
def api_region_velocity(): return jsonify(v11_sources.region_velocity())

# ── P14 Item 8: living-memory entity pipeline ─────────────────────────────────
@app.route("/api/pipeline")
@app.route("/api/pipeline/<query>")
def api_pipeline(query=None):
    return jsonify({"entities": v11_sources.entity_pipeline(query),
                    "stats": v11_sources.entity_ledger_stats(),
                    "query": query})

# ── v11 Observatory: intelligence layer (LLM-enhanced, heuristic-guaranteed) ──
import intelligence as v11_intel
# ── v15 Cognition: belief/diff/attention/consolidation (DB-only think layer) ──
import cognition as v11_cog
# ── v16 MemoryOS: curation + dual-hierarchy + multi-tier recall (DB-only) ─────
import memory as v16_mem
# ── v17 Executive Function: fuse all faculties → ranked, scored decisions ─────
import decisions as v17_dec

def _decision_context():
    """Gather the market-derived faculties (yfinance fetchers live here, in the
    expression layer) for the decision engine. decisions.py reads the DB-only
    faculties (beliefs/attention/chokepoints/lead-lag) itself."""
    ctx = {}
    for k, fn in (("implications", fetch_re_implications), ("regime", fetch_re_regime),
                  ("forecast", fetch_re_forecast), ("fear_greed", compute_fear_greed)):
        try: ctx[k] = fn()
        except Exception: ctx[k] = {}
    try:
        ctx["prices"] = {s: (v.get("price") if isinstance(v, dict) else None)
                         for s, v in fetch_all_quotes().items()}
    except Exception:
        ctx["prices"] = {}
    return ctx

# ── P15 B1/B2/B3 — Cognition routes (temporal awareness + belief revision) ────
@app.route("/api/delta/today")
def api_delta_today():
    """What changed since yesterday — the hippocampus' consolidation output."""
    return jsonify(v11_cog.get_today_delta())

@app.route("/api/delta/run", methods=["POST", "GET"])
def api_delta_run():
    """Force a consolidation pass now (manual sleep cycle)."""
    return jsonify(v11_cog.run_consolidation(force=True))

@app.route("/api/beliefs")
def api_beliefs():
    """Neuron's current understanding of key metrics, with conflicts surfaced."""
    return jsonify(v11_cog.beliefs_view())

@app.route("/api/attention")
def api_attention():
    """Unusualness flags — what the brain should attend to right now."""
    return jsonify(v11_cog.compute_attention())

@app.route("/api/self_test")
def api_self_test():
    """C2 — on-demand diagnostics: in-process invariant suite, structured result."""
    return jsonify(v11_cog.self_test())

# ── P16 MemoryOS — living memory: curation, dual-hierarchy, unified recall ────
@app.route("/api/memory/recall")
def api_memory_recall():
    """The 'what do I know about X' entrypoint — semantic+keyword+temporal fusion."""
    q = request.args.get("q", "")
    k = min(int(request.args.get("k", 8) or 8), 30)
    when = request.args.get("when") or None
    scope = request.args.get("scope", "neuron")
    return jsonify(v16_mem.recall(q, k=k, when=when, scope=scope))

@app.route("/api/memory/stats")
def api_memory_stats():
    return jsonify(v16_mem.memory_stats(request.args.get("scope")))

@app.route("/api/memory/add", methods=["POST"])
def api_memory_add():
    body = request.get_json(silent=True) or {}
    text = body.get("text", "")
    if not text:
        return jsonify({"ok": False, "error": "text required"}), 400
    res = v16_mem.add_note(text, body.get("source", "owner"), body.get("scope", "neuron"))
    return jsonify(res), (200 if res.get("ok") else 400)

# ── P15 A3 — Entity correction (editable, auditable living memory) ────────────
@app.route("/api/pipeline/entity/<entity_id>", methods=["DELETE"])
def api_pipeline_entity_delete(entity_id):
    reason = (request.get_json(silent=True) or {}).get("reason", "") if request.data else ""
    return jsonify(v11_sources.delete_entity(entity_id, reason))

@app.route("/api/pipeline/entity/<entity_id>", methods=["PATCH"])
def api_pipeline_entity_patch(entity_id):
    body = request.get_json(silent=True) or {}
    fields = body.get("fields", {k: v for k, v in body.items() if k != "reason"})
    if not fields:
        return jsonify({"ok": False, "error": "no fields to patch"}), 400
    res = v11_sources.patch_entity(entity_id, fields, body.get("reason", ""))
    return jsonify(res), (200 if res.get("ok") else 400)

@app.route("/api/intel/early_signals")
def api_early_signals(): return jsonify(v11_intel.early_signals())

@app.route("/api/intel/novelty")
def api_novelty(): return jsonify(v11_intel.novelty_radar())

@app.route("/api/chokepoints")
def api_chokepoints():
    """P16.4 — maritime chokepoint stress + India energy-import exposure."""
    return jsonify(v11_intel.chokepoint_monitor())

@app.route("/api/decisions")
def api_decisions():
    """P17 — executive function: ranked, conviction-scored, falsifiable decisions
    fused across every faculty. ?narrative=1 adds an LLM executive read."""
    narrative = request.args.get("narrative") == "1"
    return jsonify(v17_dec.synthesize_decisions(_decision_context(), narrative=narrative))

@app.route("/api/decisions/scorecard")
def api_decisions_scorecard():
    """P17 — Neuron's self-calibration: hit-rate by conviction band over time."""
    return jsonify(v17_dec.decision_scorecard())

@app.route("/api/deep_read", methods=["POST"])
def api_deep_read():
    """P19.5 — Deep-Read Agent: one article URL → Top-1% analyst one-pager."""
    body = request.get_json(silent=True) or {}
    url = body.get("url", "")
    if not url:
        return jsonify({"ok": False, "error": "url required"}), 400
    system_prompt = body.get("system_prompt")
    return jsonify(v11_intel.deep_read(url, system_prompt=system_prompt))

def _tg_esc(s):
    return str(s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

def push_strong_decisions(max_push=3):
    """P20 — proactively Telegram-push ONLY STRONG, cross-corroborated, FRESH
    decisions (each with its falsifier + honest track-record note). Hard-gated to
    avoid noise; deduped via kv so a standing call is pushed once; degrades
    silently when Telegram is unconfigured."""
    if not TELEGRAM_BOT_TOKEN:
        return {"pushed": 0, "configured": False}
    try:
        ctx = _decision_context()
        res = v17_dec.synthesize_decisions(ctx, cite=True)
        sc = v17_dec.decision_scorecard()
    except Exception as e:
        return {"pushed": 0, "error": str(e)[:80]}
    try:
        seen = set(json.loads(v11_sources.kv_get("pushed_decision_keys") or "[]"))
    except Exception:
        seen = set()
    strong = [d for d in res.get("decisions", [])
              if d.get("band") == "STRONG" and d.get("corroboration", 0) >= 2
              and d.get("key") not in seen]
    resolved = sum((c.get("CONFIRMED", 0) + c.get("INVALIDATED", 0))
                   for c in (sc.get("calibration_by_band") or {}).values())
    track = "unproven — no calls resolved yet" if resolved == 0 else f"{resolved} calls resolved to date"
    pushed = []
    for d in strong[:max_push]:
        cite = (d.get("supporting_facts") or [{}])[0].get("text", "")
        msg = (f"🜂 <b>NEURON — STRONG call</b> ({d.get('conviction')}/100 · "
               f"{d.get('corroboration')} faculties agree)\n\n<b>{_tg_esc(d.get('thesis',''))}</b>\n\n"
               f"Action: {_tg_esc(d.get('action',''))}"
               + (f" · {_tg_esc(d.get('ticker'))} {_tg_esc(d.get('direction',''))}" if d.get("ticker") else "")
               + f" · {d.get('horizon_days',0)}d\nBacked by: {_tg_esc(', '.join(d.get('faculties', [])))}\n"
               f"⚠ Wrong if: {_tg_esc(d.get('falsifier',''))}"
               + (f"\n\n📎 {_tg_esc(cite[:160])}" if cite else "")
               + f"\n\n<i>Self-scored conviction · {track}</i>")
        if send_telegram(msg):
            pushed.append(d.get("key")); seen.add(d.get("key"))
    if pushed:
        v11_sources.kv_set("pushed_decision_keys", json.dumps(list(seen)[-120:]))
    return {"pushed": len(pushed), "candidates": len(strong), "configured": True, "keys": pushed}

@app.route("/api/decisions/push", methods=["POST"])
def api_decisions_push():
    """P20 — fire the proactive STRONG-decision Telegram push (deduped, gated)."""
    return jsonify(push_strong_decisions())

@app.route("/api/intel/synthesis")
def api_synthesis():
    force = request.args.get("refresh") == "1"
    return jsonify(v11_intel.synthesis_brief(force=force))

@app.route("/api/intel/standing")
def api_standing():
    force = request.args.get("refresh") == "1"
    return jsonify(v11_intel.standing_questions(force=force))

# ── v12: stories, archive search, ask, state briefs, SSE ─────────────────────
@app.route("/api/stories")
def api_stories(): return jsonify(v11_intel.cluster_stories())

@app.route("/api/archive/search")
def api_archive_search():
    q = request.args.get("q", "")
    return jsonify({"q": q, "results": v11_intel.archive_search(q)})

@app.route("/api/ask")
def api_ask():
    return jsonify(v11_intel.ask_neuron(request.args.get("q", "")))

def _canon_state(s):
    s = re.sub(r"[^a-z]", "", (s or "").lower().replace("&", "and"))
    return (s.replace("odisha", "orissa").replace("uttarakhand", "uttaranchal")
             .replace("andamanandnicobarislands", "andamanandnicobar"))

@app.route("/api/state/<name>")
def api_state_brief(name):
    name = name[:60]
    key = _canon_state(name)
    cap = fetch_mnre_state_capacity() or {}
    cap_row, cap_name = None, name
    for st, v in (cap.get("states") or {}).items():
        if _canon_state(st) == key or key in _canon_state(st) or _canon_state(st) in key:
            cap_row, cap_name = v, st
            break
    arts = v11_intel.archive_search(name, limit=12)
    tenders = [t for t in (fetch_seci_tenders() or {}).get("tenders", [])
               if name.lower() in (t.get("title", "") or "").lower()][:5]
    # regulatory/grid chatter from the per-state SERC/DISCOM standing sources
    reg = [a for a in v11_sources.recent_articles(region="india", hours=24*7, limit=400)
           if (a["source_id"].startswith(("in_serc_", "in_discom_", "in_state_"))
               and key.startswith(_canon_state(a["source_id"].split("_", 2)[-1])[:6]))][:8]
    return jsonify({"state": cap_name, "capacity": cap_row,
                    "as_on": cap.get("as_on", ""),
                    "articles": arts, "tenders": tenders, "regulatory": reg,
                    "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M")})

@app.route("/api/stream")
def api_stream():
    """Lean SSE: pushes network stats every 45 s — sources pill updates without polling."""
    def gen():
        import json as _json
        while True:
            try:
                s = v11_sources.source_stats()
                healthy = sum(v.get("healthy", 0) for v in s["health_by_region"].values())
                a24 = sum(s["articles_24h_by_region"].values())
                payload = _json.dumps({"healthy": healthy, "total": s["registry"]["total"],
                                       "articles_24h": a24, "ts": time.time()})
                yield f"data: {payload}\n\n"
            except Exception:
                yield "data: {}\n\n"
            time.sleep(45)
    return app.response_class(gen(), mimetype="text/event-stream",
                              headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

# RE-hub generation weather — Open-Meteo, keyless, free
RE_HUBS = [
    ("Jaisalmer (RJ solar/wind)", 26.92, 70.90),
    ("Bhuj–Kutch (GJ hybrid)",    23.24, 69.67),
    ("Pavagada (KA solar)",       14.10, 77.27),
    ("Tuticorin (TN wind)",        8.76, 78.13),
    ("Anantapur (AP solar)",      14.68, 77.60),
]

def _owm_current(lat, lon):
    """OpenWeatherMap current conditions for an RE hub. None if no key / failure
    (degrade-never-break). Cloud cover is a direct near-term solar-yield signal."""
    if not OPENWEATHER_API_KEY:
        return None
    try:
        r = requests.get("https://api.openweathermap.org/data/2.5/weather", timeout=10,
                         params={"lat": lat, "lon": lon, "appid": OPENWEATHER_API_KEY,
                                 "units": "metric"})
        if r.status_code != 200:
            return None
        j = r.json()
        return {"clouds_pct": (j.get("clouds") or {}).get("all"),
                "temp_c": (j.get("main") or {}).get("temp"),
                "wind_ms": (j.get("wind") or {}).get("speed"),
                "desc": (j.get("weather") or [{}])[0].get("description", "")}
    except Exception:
        return None

@app.route("/api/observatory/weather")
def api_obs_weather():
    cached = get_cache("obs_weather", 6 * 3600)
    if cached: return jsonify(cached)
    from concurrent.futures import ThreadPoolExecutor, as_completed

    def _fetch_hub(name, lat, lon):
        try:
            r = requests.get("https://api.open-meteo.com/v1/forecast", timeout=30, params={
                "latitude": lat, "longitude": lon, "forecast_days": 7,
                "daily": "shortwave_radiation_sum,windspeed_10m_max",
                "timezone": "Asia/Kolkata"})
            d = r.json().get("daily", {})
            rad = d.get("shortwave_radiation_sum") or []
            wind = d.get("windspeed_10m_max") or []
            return {"hub": name, "dates": d.get("time") or [],
                    "radiation_mj_m2": rad, "wind_max_kmh": wind,
                    "rad_avg": round(sum(rad)/len(rad), 1) if rad else None,
                    "wind_avg": round(sum(wind)/len(wind), 1) if wind else None,
                    "current": _owm_current(lat, lon)}
        except Exception as e:
            return {"hub": name, "error": str(e)[:80]}

    # Fetch all 5 hubs in parallel (reduces cold-fetch time from ~5×30s to ~30s)
    futures = {}
    with ThreadPoolExecutor(max_workers=5) as ex:
        for name, lat, lon in RE_HUBS:
            futures[ex.submit(_fetch_hub, name, lat, lon)] = name
    hubs = [f.result() for f in sorted(futures, key=lambda f: [n for n,_,_ in RE_HUBS].index(futures[f]))]

    result = {"hubs": hubs,
              "source": "open-meteo.com (forecast) + openweathermap.org (current)"
                        if OPENWEATHER_API_KEY else "open-meteo.com (free, keyless)",
              "fetched_at": datetime.now().strftime("%Y-%m-%d %H:%M")}
    set_cache("obs_weather", result)
    return jsonify(result)

# ── P19 ALCM Ownership Atlas — true corporate structure of India's PV cell mfg ─
# Loads the authoritative artifacts produced by ALCM/run_alcm_atlas.py (group
# mapping + cross-linkages + history). mtime-keyed: re-running the script auto-
# refreshes Neuron. Group-level consolidation + HHI/CR4 concentration — the
# analysis the naive entity-level ALMM parse got wrong.
ALCM_DIR = os.path.join(os.path.dirname(__file__), "ALCM")

def _alcm_mtime():
    if not os.path.isdir(ALCM_DIR):
        return 0
    try:
        return int(max((os.path.getmtime(os.path.join(ALCM_DIR, f))
                        for f in os.listdir(ALCM_DIR) if f.endswith(".csv")), default=0))
    except Exception:
        return 0

def _alcm_seed_memory(result):
    """Persist the ownership intelligence into MemoryOS so recall() can surface
    'who controls India's cell manufacturing' (owner request — knowledge must
    survive into the new UI). Idempotent per file-revision."""
    flag = f"alcm_mem_seeded:{_alcm_mtime()}"
    if v11_sources.kv_get(flag):
        return
    try:
        m = result["summary"]
        v16_mem.add_note(
            f"ALCM cell-manufacturing concentration: {m['groups']} business groups control "
            f"{m['total_capacity_gw']} GW of MNRE-approved solar cell capacity; HHI {m['hhi']}, "
            f"CR4 {m['cr4']}%, top-5 groups {m['top5_share']}% — group-level, not entity-level.",
            source="ALCM atlas")
        for g in result["groups"][:13]:
            ents = ", ".join(e["name"] for e in g["entities"])
            v16_mem.add_note(
                f"ALCM ownership: {g['group']} controls {g['capacity_mw']:.0f} MW "
                f"({g['share_pct']}%) of India's approved solar cell capacity via {ents} — "
                f"treat as one account.", source="ALCM atlas")
        v11_sources.kv_set(flag, "1")
    except Exception:
        pass

def fetch_alcm_atlas():
    import csv
    ck = f"alcm_atlas:{_alcm_mtime()}"
    cached = get_cache(ck, 21600)
    if cached:
        return cached

    def _read(name):
        p = os.path.join(ALCM_DIR, name)
        if not os.path.exists(p):
            return []
        try:
            with open(p, encoding="utf-8") as fh:
                return list(csv.DictReader(fh))
        except Exception:
            return []

    gm = _read("alcm_group_mapping.csv")
    if not gm:
        mark_health("ALCM Atlas", False, "no group_mapping.csv")
        return {"available": False,
                "note": "ALCM atlas not generated — run ALCM/run_alcm_atlas.py to build it."}

    groups = {}
    for r in gm:
        g = (r.get("ultimate_group") or "Unknown").strip()
        try: cap = float(r.get("capacity_val") or 0)
        except (TypeError, ValueError): cap = 0.0
        e = groups.setdefault(g, {"group": g, "capacity_mw": 0.0, "entities": [],
                                  "_locs": set(), "_tech": set()})
        e["capacity_mw"] += cap
        e["entities"].append({"name": r.get("manufacturer_normalized"), "capacity_mw": cap,
                              "relationship": r.get("relationship_type"),
                              "evidence": r.get("evidence"),
                              "technologies": r.get("technologies"), "location": r.get("locations")})
        if r.get("locations"): e["_locs"].add(r["locations"].strip())
        for t in (r.get("technologies") or "").split(","):
            if t.strip(): e["_tech"].add(t.strip())

    total = sum(g["capacity_mw"] for g in groups.values())
    glist = sorted(groups.values(), key=lambda x: -x["capacity_mw"])
    for g in glist:
        g["share_pct"] = round(g["capacity_mw"] / total * 100, 2) if total else 0
        g["entity_count"] = len(g["entities"])
        g["locations"] = sorted(g.pop("_locs")); g["technologies"] = sorted(g.pop("_tech"))
        g["entities"].sort(key=lambda e: -(e["capacity_mw"] or 0))

    hhi = round(sum((g["share_pct"]) ** 2 for g in glist), 2)
    crN = lambda n: round(sum(g["share_pct"] for g in glist[:n]), 2)
    summary = {"total_capacity_mw": round(total, 1), "total_capacity_gw": round(total / 1000, 2),
               "groups": len(glist), "entities": len(gm), "hhi": hhi,
               "cr4": crN(4), "cr8": crN(8), "top5_share": crN(5),
               "concentration": ("HIGH" if hhi >= 2500 else "MODERATE" if hhi >= 1500
                                 else "GROUP-CONCENTRATED (looks fragmented on paper)")}
    result = {"available": True, "as_of": "ALCM List-II Rev-7 (30/04/2026)",
              "summary": summary, "groups": glist,
              "director_linkages": _read("alcm_director_cross_linkages.csv"),
              "shareholder_linkages": _read("alcm_shareholder_cross_linkages.csv"),
              "corporate_history": _read("alcm_corporate_history.csv"),
              "source": "ALCM/ (run_alcm_atlas.py output, mtime-keyed)",
              "fetched_at": datetime.now().strftime("%Y-%m-%d %H:%M")}
    set_cache(ck, result)
    mark_health("ALCM Atlas", True,
                f"{len(glist)} groups · {summary['total_capacity_gw']}GW · HHI {hhi}")
    _alcm_seed_memory(result)
    return result

@app.route("/api/alcm/atlas")
def api_alcm_atlas():
    """P19 — ALCM ownership atlas: group-level capacity, HHI/CR4 concentration,
    director & shareholder cross-linkages, corporate history."""
    return jsonify(fetch_alcm_atlas())

# ── P21.2 ALMM Ownership Atlas — true corporate structure of India's MODULE mfg ─
# Counterpart to the ALCM (cell) atlas. Loads ALMM/run_almm_atlas.py outputs,
# mtime-keyed. Group-level HHI/CR4 — the analysis the old PDF-parse never did.
ALMM_ATLAS_DIR = os.path.join(os.path.dirname(__file__), "ALMM")

def _almm_mtime():
    if not os.path.isdir(ALMM_ATLAS_DIR):
        return 0
    try:
        return int(max((os.path.getmtime(os.path.join(ALMM_ATLAS_DIR, f))
                        for f in os.listdir(ALMM_ATLAS_DIR) if f.endswith(".csv")), default=0))
    except Exception:
        return 0

def fetch_almm_atlas():
    import csv
    ck = f"almm_atlas:{_almm_mtime()}"
    cached = get_cache(ck, 21600)
    if cached:
        return cached

    def _read(name):
        p = os.path.join(ALMM_ATLAS_DIR, name)
        if not os.path.exists(p):
            return []
        try:
            with open(p, encoding="utf-8") as fh:
                return list(csv.DictReader(fh))
        except Exception:
            return []

    gm = _read("almm_group_mapping.csv")
    if not gm:
        mark_health("ALMM Atlas", False, "no almm_group_mapping.csv")
        return {"available": False,
                "note": "ALMM atlas not generated — run ALMM/run_almm_atlas.py to build it."}

    groups = {}
    for r in gm:
        g = (r.get("ultimate_group") or "Unknown").strip()
        try: cap = float(r.get("capacity_val") or 0)
        except (TypeError, ValueError): cap = 0.0
        e = groups.setdefault(g, {"group": g, "capacity_mw": 0.0, "entities": [],
                                  "_locs": set(), "_bis": 0})
        bis = [b.strip() for b in (r.get("bis_registrations") or "").split(",") if b.strip()]
        e["capacity_mw"] += cap
        e["_bis"] += len(bis)
        e["entities"].append({"name": r.get("manufacturer_normalized"), "capacity_mw": cap,
                              "relationship": r.get("relationship_type"), "evidence": r.get("evidence"),
                              "bis_registrations": len(bis), "location": r.get("locations")})
        if r.get("locations"): e["_locs"].add(r["locations"].strip())

    total = sum(g["capacity_mw"] for g in groups.values())
    glist = sorted(groups.values(), key=lambda x: -x["capacity_mw"])
    for g in glist:
        g["share_pct"] = round(g["capacity_mw"] / total * 100, 2) if total else 0
        g["entity_count"] = len(g["entities"]); g["bis_registrations"] = g.pop("_bis")
        g["locations"] = sorted(g.pop("_locs")); g["entities"].sort(key=lambda e: -(e["capacity_mw"] or 0))

    hhi = round(sum(g["share_pct"] ** 2 for g in glist), 2)
    crN = lambda n: round(sum(g["share_pct"] for g in glist[:n]), 2)
    summary = {"total_capacity_mw": round(total, 1), "total_capacity_gw": round(total / 1000, 2),
               "groups": len(glist), "entities": len(gm),
               "bis_registrations": sum(g["bis_registrations"] for g in glist),
               "hhi": hhi, "cr4": crN(4), "cr8": crN(8), "top10_share": crN(10),
               "concentration": ("HIGH" if hhi >= 2500 else "MODERATE" if hhi >= 1500
                                 else "COMPETITIVE (fragmented at industry level)")}
    result = {"available": True, "as_of": "ALMM Rev-XLVIII (01/05/2026)", "kind": "modules",
              "summary": summary, "groups": glist,
              "director_linkages": _read("director_cross_linkages.csv"),
              "shareholder_linkages": _read("shareholder_cross_linkages.csv"),
              "corporate_history": _read("corporate_history.csv"),
              "source": "ALMM/ (run_almm_atlas.py output, mtime-keyed)",
              "fetched_at": datetime.now().strftime("%Y-%m-%d %H:%M")}
    set_cache(ck, result)
    mark_health("ALMM Atlas", True,
                f"{len(glist)} groups · {summary['total_capacity_gw']}GW · HHI {hhi}")
    flag = f"almm_mem_seeded:{_almm_mtime()}"
    if not v11_sources.kv_get(flag):
        try:
            v16_mem.add_note(
                f"ALMM module-manufacturing: {summary['groups']} groups, {summary['total_capacity_gw']} GW "
                f"approved module capacity, HHI {hhi} (competitive — vs concentrated cells); CR4 {summary['cr4']}%, "
                f"top-10 {summary['top10_share']}%.", source="ALMM atlas")
            for g in glist[:10]:
                v16_mem.add_note(
                    f"ALMM ownership: {g['group']} controls {g['capacity_mw']:.0f} MW ({g['share_pct']}%) of "
                    f"India's approved solar module capacity.", source="ALMM atlas")
            v11_sources.kv_set(flag, "1")
        except Exception:
            pass
    return result

@app.route("/api/almm/atlas")
def api_almm_atlas():
    """P21.2 — ALMM (module) ownership atlas: group-level capacity, HHI/CR4,
    BIS registrations, cross-linkages. Counterpart to /api/alcm/atlas (cells)."""
    return jsonify(fetch_almm_atlas())

# ── P21.4 Word (.docx) newsletter export — editorial intelligence briefing ────
def build_newsletter_docx():
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    GOLD = RGBColor(0x9A, 0x6A, 0x00); INK = RGBColor(0x1B, 0x17, 0x12); MUT = RGBColor(0x6A, 0x62, 0x53)
    doc = Document()
    def _center(run_text, size, color, bold=True):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(run_text); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color; return p
    def _h(text):
        doc.add_paragraph()
        p = doc.add_paragraph(); r = p.add_run(text.upper()); r.bold = True
        r.font.size = Pt(12); r.font.color.rgb = GOLD; return p
    def _table(headers, rows):
        tb = doc.add_table(rows=1, cols=len(headers))
        try: tb.style = "Light List Accent 1"
        except Exception: pass
        for i, hh in enumerate(headers):
            c = tb.rows[0].cells[i]; c.paragraphs[0].text = ""
            rr = c.paragraphs[0].add_run(str(hh)); rr.bold = True; rr.font.size = Pt(8)
        for row in rows:
            cells = tb.add_row().cells
            for i, val in enumerate(row):
                cells[i].paragraphs[0].text = ""
                rr = cells[i].paragraphs[0].add_run("" if val is None else str(val)); rr.font.size = Pt(9)
        return tb

    _center("NEURON", 28, INK)
    _center("THE FOOL'S CREATION · INDIA RE INTELLIGENCE · OBSERVATORY", 9, GOLD)
    _center(datetime.now().strftime("Intelligence Briefing · %d %B %Y · %H:%M IST"), 9, MUT, bold=False)

    # 1. Executive decisions — use cached result, skip re-run to stay fast
    try:
        dec = get_cache("decisions_result", 3600) or {}
        ds = (dec.get("decisions") or [])[:6]
        if ds:
            _h(f"Executive Read — {dec.get('regime','')}")
            _table(["Band", "Conv", "Thesis", "Wrong if"],
                   [[d.get("band"), round(d.get("conviction", 0)), (d.get("thesis") or "")[:120],
                     (d.get("falsifier") or "")[:90]] for d in ds])
    except Exception:
        pass
    # 2. India capacity
    try:
        m = fetch_mnre_live(); D = m.get("data") or {}
        def _pk(k):
            v = D.get(k);  return (v if isinstance(v, (int, float)) else
                                   (v.get("cumulative_mw") or v.get("total_mw") or 0) if isinstance(v, dict) else 0)
        _h("India RE Capacity (MNRE)")
        _table(["Metric", "GW"], [["Total RE", round((m.get("total_re_mw") or 0)/1000, 1)],
               ["Solar", round(_pk("Solar Power")/1000, 1)], ["Wind", round(_pk("Wind Power")/1000, 1)],
               ["Hydro", round((_pk("Large Hydro")+_pk("Small Hydro Power"))/1000, 1)]])
    except Exception:
        pass
    # 3. Manufacturing ownership (modules + cells)
    for label, fn in (("Module manufacturing (ALMM)", fetch_almm_atlas),
                      ("Cell manufacturing (ALCM)", fetch_alcm_atlas)):
        try:
            a = fn()
            if a.get("available"):
                s = a["summary"]
                _h(f"{label} — {s['total_capacity_gw']} GW · {s['groups']} groups · HHI {s['hhi']}")
                _table(["Group", "MW", "Share"],
                       [[g["group"], round(g["capacity_mw"]), f"{g['share_pct']}%"] for g in a["groups"][:6]])
        except Exception:
            pass
    # 4. Markets — top movers
    try:
        qs = [q for q in fetch_all_quotes().values() if isinstance(q, dict) and q.get("change_pct") is not None]
        qs.sort(key=lambda q: q.get("change_pct", 0), reverse=True)
        movers = qs[:5] + qs[-5:]
        if movers:
            _h("Markets — RE equity movers")
            _table(["Stock", "Price ₹", "Chg %"],
                   [[q.get("name"), q.get("price"), f"{q.get('change_pct'):+.2f}"] for q in movers])
    except Exception:
        pass
    # 5. Chokepoints — kv-cached inside intelligence.py; use it but cap wait time
    try:
        import threading as _thr
        _cpres = [None]
        def _cpfetch(): _cpres[0] = v11_intel.chokepoint_monitor()
        t = _thr.Thread(target=_cpfetch, daemon=True); t.start(); t.join(timeout=18)
        cps = (_cpres[0] or {}).get("chokepoints", [])
        if cps:
            _h("Maritime Chokepoints → India import exposure")
            _table(["Chokepoint", "Status", "Exposure"],
                   [[c.get("name"), c.get("status"), (c.get("india_exposure") or "")[:110]] for c in cps])
    except Exception:
        pass

    doc.add_paragraph()
    _center("Generated by NEURON · sources: MNRE/CEA, IRENA, IMF, ALMM/ALCM, Yahoo Finance, 540+ feeds", 7, MUT, bold=False)
    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf

@app.route("/api/export/docx")
def api_export_docx():
    """P21.4 — download the intelligence briefing as a Word newsletter."""
    from flask import send_file
    try:
        buf = build_newsletter_docx()
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)[:120]}), 500
    fname = datetime.now().strftime("NEURON_Briefing_%Y-%m-%d.docx")
    return send_file(buf, as_attachment=True, download_name=fname,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ── P20 India macro/trade beliefs — the fundamentals leg for the decider ──────
# IMF DataMapper (keyless) + Yahoo FX, seeded into the v15 belief layer so the
# decision engine reasons over fundamentals, not just news/flow. EIA + Trading-
# Economics are optional (env-keyed); absent ⇒ honest gap, never fabricated.
_IMF_INDICATORS = [
    ("india_gdp_growth",      "NGDP_RPCH", "%",     "India GDP growth (IMF)"),
    ("india_inflation",       "PCPIPCH",   "%",     "India CPI inflation (IMF)"),
    ("india_current_account", "BCA_NGDPD", "% GDP", "India current account (IMF)"),
    ("india_govt_debt",       "GGXWDG_NGDP", "% GDP", "India general govt debt (IMF)"),
]

def _imf_latest(indicator):
    try:
        r = requests.get(f"https://www.imf.org/external/datamapper/api/v1/{indicator}/IND",
                         timeout=12)
        vals = (r.json().get("values", {}).get(indicator, {}) or {}).get("IND", {})
        if not vals:
            return None, None
        # Prefer the latest year <= the current year (IMF series include multi-year
        # projections — a belief should reflect 'now', not a 2031 forecast).
        cur = datetime.now().year
        yrs = sorted(int(y) for y in vals)
        pick = next((y for y in reversed(yrs) if y <= cur), yrs[-1])
        return round(float(vals[str(pick)]), 2), str(pick)
    except Exception:
        return None, None

def fetch_india_macro_plus():
    cached = get_cache("india_macro_plus", 12 * 3600)
    if cached:
        return cached
    metrics = []
    for metric, ind, unit, label in _IMF_INDICATORS:
        val, yr = _imf_latest(ind)
        if val is not None:
            metrics.append({"metric": metric, "value": val, "unit": unit,
                            "label": label, "source": "IMF", "as_of": yr})
    # USD/INR from the already-cached quote layer (import-cost pressure proxy).
    try:
        inr = fetch_all_quotes().get("USDINR=X", {})
        if isinstance(inr, dict) and inr.get("price"):
            metrics.append({"metric": "usd_inr", "value": round(float(inr["price"]), 2),
                            "unit": "₹/$", "label": "USD/INR", "source": "Yahoo",
                            "as_of": datetime.now().strftime("%Y-%m-%d")})
    except Exception:
        pass
    # Optional EIA India crude consumption (import-bill proxy) — only if key set.
    if EIA_API_KEY:
        try:
            er = requests.get("https://api.eia.gov/v2/international/data/", timeout=12, params={
                "api_key": EIA_API_KEY, "frequency": "annual",
                "data[0]": "value", "facets[countryRegionId][]": "IND",
                "facets[productId][]": "5", "facets[activityId][]": "2",
                "sort[0][column]": "period", "sort[0][direction]": "desc", "length": "1"})
            d0 = (er.json().get("response", {}).get("data") or [{}])[0]
            if d0.get("value") is not None:
                metrics.append({"metric": "india_oil_consumption", "value": round(float(d0["value"]), 1),
                                "unit": d0.get("unit", "Mbbl/d"), "label": "India oil consumption (EIA)",
                                "source": "EIA", "as_of": str(d0.get("period", ""))})
        except Exception:
            pass
    # Seed every metric into the belief layer (cognition owns revision/conflict).
    seeded = []
    for m in metrics:
        try:
            r = v11_cog.upsert_belief(m["metric"], m["value"], m["unit"], m["label"],
                                      m["source"], m["as_of"], confidence="MEDIUM")
            seeded.append({"metric": m["metric"], "action": r.get("action")})
        except Exception:
            pass
    result = {"metrics": metrics, "seeded": seeded,
              "sources": "IMF DataMapper + Yahoo" + (" + EIA" if EIA_API_KEY else "")
                         + (" + TradingEconomics" if TRADINGECONOMICS_KEY else ""),
              "note": None if (EIA_API_KEY and TRADINGECONOMICS_KEY) else
                      "Set EIA_API_KEY / TRADINGECONOMICS_KEY in .env to add import-bill & calendar.",
              "fetched_at": datetime.now().strftime("%Y-%m-%d %H:%M")}
    set_cache("india_macro_plus", result)
    mark_health("India Macro+", bool(metrics), f"{len(metrics)} metrics seeded")
    return result

@app.route("/api/india_macro_plus")
def api_india_macro_plus():
    """P20 — India macro/trade fundamentals (IMF + FX), seeded as beliefs."""
    return jsonify(fetch_india_macro_plus())

# ── P20 Forecast markets — Polymarket cross-check (keyless Gamma API) ──────────
# Market-implied probabilities for energy/geopolitics/macro events relevant to
# India RE — a market-money second opinion alongside Neuron's own re_forecast.
_POLY_KW = ("oil", "opec", "crude", "brent", "natural gas", " lng", "coal", "gasoline",
            "energy", "power grid", "blackout", "solar", "renewable", "climate", "carbon",
            "global temperature", "heat record", "warming", "india", "modi", "rupee",
            "rate cut", "rate hike", "fed ", "inflation", "recession", "gdp", "tariff",
            "trade war", "hormuz", "strait", "red sea", "houthi", "saudi", "sanction",
            "nuclear deal", "invade", "ceasefire", "war with", "taiwan", "semiconductor")
# Exclude sports / celebrity / entertainment noise that shares loose keywords.
_POLY_EXCLUDE = ("fifa", "world cup", "super bowl", "nba", "nfl", "mvp", "grammy", "oscar",
                 "kardashian", "lebron", "ramaswamy", "walz", "newsom", "taylor swift",
                 "movie", "box office", "album", "premier league", "champions league",
                 "ballon", "olympic", "cricket", " ipl", "wwe", "ufc", "heisman")

def fetch_forecast_markets():
    cached = get_cache("forecast_markets", 3 * 3600)
    if cached:
        return cached
    out = []
    try:
        req = requests.get("https://gamma-api.polymarket.com/markets", timeout=15,
                           headers={"User-Agent": "Mozilla/5.0 (NEURON)"},
                           params={"active": "true", "closed": "false", "archived": "false",
                                   "order": "volumeNum", "ascending": "false", "limit": "150"})
        mkts = req.json()
        mkts = mkts if isinstance(mkts, list) else mkts.get("data", [])
        for m in mkts:
            q = (m.get("question") or "").strip()
            ql = q.lower()
            if not q or not any(k in ql for k in _POLY_KW) or any(x in ql for x in _POLY_EXCLUDE):
                continue
            try: prices = json.loads(m.get("outcomePrices") or "[]")
            except Exception: prices = []
            try: outs = json.loads(m.get("outcomes") or "[]")
            except Exception: outs = []
            prob = None
            if prices:
                try: prob = round(float(prices[0]) * 100, 1)
                except Exception: prob = None
            try: vol = float(m.get("volume") or m.get("volumeNum") or 0)
            except Exception: vol = 0.0
            slug = m.get("slug") or ""
            out.append({"question": q[:150], "prob_pct": prob,
                        "outcome": (outs[0] if outs else "Yes"), "volume": round(vol),
                        "ends": (m.get("endDate") or "")[:10],
                        "url": "https://polymarket.com/event/" + slug if slug else "https://polymarket.com"})
        out.sort(key=lambda x: -(x["volume"] or 0))
        out = out[:12]
    except Exception as e:
        result = {"markets": [], "error": str(e)[:90], "source": "Polymarket Gamma (keyless)",
                  "fetched_at": datetime.now().strftime("%Y-%m-%d %H:%M")}
        set_cache("forecast_markets", result)
        mark_health("Forecast Markets", False, str(e)[:60])
        return result
    result = {"markets": out, "count": len(out), "source": "Polymarket Gamma (keyless)",
              "note": "Market-implied probabilities for energy/geo/macro events relevant to India RE.",
              "fetched_at": datetime.now().strftime("%Y-%m-%d %H:%M")}
    set_cache("forecast_markets", result)
    mark_health("Forecast Markets", bool(out), f"{len(out)} markets")
    return result

@app.route("/api/forecast_markets")
def api_forecast_markets():
    """P20 — Polymarket-implied probabilities for events relevant to India RE."""
    return jsonify(fetch_forecast_markets())

@app.route("/api/statewise")
def api_statewise(): return jsonify(fetch_statewise())

@app.route("/api/mnre_state_capacity")
def api_mnre_state_capacity(): return jsonify(fetch_mnre_state_capacity())

@app.route("/api/cea_history")
def api_cea_history():
    """Monthly snapshots of India RE capacity — used for growth chart (P6.3)."""
    try:
        con = sqlite3.connect(DB_PATH)
        rows = con.execute(
            "SELECT snap_date,re_total_mw,solar_mw,wind_mw,hydro_mw FROM cea_national_snap ORDER BY snap_date"
        ).fetchall()
        con.close()
        return jsonify([{"date":r[0],"re_total":r[1],"solar":r[2],"wind":r[3],"hydro":r[4]} for r in rows])
    except Exception as e:
        return jsonify({"error": str(e)})

@app.route("/api/pm_surya_ghar")
def api_pm_surya_ghar(): return jsonify(fetch_pm_surya_ghar())

@app.route("/api/pm_kusum")
def api_pm_kusum(): return jsonify(fetch_pm_kusum())

@app.route("/api/alerts")
def api_alerts(): return jsonify(get_alerts())

@app.route("/api/alerts/history")
def api_alerts_history(): return jsonify(get_alert_history(200))

@app.route("/api/seci_tenders")
def api_seci_tenders(): return jsonify(fetch_seci_tenders())

@app.route("/api/correlation")
def api_correlation(): return jsonify(fetch_correlation())

@app.route("/api/live_channels")
def api_live_channels(): return jsonify(LIVE_CHANNELS)

@app.route("/api/youtube_live/<channel_id>")
def api_youtube_live(channel_id):
    return jsonify(get_youtube_live(re.sub(r"[^A-Za-z0-9_-]","",channel_id)))

@app.route("/api/worldbank")
def api_worldbank(): return jsonify({"renewable_pct":india_energy_worldbank()})

@app.route("/api/history/<symbol>")
def api_history(symbol):
    sym    = symbol if ("." in symbol or symbol in _NON_NSE_TICKERS) else symbol+".NS"
    cached = get_cache(f"hist_{sym}")
    if cached: return jsonify(cached)
    data   = fetch_history(sym); set_cache(f"hist_{sym}",data); return jsonify(data)

@app.route("/api/analysis/<symbol>")
def api_analysis(symbol):
    sym    = symbol if ("." in symbol or symbol in _NON_NSE_TICKERS) else symbol+".NS"
    prices = fetch_history(sym,"2y")
    return jsonify({"symbol":sym,"technicals":compute_technicals(prices),
                    "projection":compute_projection(prices),"history":prices[-180:]})

# ── P21.3 Quant Signals — ranked panel (RSI/MACD/momentum z/vol/beta/drawdown) ─
def _quant_score_stock(sym, nifty_closes):
    """Compute a composite quant score [0–100] for one stock."""
    try:
        prices = fetch_history(sym, "1y")
        if len(prices) < 60:
            return None
        closes = np.array([p["close"] for p in prices], dtype=float)
        tech = compute_technicals(prices)
        # RSI signal: neutral=50→score 0, overbought/oversold extremes score higher
        rsi = tech.get("rsi", 50) or 50
        rsi_score = max(0, min(1, abs(rsi - 50) / 50))  # 0=neutral, 1=extreme
        # Momentum z-score: 20-day return vs 252-day rolling std
        ret20 = (closes[-1] / closes[-21] - 1) if len(closes) >= 21 else 0
        std252 = float(np.std(np.diff(closes[-252:]) / closes[-252:-1])) if len(closes) >= 252 else float(np.std(np.diff(closes) / closes[:-1]))
        mom_z = abs(ret20 / (std252 * (252**0.5))) if std252 > 0 else 0
        # Volatility (annualised): higher = more signal
        vol_ann = std252 * (252**0.5)
        # Beta vs Nifty (if nifty_closes available)
        beta = 1.0
        if nifty_closes is not None and len(nifty_closes) >= 30:
            n = min(len(closes), len(nifty_closes), 252)
            sr = np.diff(closes[-n:]) / closes[-n:-1]
            nr = np.diff(nifty_closes[-n:]) / nifty_closes[-n:-1]
            if np.std(nr) > 0:
                beta = float(np.cov(sr, nr)[0, 1] / np.var(nr))
        # Max drawdown (1y)
        peak = np.maximum.accumulate(closes)
        dd = float(np.min((closes - peak) / peak))  # negative
        # Composite: weighted combination → normalize 0-100
        composite = (
            rsi_score * 20 +
            min(mom_z, 3) / 3 * 30 +
            min(vol_ann, 0.8) / 0.8 * 15 +
            min(abs(beta), 2.5) / 2.5 * 15 +
            min(abs(dd), 0.5) / 0.5 * 20
        )
        # Directional label
        macd = tech.get("macd", 0) or 0
        if rsi > 70 or (mom_z > 1 and macd > 0):
            direction = "BULLISH"
        elif rsi < 30 or (mom_z > 1 and macd < 0):
            direction = "BEARISH"
        else:
            direction = "NEUTRAL"
        return {
            "symbol": sym.replace(".NS", ""),
            "name": RE_STOCKS.get(sym, sym.replace(".NS", "")),
            "score": round(composite, 1),
            "direction": direction,
            "rsi": round(rsi, 1),
            "macd": round(macd, 2),
            "mom_z": round(float(mom_z), 2),
            "vol_ann_pct": round(vol_ann * 100, 1),
            "beta": round(beta, 2),
            "drawdown_pct": round(dd * 100, 1),
            "price": prices[-1]["close"] if prices else None,
        }
    except Exception:
        return None

def fetch_quant_signals():
    ck = "quant_signals"
    cached = get_cache(ck, 1800)
    if cached:
        return cached
    # Get Nifty50 as benchmark
    nifty_closes = None
    try:
        np_prices = fetch_history("^NSEI", "1y")
        if np_prices:
            nifty_closes = np.array([p["close"] for p in np_prices], dtype=float)
    except Exception:
        pass
    signals = []
    with ThreadPoolExecutor(max_workers=6) as ex:
        futs = {ex.submit(_quant_score_stock, sym, nifty_closes): sym for sym in RE_STOCKS}
        for f in as_completed(futs):
            r = f.result()
            if r:
                signals.append(r)
    signals.sort(key=lambda x: x["score"] if x["score"] is not None else -999.0, reverse=True)
    result = {"signals": signals, "count": len(signals),
              "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M")}
    set_cache(ck, result)
    return result

@app.route("/api/quant_signals")
def api_quant_signals():
    """P21.3 — ranked quant signals across RE watchlist (RSI/MACD/mom-z/vol/beta/drawdown)."""
    return jsonify(fetch_quant_signals())

# ── P21 Tender Tracker ───────────────────────────────────────────────────────
_TENDER_KEYWORDS = [
    "tender", "rfp", "eoi", "bid", "request for proposal", "expression of interest",
    "solicitation", "procurement", "auction", "tariff discovery", "reverse auction",
    "capacity addition", "mw solar", "mw wind", "mw bess", "mw storage",
    "mw green hydrogen", "mw electrolysis",
]
_TENDER_ENTITIES = [
    "seci", "ntpc", "recl", "pfc", "ireda", "sjvn", "nhpc", "thdc", "neepco",
    "msedcl", "tneb", "tnerc", "kseb", "bescom", "aptransco", "tsgenco",
    "rrecl", "guvnl", "mnre", "cea", "discoms", "discom",
]

def _parse_mw_from_title(title):
    """Extract capacity MW from tender headline. Returns float or None."""
    for pat in [r"(\d[\d,]*(?:\.\d+)?)\s*GW", r"(\d[\d,]*(?:\.\d+)?)\s*MW"]:
        m = re.search(pat, title, re.I)
        if m:
            v = float(m.group(1).replace(",",""))
            if "GW" in m.group(0).upper(): v *= 1000
            return v
    return None

def _detect_sector(title):
    t = title.lower()
    if any(w in t for w in ["bess","battery","storage","btm"]): return "BESS"
    if any(w in t for w in ["green hydrogen","gh2","electrolyser","electrolysis","ammonia"]): return "GH"
    if "wind" in t and "solar" not in t: return "Wind"
    if "solar" in t and "wind" not in t: return "Solar"
    if "hydro" in t: return "Hydro"
    if "hybrid" in t or ("solar" in t and "wind" in t): return "Hybrid"
    return "RE"

def _detect_state(title):
    states = ["Rajasthan","Gujarat","Maharashtra","Tamil Nadu","Andhra Pradesh",
              "Telangana","Karnataka","Madhya Pradesh","Uttar Pradesh","Punjab",
              "Haryana","Odisha","West Bengal","Kerala","Himachal Pradesh","Uttarakhand",
              "Jharkhand","Bihar","Chhattisgarh","Assam","Goa","Jammu","Kashmir"]
    for s in states:
        if s.lower() in title.lower(): return s
    return None

def ingest_tenders_from_news():
    """Scan recent news archive for tender signals → upsert into v21_tenders."""
    try:
        con = sqlite3.connect(DB_PATH)
        con.execute("PRAGMA journal_mode=WAL")
        # Pull last 7 days of news with tender keywords
        cutoff = time.time() - 7 * 86400
        rows = con.execute(
            "SELECT uid, title, link, published_dt FROM news_archive WHERE ts >= ? ORDER BY ts DESC LIMIT 400",
            (cutoff,)).fetchall()
        added = 0
        for uid, title, link, pub_dt in rows:
            if not title: continue
            tl = title.lower()
            # Must match a keyword + an entity (or contain MW/GW)
            kw_hit = any(k in tl for k in _TENDER_KEYWORDS)
            entity_hit = any(e in tl for e in _TENDER_ENTITIES)
            mw = _parse_mw_from_title(title)
            if not (kw_hit and (entity_hit or mw)): continue
            # Avoid duplicates by source_title similarity
            existing = con.execute(
                "SELECT id FROM v21_tenders WHERE source_title=?", (title[:200],)).fetchone()
            if existing: continue
            sector = _detect_sector(title)
            state = _detect_state(title)
            # Extract entity from title
            entity = next((e.upper() for e in _TENDER_ENTITIES if e in tl), "UNKNOWN")
            con.execute("""INSERT INTO v21_tenders
                (entity,project_name,tender_type,capacity_mw,sector,state,
                 announced_date,status,source_url,source_title,ts)
                VALUES(?,?,?,?,?,?,?,?,?,?,?)""",
                (entity, title[:140], "Tender", mw, sector, state,
                 pub_dt, "OPEN", link, title[:200], time.time()))
            added += 1
        con.commit(); con.close()
        return added
    except Exception as e:
        return 0

@app.route("/api/tenders")
def api_tenders():
    """P21 — tender tracker: all recorded tenders with filters."""
    sector = request.args.get("sector","")
    state  = request.args.get("state","")
    status = request.args.get("status","")
    limit  = min(int(request.args.get("limit","100")), 500)
    # Auto-ingest new tenders from recent news (fast, deduped)
    try: ingest_tenders_from_news()
    except Exception: pass
    con = sqlite3.connect(DB_PATH)
    q = "SELECT id,entity,project_name,tender_type,capacity_mw,sector,state,announced_date,bid_deadline,ppa_signed_date,status,source_url FROM v21_tenders WHERE 1=1"
    params = []
    if sector: q += " AND sector=?"; params.append(sector)
    if state:  q += " AND state LIKE ?"; params.append(f"%{state}%")
    if status: q += " AND status=?"; params.append(status)
    q += " ORDER BY ts DESC LIMIT ?"
    params.append(limit)
    rows = con.execute(q, params).fetchall()
    con.close()
    cols = ["id","entity","project_name","tender_type","capacity_mw","sector","state",
            "announced_date","bid_deadline","ppa_signed_date","status","source_url"]
    tenders = [dict(zip(cols, r)) for r in rows]
    total_mw = sum(t["capacity_mw"] or 0 for t in tenders)
    return jsonify({"tenders": tenders, "count": len(tenders),
                    "total_mw": round(total_mw, 1),
                    "sectors": list({t["sector"] for t in tenders if t["sector"]})})

@app.route("/api/tenders/ingest", methods=["POST"])
def api_tenders_ingest():
    """Trigger tender ingestion from news."""
    added = ingest_tenders_from_news()
    return jsonify({"ok": True, "added": added})

@app.route("/api/tenders/add", methods=["POST"])
def api_tenders_add():
    """Manually add or update a tender record."""
    b = request.get_json(silent=True) or {}
    required = ["entity","project_name"]
    missing = [k for k in required if not b.get(k)]
    if missing: return jsonify({"error": f"missing: {missing}"}), 400
    con = sqlite3.connect(DB_PATH)
    con.execute("""INSERT INTO v21_tenders
        (entity,project_name,tender_type,capacity_mw,sector,state,
         announced_date,bid_deadline,ppa_signed_date,status,source_url,source_title,ts)
        VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (b.get("entity",""), b.get("project_name","")[:140], b.get("tender_type","Tender"),
         b.get("capacity_mw"), b.get("sector","RE"), b.get("state"),
         b.get("announced_date"), b.get("bid_deadline"), b.get("ppa_signed_date"),
         b.get("status","OPEN"), b.get("source_url"), b.get("source_title",""), time.time()))
    con.commit(); con.close()
    return jsonify({"ok": True})

@app.route("/api/tenders/<int:tid>/ppa", methods=["POST"])
def api_tenders_ppa(tid):
    """Mark PPA signed date on an existing tender."""
    b = request.get_json(silent=True) or {}
    ppa_date = b.get("ppa_signed_date")
    if not ppa_date: return jsonify({"error": "ppa_signed_date required"}), 400
    con = sqlite3.connect(DB_PATH)
    con.execute("UPDATE v21_tenders SET ppa_signed_date=?, status='PPA_SIGNED' WHERE id=?",
                (ppa_date, tid))
    con.commit(); con.close()
    return jsonify({"ok": True})

# ── P21 CEA Daily Generation (enhanced with region) ──────────────────────────
def fetch_cea_daily_generation():
    """Pull today's generation breakdown from CEA — solar/wind/hydro/thermal/total."""
    cached = get_cache("cea_daily_gen", 3600)
    if cached: return cached
    try:
        r = gov_get("https://vidyut.cea.gov.in/dashboard", timeout=20)
        soup = BeautifulSoup(r.text, "html.parser")
        text = soup.get_text(" ", strip=True)
        # Parse MU figures from the page
        result = {"sectors": {}, "fetched_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
                  "source": "vidyut.cea.gov.in"}
        for tech, patterns in [
            ("Solar",   [r"Solar[^\d]{0,20}([\d,\.]+)\s*MU", r"([\d,\.]+)\s*MU.*?Solar"]),
            ("Wind",    [r"Wind[^\d]{0,20}([\d,\.]+)\s*MU",  r"([\d,\.]+)\s*MU.*?Wind"]),
            ("Hydro",   [r"Hydro[^\d]{0,20}([\d,\.]+)\s*MU"]),
            ("Thermal", [r"Thermal[^\d]{0,20}([\d,\.]+)\s*MU"]),
            ("Nuclear", [r"Nuclear[^\d]{0,20}([\d,\.]+)\s*MU"]),
            ("Total",   [r"Total[^\d]{0,20}([\d,\.]+)\s*MU", r"Grand Total[^\d]{0,20}([\d,\.]+)"]),
        ]:
            for pat in patterns:
                m = re.search(pat, text, re.I)
                if m:
                    try: result["sectors"][tech] = float(m.group(1).replace(",",""))
                    except: pass
                    break
        # Fallback to cea_generation endpoint data
        if not result["sectors"]:
            gen = fetch_cea_generation()
            for k, v in (gen.get("data") or {}).items():
                if isinstance(v, dict): result["sectors"][k.title()] = v.get("generation_mu",0)
        set_cache("cea_daily_gen", result)
        return result
    except Exception as e:
        return {"sectors": {}, "error": str(e)[:80], "fetched_at": datetime.now().strftime("%Y-%m-%d")}

@app.route("/api/cea_daily_gen")
def api_cea_daily_gen():
    return jsonify(fetch_cea_daily_generation())

# ── P21 ALMM/ALCM Wp Bucket Analysis ─────────────────────────────────────────
def fetch_almm_wp_buckets():
    """Group ALMM List-II by Wp bucket: <400, 400-500, 500-600, 600+ with count/top eff."""
    cached = get_cache("almm_wp_buckets", 21600)
    if cached: return cached
    try:
        con = sqlite3.connect(DB_PATH)
        rows = con.execute(
            "SELECT parent_company, capacity_mw_yr, efficiency, module_wp, module_type "
            "FROM almm_modules WHERE module_wp IS NOT NULL ORDER BY efficiency DESC"
        ).fetchall()
        con.close()
        if not rows:
            return {"available": False, "note": "ALMM modules not parsed yet"}
        buckets = {}
        for parent, cap, eff, wp, mtype in rows:
            if wp < 400: bk = "<400 Wp"
            elif wp < 500: bk = "400–500 Wp"
            elif wp < 600: bk = "500–600 Wp"
            else: bk = "600+ Wp"
            if bk not in buckets:
                buckets[bk] = {"count": 0, "top_eff": 0, "total_cap_mw": 0,
                               "top_company": "", "top_wp": 0, "techs": set()}
            b = buckets[bk]
            b["count"] += 1
            b["total_cap_mw"] = round((b["total_cap_mw"] or 0) + (cap or 0), 1)
            if (eff or 0) > b["top_eff"]:
                b["top_eff"] = round(eff, 2)
                b["top_company"] = parent or ""
                b["top_wp"] = wp
            if mtype: b["techs"].add(mtype)
        for b in buckets.values(): b["techs"] = list(b["techs"])[:5]
        result = {"buckets": [{"label": k, **v} for k, v in sorted(buckets.items())],
                  "total_models": len(rows), "available": True,
                  "as_of": datetime.now().strftime("%Y-%m-%d")}
        set_cache("almm_wp_buckets", result)
        return result
    except Exception as e:
        return {"available": False, "error": str(e)[:80]}

@app.route("/api/almm/wp_buckets")
def api_almm_wp_buckets():
    return jsonify(fetch_almm_wp_buckets())

def fetch_alcm_wp_buckets():
    """Group ALCM List-I by Wp bucket: <450, 450-550, 550+ with count/top eff."""
    cached = get_cache("alcm_wp_buckets", 21600)
    if cached: return cached
    try:
        con = sqlite3.connect(DB_PATH)
        rows = con.execute(
            "SELECT parent_company, capacity_val, efficiency, technologies "
            "FROM almm_list WHERE efficiency IS NOT NULL ORDER BY efficiency DESC"
        ).fetchall()
        con.close()
        if not rows:
            return {"available": False, "note": "ALCM data not parsed yet"}
        # ALCM capacity_val is in MW/yr; no per-cell Wp — use capacity range as proxy
        total_cap = sum(float(r[1] or 0) for r in rows)
        companies = {}
        for parent, cap, eff, tech in rows:
            p = parent or "Unknown"
            if p not in companies:
                companies[p] = {"cap_mw": 0, "top_eff": 0, "techs": set()}
            companies[p]["cap_mw"] = round(companies[p]["cap_mw"] + float(cap or 0), 1)
            if (eff or 0) > companies[p]["top_eff"]:
                companies[p]["top_eff"] = round(eff, 2)
            if tech: companies[p]["techs"].add(str(tech))
        top = sorted(companies.items(), key=lambda x: x[1]["cap_mw"], reverse=True)[:12]
        result = {"companies": [{"name": k, **{**v, "techs": list(v["techs"])[:3]}} for k, v in top],
                  "total_models": len(rows), "total_cap_mw": round(total_cap, 1),
                  "available": True, "as_of": datetime.now().strftime("%Y-%m-%d")}
        set_cache("alcm_wp_buckets", result)
        return result
    except Exception as e:
        return {"available": False, "error": str(e)[:80]}

@app.route("/api/alcm/wp_buckets")
def api_alcm_wp_buckets():
    return jsonify(fetch_alcm_wp_buckets())

@app.route("/api/health")
def api_health():
    # P15 A5/B2/B4 — surface worker liveness, belief conflicts, consolidation
    # freshness and prompt-guard activity. Anything degraded shows here loudly.
    worker, beliefs, consolidation, guard, last_test = {}, {}, {}, {}, None
    try: worker = v11_sources.worker_health()
    except Exception as e: worker = {"status": "UNKNOWN", "error": str(e)[:80]}
    try:
        bv = v11_cog.beliefs_view()
        beliefs = {"count": bv["count"], "conflict_count": bv["conflict_count"],
                   "conflicts": bv["conflicts"]}
    except Exception as e: beliefs = {"error": str(e)[:80]}
    try: consolidation = v11_cog.consolidation_status()
    except Exception as e: consolidation = {"error": str(e)[:80]}
    try: guard = v11_intel.prompt_guard_stats()
    except Exception as e: guard = {"error": str(e)[:80]}
    try:
        raw = v11_sources.kv_get("last_self_test")
        last_test = json.loads(raw) if raw else None
    except Exception: last_test = None
    return jsonify({"sources": SOURCE_HEALTH, "cache_keys": list(cache.keys()),
                    "seen_alerts": len(SEEN_ALERTS),
                    "worker": worker, "beliefs": beliefs,
                    "consolidation": consolidation, "prompt_guard": guard,
                    "last_self_test": last_test,
                    "ts": datetime.now().strftime("%Y-%m-%d %H:%M:%S")})

@app.route("/api/watchlist", methods=["GET"])
def api_watchlist_get():
    return jsonify(RE_STOCKS)

@app.route("/api/watchlist", methods=["POST"])
def api_watchlist_post():
    body = request.get_json(silent=True) or {}
    action = body.get("action","")
    ticker = body.get("ticker","").upper().strip()
    name   = body.get("name","").strip()
    if not ticker: return jsonify({"error":"ticker required"}),400
    if action == "add":
        sym = ticker if "." in ticker else ticker+".NS"
        if not name: name = ticker.replace(".NS","")
        RE_STOCKS[sym] = name
        save_watchlist(RE_STOCKS)
        cache.pop("all_quotes", None)
        return jsonify({"ok":True,"added":sym})
    elif action == "remove":
        sym = ticker if "." in ticker else ticker+".NS"
        removed = RE_STOCKS.pop(sym, None)
        save_watchlist(RE_STOCKS)
        cache.pop("all_quotes", None)
        return jsonify({"ok":True,"removed":sym if removed else None})
    return jsonify({"error":"action must be add or remove"}), 400

@app.route("/api/brief")
def api_brief(): return jsonify(fetch_daily_brief())

@app.route("/api/signal_score")
def api_signal_score(): return jsonify(compute_signal_score())

@app.route("/api/iex_power")
def api_iex_power(): return jsonify(fetch_iex_power())

@app.route("/api/intel_engine")
def api_intel_engine(): return jsonify(fetch_intel_engine())

@app.route("/api/intel_engine/refresh")
def api_intel_engine_refresh():
    cache.pop("intel_engine", None)
    return jsonify(fetch_intel_engine())

@app.route("/api/pulse_history")
def api_pulse_history():
    try:
        con = sqlite3.connect(DB_PATH)
        rows = con.execute("SELECT date, pulse, label, hot_topics FROM pulse_history ORDER BY date DESC LIMIT 60").fetchall()
        con.close()
        return jsonify([{"date": r[0], "pulse": r[1], "label": r[2],
                         "topics": (r[3] or "").split(",")[:3]} for r in rows])
    except Exception as e:
        return jsonify({"error": str(e)})

@app.route("/api/company_news")
def api_company_news(): return jsonify(fetch_company_news())

@app.route("/api/mnre_notifications")
def api_mnre_notifications(): return jsonify(fetch_mnre_notifications())

# /api/ntpc_tenders, /api/nhpc_tenders, /api/sjvn_tenders removed (P6.1)
# These scrapers are used internally by fetch_intel_engine() — not needed as standalone endpoints

@app.route("/api/india_macro")
def api_india_macro(): return jsonify(fetch_india_macro())

@app.route("/api/sector_history")
def api_sector_history():
    period = request.args.get("period","1mo")
    return jsonify(fetch_sector_history(period))

@app.route("/api/focus_stock")
def api_focus_stock(): return jsonify(fetch_focus_stock_deep())

@app.route("/api/pv_prices")
def api_pv_prices(): return jsonify(fetch_pv_prices())

@app.route("/api/seci_results")
def api_seci_results(): return jsonify(fetch_seci_results())

@app.route("/api/cea_generation")
def api_cea_generation(): return jsonify(fetch_cea_generation())

@app.route("/api/dashboard")
def api_dashboard():
    return jsonify({
        "quotes":     fetch_all_quotes(),
        "news":       fetch_news()[:15],
        "commodities":fetch_commodities(),
        "global_re":  fetch_global_re(),
        "alerts":     get_alerts()[:10],
        "health":     SOURCE_HEALTH
    })

# ── ALMM/ALCM PDF Parser (P1.4) ──────────────────────────────────────────────
ALMM_PAGE_URL = "https://mnre.gov.in/approved-list-of-models-and-manufacturers-almm/"

ALMM_PARENT_MAP = {
    # Adani Group (Mundra Solar = Adani's cell/module mfg arm)
    "adani solar":               "Adani Solar",
    "adani green":               "Adani Solar",
    "mundra solar":              "Adani Solar",
    # Waaree Energies + subsidiaries
    "waaree energies":           "Waaree Energies",
    "waaree":                    "Waaree Energies",
    "fs india solar":            "Waaree Energies",   # wholly-owned subsidiary of Waaree
    # Vikram Solar
    "vikram solar":              "Vikram Solar",
    # Tata Power Solar
    "tata power solar":          "Tata Power Solar",
    "tata power renewable":      "Tata Power Solar",
    "tata":                      "Tata Power Solar",
    # Premier Energies (two manufacturing entities)
    "premier energies":          "Premier Energies",
    # Goldi Solar (two entities: Goldi Sun + Goldi Solar Pvt)
    "goldi":                     "Goldi Solar",
    # Saatvik Green Energy
    "saatvik":                   "Saatvik Green Energy",
    # RenewSys India
    "renewsys":                  "RenewSys India",
    # Emmvee Solar (two entities: Emmvee Energy + Emmvee Photovoltaic Power)
    "emmvee":                    "Emmvee Solar",
    # ReNew Power
    "renew photovoltaics":       "ReNew Power",
    # Avaada Energy
    "avaada":                    "Avaada Energy",
    # Greenko
    "greenko":                   "Greenko",
    # Websol Energy
    "websol":                    "Websol Energy",
    # Sova Solar (JBM Group subsidiary)
    "sova solar":                "Sova Solar",
    # Jupiter International
    "jupiter international":     "Jupiter International",
    # Insolation Energy
    "insolation energy":         "Insolation Energy",
    # Solex Energy
    "solex energy":              "Solex Energy",
}

def _clean_mfr(name):
    """Strip M/S., M/s., M/S prefixes from manufacturer names."""
    return re.sub(r'^M\s*/\s*[Ss][\.\s]+', '', (name or '')).strip()

def _map_parent(mfr_name):
    if not mfr_name: return mfr_name
    cleaned = _clean_mfr(mfr_name)
    lo = cleaned.lower()
    for k, v in ALMM_PARENT_MAP.items():
        if k in lo: return v
    return cleaned.title()

def _fetch_almm_pdf_url():
    """Scrape MNRE ALMM page, return (pdf_url, pub_date_str) of latest list."""
    hdr = {"User-Agent": "Mozilla/5.0 Chrome/124 Safari/537.36"}
    r   = gov_get(ALMM_PAGE_URL, timeout=15)
    soup = BeautifulSoup(r.text, "html.parser")
    best_url, best_date = None, ""
    for a in soup.find_all("a", href=True):
        href = a["href"]
        text = a.get_text(strip=True)
        if re.search(r"\.pdf", href, re.I) and re.search(r"almm|module|manufacturer", href + text, re.I):
            date_m = re.search(r"(\d{2}[.\-/]\d{2}[.\-/]\d{4}|\d{4}[.\-/]\d{2}[.\-/]\d{2})", text + href)
            pub_date = date_m.group(1) if date_m else ""
            if not best_url or pub_date > best_date:
                best_url  = href if href.startswith("http") else "https://mnre.gov.in" + href
                best_date = pub_date
    return best_url, best_date

def _parse_almm_pdf(pdf_path):
    """Parse ALMM PDF robustly. Handles multiline headers and S.No. in col 0."""
    import pdfplumber
    records = []
    # Column indices — detected from first header-bearing table, reused for all pages
    col_mfr = col_model = col_cap = col_eff = col_tech = col_val = None

    def _norm(cell):
        """Normalize a cell value: collapse newlines and spaces."""
        if not cell: return ""
        return re.sub(r'\s+', ' ', str(cell)).strip().lower()

    def _cell(row, i):
        if i is None or i < 0 or i >= len(row) or not row[i]: return ""
        return re.sub(r'\s+', ' ', str(row[i])).strip()

    def _detect_cols(header):
        h = [_norm(c) for c in header]
        # Manufacturer name column — avoid s.no / serial
        mfr = next((i for i,x in enumerate(h)
                    if any(k in x for k in ("manufactur","company name","name of manuf"))), None)
        if mfr is None:
            # fallback: 'name' but not s.no or serial
            mfr = next((i for i,x in enumerate(h)
                        if "name" in x and not any(k in x for k in ("s.no","s. no","serial","sl."))), None)
        if mfr is None:
            # last resort: col 1 (col 0 is almost always S.No.)
            mfr = 1

        model = next((i for i,x in enumerate(h) if "model" in x), mfr + 1)
        cap   = next((i for i,x in enumerate(h)
                      if any(k in x for k in ("rated power","watt","wp","capacity","power"))), model + 1)
        eff   = next((i for i,x in enumerate(h) if "effic" in x), -1)
        tech  = next((i for i,x in enumerate(h)
                      if any(k in x for k in ("tech","cell type","module type","technology"))), -1)
        val   = next((i for i,x in enumerate(h) if any(k in x for k in ("valid","expir","date"))), -1)
        return mfr, model, cap, eff, tech, val

    _SKIP = {"", "nan", "none", "name", "name of manufacturer", "manufacturer",
             "manufactur", "s.no", "s. no", "s.no.", "sl.no", "sl no", "serial no"}

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                if not table or len(table) < 2: continue

                # Find the actual header row (may be table[0] or table[1] after a merged title)
                # A title row has only 1 non-None cell; the real header has "manufactur" or "name"
                header_idx = 0
                for ri, row in enumerate(table[:4]):
                    h_norm = [_norm(c) for c in row]
                    if any(any(k in x for k in ("manufactur","name","model","wp","watt","rated","enlisted"))
                           for x in h_norm):
                        header_idx = ri
                        break

                raw_h = table[header_idx]
                col_mfr, col_model, col_cap, col_eff, col_tech, col_val = _detect_cols(raw_h)

                # Data starts after header + any sub-header rows
                # Skip rows until we find one whose first numeric-looking cell matches S.No.
                start = header_idx + 1
                # Skip sub-header rows (no non-None cell in col_mfr position)
                while start < len(table) and not _cell(table[start], col_mfr):
                    start += 1

                for row in table[start:]:
                    if not row or not any(row): continue
                    mfr = _clean_mfr(_cell(row, col_mfr))
                    if not mfr: continue
                    mfr_lo = mfr.lower().strip()
                    # Skip header-lookalike rows and serial numbers
                    if mfr_lo in _SKIP: continue
                    if re.match(r'^\d+\.?\s*$', mfr): continue  # pure S.No. value
                    if any(k in mfr_lo for k in ("name of manuf", "manufactur", "s.no")): continue

                    cap_str = _cell(row, col_cap)
                    cap_m   = re.search(r"([\d,]+(?:\.\d+)?)", cap_str.replace(",",""))
                    cap_wp  = float(cap_m.group(1).replace(",","")) if cap_m else 0.0

                    eff_str = _cell(row, col_eff) if col_eff is not None and col_eff >= 0 else ""
                    eff_m   = re.search(r"([\d]+(?:\.\d+)?)", eff_str)
                    eff     = float(eff_m.group(1)) if eff_m else 0.0

                    records.append({
                        "mfr":            mfr[:120],
                        "model":          _cell(row, col_model)[:120],
                        "capacity_wp":    cap_wp,
                        "efficiency":     eff,
                        "technology":     _cell(row, col_tech)[:60] if col_tech is not None and col_tech >= 0 else "",
                        "validity_date":  _cell(row, col_val)[:30]  if col_val  is not None and col_val  >= 0 else "",
                        "parent_company": _map_parent(mfr),
                    })
    return records

@serialized
def fetch_almm():
    """Return ALMM data from SQLite. Re-parse only if new PDF detected."""
    con = sqlite3.connect(DB_PATH)
    try:
        # Check if we have a recent parse (re-check PDF URL every 6hr)
        version_cached = get_cache("almm_version_check", 21600)
        if not version_cached:
            try:
                pdf_url, pub_date = _fetch_almm_pdf_url()
                set_cache("almm_version_check", {"pdf_url": pdf_url, "pub_date": pub_date})
            except Exception as e:
                pdf_url, pub_date = None, ""
                set_cache("almm_version_check", {"pdf_url": None, "pub_date": "", "error": str(e)})
        else:
            pdf_url  = version_cached.get("pdf_url")
            pub_date = version_cached.get("pub_date", "")

        # Check if this version is already parsed
        row = con.execute("SELECT id, record_count FROM almm_meta WHERE pdf_url=? ORDER BY parsed_at DESC LIMIT 1",
                          (pdf_url,)).fetchone() if pdf_url else None
        if not row and pdf_url:
            # Download + parse
            try:
                import tempfile, pdfplumber
                r = gov_get(pdf_url, timeout=60)
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                    tmp.write(r.content)
                    tmp_path = tmp.name
                records = _parse_almm_pdf(tmp_path)
                os.unlink(tmp_path)
                if records:
                    con.execute("DELETE FROM almm_list")
                    con.execute("DELETE FROM almm_meta")
                    con.execute("INSERT INTO almm_meta(pdf_url,pub_date,parsed_at,record_count) VALUES(?,?,?,?)",
                                (pdf_url, pub_date, time.time(), len(records)))
                    meta_id = con.execute("SELECT last_insert_rowid()").fetchone()[0]
                    for rec in records:
                        con.execute("INSERT INTO almm_list(mfr,model,capacity_wp,efficiency,technology,validity_date,parent_company,meta_id) VALUES(?,?,?,?,?,?,?,?)",
                                    (rec["mfr"],rec["model"],rec["capacity_wp"],rec["efficiency"],
                                     rec["technology"],rec["validity_date"],rec["parent_company"],meta_id))
                    con.commit()
                    mark_health("ALMM", True, f"{len(records)} records parsed")
            except Exception as e:
                mark_health("ALMM", False, str(e))

        # Aggregate: per-manufacturer capacity (MAX per mfr entity) then SUM per parent group
        # This correctly handles: Adani (Mundra Energy + Mundra PV), Premier (2 entities), etc.
        rows = con.execute("""
            SELECT parent_company, SUM(mfr_cap) as total_mw_yr, SUM(model_count) as models
            FROM (
                SELECT parent_company, mfr, MAX(capacity_wp) as mfr_cap, COUNT(*) as model_count
                FROM almm_list WHERE parent_company != ''
                GROUP BY parent_company, mfr
            )
            GROUP BY parent_company ORDER BY total_mw_yr DESC
        """).fetchall()
        meta = con.execute("SELECT pdf_url, pub_date, parsed_at, record_count FROM almm_meta ORDER BY parsed_at DESC LIMIT 1").fetchone()

        top20 = [{"parent": r[0], "models": r[2], "total_mw_yr": round(r[1], 0)} for r in rows[:20]]
        others_mw = sum(r[1] for r in rows[20:])
        if others_mw: top20.append({"parent": "Others", "models": sum(r[2] for r in rows[20:]), "total_mw_yr": round(others_mw, 0)})

        return {
            "top20":        top20,
            "pdf_url":      meta[0] if meta else None,
            "pub_date":     meta[1] if meta else None,
            "parsed_at":    datetime.fromtimestamp(meta[2]).strftime("%Y-%m-%d %H:%M") if meta else None,
            "record_count": meta[3] if meta else 0,
            "fetched_at":   datetime.now().strftime("%Y-%m-%d %H:%M"),
        }
    finally:
        con.close()

def check_almm_update():
    """Force re-check of ALMM PDF version. Returns version info."""
    cache.pop("almm_version_check", None)
    try:
        pdf_url, pub_date = _fetch_almm_pdf_url()
        con = sqlite3.connect(DB_PATH)
        row = con.execute("SELECT pub_date, parsed_at, record_count FROM almm_meta WHERE pdf_url=? LIMIT 1",
                          (pdf_url,)).fetchone()
        con.close()
        return {
            "pdf_url":    pdf_url,
            "pub_date":   pub_date,
            "already_parsed": bool(row),
            "last_parsed": datetime.fromtimestamp(row[1]).strftime("%Y-%m-%d %H:%M") if row else None,
        }
    except Exception as e:
        return {"error": str(e)}

@app.route("/api/news/archive")
def api_news_archive(): return jsonify(fetch_news_archive())

@app.route("/api/almm")
def api_almm(): return jsonify(fetch_almm())

@app.route("/api/almm/check_update")
def api_almm_check_update(): return jsonify(check_almm_update())

@app.route("/api/almm/force_reparse")
def api_almm_force_reparse():
    """Delete all cached ALMM data and re-parse from scratch."""
    try:
        con = sqlite3.connect(DB_PATH)
        con.execute("DELETE FROM almm_list")
        con.execute("DELETE FROM almm_meta")
        con.commit()
        con.close()
        cache.pop("almm_version_check", None)
        result = fetch_almm()
        return jsonify({"status": "ok", "records": result.get("record_count", 0),
                        "pub_date": result.get("pub_date"), "sample": result.get("top20", [])[:3]})
    except Exception as e:
        return jsonify({"status": "error", "error": str(e)})

# ── ALMM List-I — Solar PV Modules (P5.1) ────────────────────────────────────
def _fetch_almm_modules_pdf_url():
    hdr = {"User-Agent": "Mozilla/5.0 Chrome/124 Safari/537.36"}
    r   = gov_get(ALMM_PAGE_URL, timeout=15)
    soup = BeautifulSoup(r.text, "html.parser")
    for a in soup.find_all("a", href=True):
        txt  = a.get_text(strip=True)
        href = a["href"]
        if "list-i" in txt.lower() and "module" in txt.lower() and href.endswith(".pdf"):
            # Extract date from text like "Updated (01.05.2026)"
            m = re.search(r'\((\d{2}\.\d{2}\.\d{4})\)', txt)
            pub_date = m.group(1) if m else ""
            return href, pub_date
    raise ValueError("ALMM List-I PDF not found")

def _parse_almm_modules_pdf(pdf_path):
    import pdfplumber
    records = []
    current_mfr = ""
    current_cap = 0.0

    HDR_KEYWORDS = ("name of the manufact", "enlisted capacity", "enlisted models",
                    "module efficiency", "type of module", "s. no")

    def _nc(cell):
        if not cell: return ""
        return re.sub(r'\s+', ' ', str(cell)).strip()

    def _is_header(row):
        joined = " ".join(_nc(c).lower() for c in row if c)
        return any(k in joined for k in HDR_KEYWORDS)

    def _extract_wp(text):
        m = re.search(r'\((\d{3,4})\s*[Ww][Pp]?\)', text or "")
        return float(m.group(1)) if m else 0.0

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                if not table: continue
                for row in table:
                    if not row or not any(row): continue
                    if _is_header(row): continue
                    # Update current manufacturer when col 1 has data
                    mfr_cell = _nc(row[1]) if len(row) > 1 else ""
                    cap_cell  = _nc(row[4]) if len(row) > 4 else ""
                    mod_cell  = _nc(row[8]) if len(row) > 8 else ""
                    eff_cell  = _nc(row[9]) if len(row) > 9 else ""
                    typ_cell  = _nc(row[6]) if len(row) > 6 else ""
                    app_cell  = _nc(row[7]) if len(row) > 7 else ""
                    vfrom     = _nc(row[12]) if len(row) > 12 else ""
                    vto       = _nc(row[13]) if len(row) > 13 else ""

                    if mfr_cell and not re.match(r'^[A-Z]\.|^\d+$', mfr_cell):
                        # Looks like a real mfr name update
                        lo = mfr_cell.lower()
                        if not any(k in lo for k in ("name of", "s. no", "updated", "addition", "list-i", "almm")):
                            current_mfr = _clean_mfr(mfr_cell)[:120]
                    if cap_cell:
                        cm = re.search(r'([\d,]+(?:\.\d+)?)', cap_cell.replace(",", ""))
                        if cm:
                            try: current_cap = float(cm.group(1))
                            except: pass

                    if not current_mfr or not mod_cell: continue
                    # Skip section headers like "A.", "B."
                    if re.match(r'^[A-Z]\.\s', mod_cell): continue

                    eff_m = re.search(r'([\d]+(?:\.\d+)?)', eff_cell)
                    eff   = float(eff_m.group(1)) if eff_m else 0.0
                    wp    = _extract_wp(app_cell)

                    records.append({
                        "mfr":          current_mfr,
                        "model":        mod_cell[:80],
                        "capacity_mw_yr": current_cap,
                        "efficiency":   eff,
                        "module_type":  typ_cell[:60],
                        "module_wp":    wp,
                        "validity_from": vfrom[:20],
                        "validity_to":  vto[:30],
                        "parent_company": _map_parent(current_mfr),
                    })
    return records

@serialized
def fetch_almm_modules():
    con = sqlite3.connect(DB_PATH)
    try:
        version_cached = get_cache("almm_modules_version", 21600)
        if not version_cached:
            try:
                pdf_url, pub_date = _fetch_almm_modules_pdf_url()
                set_cache("almm_modules_version", {"pdf_url": pdf_url, "pub_date": pub_date})
            except Exception as e:
                pdf_url, pub_date = None, ""
                set_cache("almm_modules_version", {"pdf_url": None, "pub_date": "", "error": str(e)})
        else:
            pdf_url  = version_cached.get("pdf_url")
            pub_date = version_cached.get("pub_date", "")

        row = con.execute("SELECT id, record_count FROM almm_modules_meta WHERE pdf_url=? ORDER BY parsed_at DESC LIMIT 1",
                          (pdf_url,)).fetchone() if pdf_url else None
        if not row and pdf_url:
            try:
                import tempfile
                r = gov_get(pdf_url, timeout=120)
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                    tmp.write(r.content); tmp_path = tmp.name
                records = _parse_almm_modules_pdf(tmp_path)
                os.unlink(tmp_path)
                if records:
                    con.execute("DELETE FROM almm_modules")
                    con.execute("DELETE FROM almm_modules_meta")
                    con.execute("INSERT INTO almm_modules_meta(pdf_url,pub_date,parsed_at,record_count) VALUES(?,?,?,?)",
                                (pdf_url, pub_date, time.time(), len(records)))
                    meta_id = con.execute("SELECT last_insert_rowid()").fetchone()[0]
                    for rec in records:
                        con.execute("""INSERT INTO almm_modules(mfr,model,capacity_mw_yr,efficiency,
                                       module_type,module_wp,validity_from,validity_to,parent_company,meta_id)
                                       VALUES(?,?,?,?,?,?,?,?,?,?)""",
                                    (rec["mfr"],rec["model"],rec["capacity_mw_yr"],rec["efficiency"],
                                     rec["module_type"],rec["module_wp"],rec["validity_from"],
                                     rec["validity_to"],rec["parent_company"],meta_id))
                    con.commit()
                    mark_health("ALMM Modules", True, f"{len(records)} models parsed")
            except Exception as e:
                mark_health("ALMM Modules", False, str(e))

        # SUM of each distinct mfr entity's capacity per parent group
        # e.g. Goldi = Goldi Sun (1474) + Goldi Solar Pvt (1000) = 2474 MW
        # e.g. Waaree = Waaree Energies (1095) + FS India Solar (3212) = 4307 MW
        rows = con.execute("""
            SELECT parent_company,
                   SUM(mfr_cap)    as cap_mw_yr,
                   AVG(avg_eff)    as avg_eff,
                   AVG(avg_wp)     as avg_wp,
                   SUM(model_count) as models
            FROM (
                SELECT parent_company, mfr,
                       MAX(capacity_mw_yr) as mfr_cap,
                       AVG(efficiency)     as avg_eff,
                       AVG(module_wp)      as avg_wp,
                       COUNT(*)            as model_count
                FROM almm_modules WHERE parent_company != ''
                GROUP BY parent_company, mfr
            )
            GROUP BY parent_company ORDER BY cap_mw_yr DESC
        """).fetchall()
        meta = con.execute("SELECT pdf_url, pub_date, parsed_at, record_count FROM almm_modules_meta ORDER BY parsed_at DESC LIMIT 1").fetchone()

        top20 = [{"parent": r[0],
                  "capacity_mw_yr": round(r[1] or 0, 0),
                  "avg_eff": round(r[2] or 0, 2),
                  "avg_wp": round(r[3] or 0, 0),
                  "models": r[4]} for r in rows[:25]]

        return {
            "top20":        top20,
            "pdf_url":      meta[0] if meta else None,
            "pub_date":     meta[1] if meta else None,
            "parsed_at":    datetime.fromtimestamp(meta[2]).strftime("%Y-%m-%d %H:%M") if meta else None,
            "record_count": meta[3] if meta else 0,
            "fetched_at":   datetime.now().strftime("%Y-%m-%d %H:%M"),
        }
    finally:
        con.close()

@app.route("/api/almm/modules")
def api_almm_modules(): return jsonify(fetch_almm_modules())

def _norm_tech(raw):
    """Normalize raw ALMM module_type strings into canonical tech categories."""
    t = (raw or "").lower()
    if "cdte" in t or "cadmium" in t or "thin film" in t:
        return "CdTe Thin Film"
    if "hjt" in t or "heterojunction" in t:
        return "HJT"
    if "topcon" in t or "n-type" in t or "n type" in t or "ntype" in t:
        if "bifacial" in t:
            return "Bifacial N-Type TOPCon"
        return "N-Type TOPCon"
    if "bifacial" in t and "perc" in t:
        return "Bifacial Mono PERC"
    if "perc" in t or "mono" in t or "c-si" in t or "c si" in t:
        return "Mono PERC"
    if "poly" in t or "polycrystalline" in t:
        return "Polycrystalline"
    return "Other"

@app.route("/api/almm/tech_mix")
def api_almm_tech_mix():
    """Technology distribution from ALMM List-I (modules) and List-II (cells), normalised."""
    try:
        con = sqlite3.connect(DB_PATH)
        mod_rows = con.execute("""
            SELECT module_type, COUNT(*) as cnt,
                   SUM(capacity_mw_yr) as cap_mw
            FROM almm_modules WHERE module_type != '' AND module_type IS NOT NULL
            GROUP BY module_type
        """).fetchall()
        cell_rows = con.execute("""
            SELECT technology, COUNT(*) as cnt
            FROM almm_list WHERE technology != '' AND technology IS NOT NULL
            GROUP BY technology
        """).fetchall()
        con.close()
        # Aggregate into canonical categories
        mod_agg = {}
        for tech_raw, cnt, cap in mod_rows:
            canon = _norm_tech(tech_raw)
            if canon not in mod_agg: mod_agg[canon] = {"count": 0, "cap_mw": 0}
            mod_agg[canon]["count"] += cnt
            mod_agg[canon]["cap_mw"] += (cap or 0)
        cell_agg = {}
        for tech_raw, cnt in cell_rows:
            canon = _norm_tech(tech_raw)
            if canon not in cell_agg: cell_agg[canon] = 0
            cell_agg[canon] += cnt
        return jsonify({
            "modules": sorted([{"tech": k, "count": v["count"], "cap_mw": round(v["cap_mw"])}
                               for k, v in mod_agg.items()], key=lambda x: -x["count"]),
            "cells":   sorted([{"tech": k, "count": v} for k, v in cell_agg.items()], key=lambda x: -x["count"]),
        })
    except Exception as e:
        return jsonify({"error": str(e), "modules": [], "cells": []})

@app.route("/api/almm/modules/force_reparse")
def api_almm_modules_force():
    try:
        con = sqlite3.connect(DB_PATH)
        con.execute("DELETE FROM almm_modules"); con.execute("DELETE FROM almm_modules_meta")
        con.commit(); con.close()
        cache.pop("almm_modules_version", None)
        result = fetch_almm_modules()
        return jsonify({"status": "ok", "records": result.get("record_count", 0)})
    except Exception as e:
        return jsonify({"status": "error", "error": str(e)})

@app.route("/api/telegram/setup")
def api_telegram_setup():
    """Auto-discover chat_id from getUpdates — call AFTER sending the bot any message."""
    if not TELEGRAM_BOT_TOKEN:
        # Degrade-never-break: don't fire a doomed external call when unconfigured.
        return jsonify({"ok": False, "configured": False,
                        "hint": "Set TELEGRAM_BOT_TOKEN in .env first (no external call made)."})
    try:
        r = requests.get(
            f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/getUpdates",
            timeout=8
        )
        data = r.json()
        results = data.get("result", [])
        if not results:
            return jsonify({"ok": False, "hint": "Send any message to the bot first, then call this endpoint."})
        # Take the most recent update's chat_id
        chat_id = str(results[-1]["message"]["chat"]["id"])
        _save_chat_id(chat_id)
        # Send confirmation
        ok = send_telegram(
            f"🔆 <b>NEURON v7 — Telegram Setup Complete</b>\n"
            f"Chat ID <code>{chat_id}</code> has been saved.\n"
            f"<i>  the Fool now has a channel above the fog.</i>"
        )
        return jsonify({"ok": True, "chat_id": chat_id, "message_sent": ok})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})

@app.route("/api/telegram/test")
def api_telegram_test():
    """Send a test Telegram message and return full diagnostic info."""
    chat_id = _get_chat_id()
    url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage"
    msg = (
        "⚡ NEURON v7 — Telegram Active\n"
        "Bot: @Neucardbot · User: Selenophileus\n"
        "  the Fool sees all above the fog.\n"
        "Alerts: Intel Pulse ≥80, SECI new tenders, Daily 9am IST"
    )
    try:
        r = requests.post(url, json={"chat_id": chat_id, "text": msg}, timeout=10)
        ok = r.status_code == 200
        detail = r.json() if r.content else {}
    except Exception as e:
        ok = False; detail = {"error": str(e)}
    return jsonify({"sent": ok, "chat_id_active": chat_id or "not set",
                    "token_prefix": TELEGRAM_BOT_TOKEN[:12]+"…",
                    "api_response": detail,
                    "hint": "OK" if ok else "Check api_response for Telegram error"})

@app.route("/api/telegram/config")
def api_telegram_config():
    chat_id = _get_chat_id()
    return jsonify({"configured": bool(TELEGRAM_BOT_TOKEN and chat_id),
                    "token_set": bool(TELEGRAM_BOT_TOKEN),
                    "chat_id": chat_id or None,
                    "setup_url": "/api/telegram/setup"})

# ════════════════════════════════════════════════════════════════════════════
# P7 — WorldMonitor Transfer Features
# ════════════════════════════════════════════════════════════════════════════

# ── P7.2 RE Sector Breadth (% of RE stocks above SMA) ───────────────────────
def fetch_sector_breadth():
    cached = get_cache("sector_breadth", 3600)
    if cached: return cached
    try:
        total = 0; above20 = 0; above50 = 0; above200 = 0
        stock_detail = []
        for sym in list(RE_STOCKS.keys()):
            try:
                h = yf.Ticker(sym).history(period="1y")
                if h.empty or len(h) < 21: continue
                closes = [float(v) for v in h["Close"].values]
                price  = closes[-1]
                sma20  = float(np.mean(closes[-20:]))
                sma50  = float(np.mean(closes[-50:])) if len(closes) >= 50 else None
                sma200 = float(np.mean(closes[-200:])) if len(closes) >= 200 else None
                total += 1
                if price > sma20: above20 += 1
                if sma50  and price > sma50:  above50  += 1
                if sma200 and price > sma200: above200 += 1
                stock_detail.append({
                    "sym": sym, "name": RE_STOCKS.get(sym, sym),
                    "price": round(price, 2),
                    "above_sma20":  price > sma20,
                    "above_sma50":  bool(sma50  and price > sma50),
                    "above_sma200": bool(sma200 and price > sma200),
                })
            except Exception:
                pass
        b20 = round(above20  / total * 100, 1) if total else 0
        b50 = round(above50  / total * 100, 1) if total else 0
        b200= round(above200 / total * 100, 1) if total else 0
        # Store daily snapshot
        try:
            _c = sqlite3.connect(DB_PATH)
            snap_date = datetime.now().strftime("%Y-%m-%d")
            if not _c.execute("SELECT 1 FROM breadth_history WHERE snap_date=?", (snap_date,)).fetchone():
                _c.execute("INSERT INTO breadth_history(ts,snap_date,breadth_20,breadth_50,breadth_200) VALUES(?,?,?,?,?)",
                           (time.time(), snap_date, b20, b50, b200))
                _c.commit()
            _c.close()
        except Exception:
            pass
        result = {"sma20": b20, "sma50": b50, "sma200": b200, "total": total,
                  "stocks": stock_detail, "fetched_at": datetime.now().strftime("%H:%M")}
        set_cache("sector_breadth", result)
        return result
    except Exception as e:
        return {"error": str(e), "sma20": 0, "sma50": 0, "sma200": 0, "total": 0}

@app.route("/api/sector_breadth")
def api_sector_breadth(): return jsonify(fetch_sector_breadth())

@app.route("/api/breadth_history")
def api_breadth_history():
    try:
        con  = sqlite3.connect(DB_PATH)
        rows = con.execute(
            "SELECT snap_date,breadth_20,breadth_50,breadth_200 FROM breadth_history ORDER BY snap_date DESC LIMIT 60"
        ).fetchall()
        con.close()
        return jsonify([{"date":r[0],"sma20":r[1],"sma50":r[2],"sma200":r[3]} for r in reversed(rows)])
    except Exception as e:
        return jsonify({"error": str(e)})

# ── P7.3 RE Fear & Greed Index ───────────────────────────────────────────────
def compute_fear_greed():
    cached = get_cache("fear_greed", 3600)
    if cached: return cached
    scores = {}
    # 1. Sector breadth — % above SMA50 → 0-25 pts
    try:
        br = fetch_sector_breadth()
        scores["breadth"] = round(br.get("sma50", 50) / 100 * 25, 1)
    except Exception:
        scores["breadth"] = 12.5
    # 2. Intel pulse → 0-25 pts
    try:
        ie = fetch_intel_engine()
        scores["pulse"] = round(min(ie.get("industry_pulse", 50), 100) / 100 * 25, 1)
    except Exception:
        scores["pulse"] = 12.5
    # 3. Nifty Energy RSI → 0-15 pts
    try:
        h = yf.Ticker("^CNXENERGY").history(period="1mo")
        if not h.empty and len(h) >= 15:
            closes = np.array([float(v) for v in h["Close"].values])
            delta  = np.diff(closes)
            gain   = np.where(delta > 0, delta, 0)
            loss   = np.where(delta < 0, -delta, 0)
            avg_g  = np.mean(gain[-14:]); avg_l = np.mean(loss[-14:])
            rsi    = 100 - (100 / (1 + avg_g / avg_l)) if avg_l else 50.0
            scores["rsi"] = round(float(rsi) / 100 * 15, 1)
        else:
            scores["rsi"] = 7.5
    except Exception:
        scores["rsi"] = 7.5
    # 4. Brent direction — high oil = RE more competitive → 0-15 pts
    try:
        comm  = fetch_commodities()
        chg_p = (comm.get("CL=F",{}) or {}).get("change_pct", 0) or 0
        scores["brent"] = round(max(0, min(15, 7.5 + chg_p * 0.5)), 1)
    except Exception:
        scores["brent"] = 7.5
    # 5. SECI pipeline MW → 0-20 pts
    try:
        seci = fetch_seci_tenders()
        mw   = seci.get("total_mw", 0) or 0
        scores["seci"] = round(min(mw / 10000 * 20, 20), 1)
    except Exception:
        scores["seci"] = 10.0

    total = round(sum(scores.values()), 1)
    if   total >= 75: label = "EXTREME GREED"
    elif total >= 55: label = "GREED"
    elif total >= 45: label = "NEUTRAL"
    elif total >= 25: label = "FEAR"
    else:             label = "EXTREME FEAR"

    result = {"score": total, "label": label, "components": scores,
              "fetched_at": datetime.now().strftime("%H:%M")}
    try:
        _c = sqlite3.connect(DB_PATH)
        snap = datetime.now().strftime("%Y-%m-%d")
        if not _c.execute("SELECT 1 FROM fear_greed_history WHERE snap_date=?", (snap,)).fetchone():
            _c.execute("INSERT INTO fear_greed_history(ts,snap_date,score,label,breadth,pulse_dir,rsi_nifty,brent_dir) VALUES(?,?,?,?,?,?,?,?)",
                       (time.time(), snap, total, label,
                        scores.get("breadth",0), scores.get("pulse",0),
                        scores.get("rsi",0), scores.get("brent",0)))
            _c.commit()
        _c.close()
    except Exception:
        pass
    set_cache("fear_greed", result)
    return result

@app.route("/api/fear_greed")
def api_fear_greed(): return jsonify(compute_fear_greed())

@app.route("/api/fear_greed/history")
def api_fear_greed_history():
    try:
        con  = sqlite3.connect(DB_PATH)
        rows = con.execute("SELECT snap_date,score,label FROM fear_greed_history ORDER BY snap_date DESC LIMIT 60").fetchall()
        con.close()
        return jsonify([{"date":r[0],"score":r[1],"label":r[2]} for r in reversed(rows)])
    except Exception as e:
        return jsonify({"error": str(e)})

# ── P7.10 RE Market Implications — Rule Engine ───────────────────────────────
_IMPLICATION_RULES = [
    ("TARIFF_SIGNAL",       ["bcd","basic customs duty","anti-dumping"],
     "WAAREEENER.NS","LONG","HIGH","1-3 months",
     "BCD/AD enforcement tightens ALMM-certified domestic module advantage — margin expansion for Indian mfrs",
     ["BCD enforcement","→ imported panels expensive","→ domestic mfrs gain share","→ margin expansion"]),
    ("TARIFF_SIGNAL",       ["bcd","basic customs duty","anti-dumping"],
     "PREMIERENE.NS","LONG","HIGH","1-3 months",
     "BCD enforcement benefits domestic solar mfrs — Premier positioned as ALMM-listed supplier",
     ["BCD enforcement","→ domestic module demand","→ Premier order book growth"]),
    ("SUPPLY_CHAIN",        ["polysilicon","module price","cell price","input cost"],
     "ADANIGREEN.NS","SHORT","MEDIUM","1-6 months",
     "Rising module costs squeeze project IRR for large IPPs with near-term commissioning",
     ["Module costs up","→ project IRR falls","→ tariff discovery pressure","→ IPP margins compress"]),
    ("POLICY_NOTIFICATION", ["rpo","renewable purchase obligation","solar obligation"],
     "IREDA.NS","LONG","MEDIUM","3-6 months",
     "RPO targets drive discoms to procure RE capacity — IREDA loan book expansion",
     ["RPO mandate","→ discom procurement","→ IREDA financing demand","→ AUM growth"]),
    ("COMMISSIONING",       ["commissioned","operationalised","goes live","inaugurates"],
     "NHPC.NS","LONG","MEDIUM","immediate",
     "New capacity commissioning confirms execution — revenue visibility improves",
     ["Commissioning event","→ power off-take begins","→ CUF-based revenue lock-in"]),
    ("FUNDING",             ["ireda","green bond","ncd","debt","loan"],
     "IREDA.NS","LONG","HIGH","1-3 months",
     "Green financing surge signals policy support — IREDA as primary conduit benefits",
     ["Funding flow","→ IREDA disbursements","→ NIM expansion"]),
    ("EXPANSION",           ["giga","gigawatt","new plant","manufacturing","greenfield"],
     "WAAREEENER.NS","LONG","MEDIUM","6-12 months",
     "Manufacturing capacity expansion underpins ability to meet ALMM demand at scale",
     ["Giga expansion","→ ALMM capacity grows","→ order fulfillment","→ revenue upside"]),
    ("GLOBAL_MACRO",        ["china","polysilicon","module oversupply","panel price"],
     "BORORENEW.NS","SHORT","LOW","3-6 months",
     "Global module oversupply from China depresses module ASP — solar glass pricing pressure",
     ["China oversupply","→ global module ASP falls","→ solar glass demand stable but pricing pressure"]),
    ("TARIFF_SIGNAL",       ["tariff","₹/kwh","l1","bid tariff"],
     "ADANIGREEN.NS","LONG","MEDIUM","3-12 months",
     "Falling L1 tariff discovery increases project viability — large IPPs gain pipeline advantage",
     ["Low tariff discovery","→ more projects viable","→ pipeline conversion","→ capacity growth"]),
    ("REGULATORY_RISK",     ["curtailment","grid","discom default","payment delay"],
     "ADANIGREEN.NS","SHORT","HIGH","immediate",
     "Grid curtailment / DISCOM default risk materially impacts RE revenue realization",
     ["Curtailment event","→ PLF/CUF falls","→ revenue shortfall","→ IPP cash flow stress"]),
]

def fetch_re_implications():
    cached = get_cache("re_implications", 1800)
    if cached: return cached
    try:
        ie     = fetch_intel_engine()
        stream = ie.get("intel_stream", [])
        cards  = []; seen = set()
        for item in stream[:15]:
            cat   = item.get("category","")
            title = item.get("title","").lower()
            for (rule_cat, kws, ticker, direction, confidence, timeframe, narrative, chain) in _IMPLICATION_RULES:
                key = f"{rule_cat}_{ticker}"
                if key in seen: continue
                if cat == rule_cat and any(k in title for k in kws):
                    seen.add(key)
                    cards.append({
                        "ticker":     ticker,
                        "name":       RE_STOCKS.get(ticker, ticker.replace(".NS","")),
                        "direction":  direction,
                        "confidence": confidence,
                        "timeframe":  timeframe,
                        "narrative":  narrative,
                        "chain":      chain,
                        "trigger":    item.get("title","")[:80],
                        "score":      item.get("score", 50),
                    })
            if len(cards) >= 8: break
        result = {"cards": cards, "count": len(cards),
                  "fetched_at": datetime.now().strftime("%H:%M IST")}
        set_cache("re_implications", result)
        return result
    except Exception as e:
        return {"error": str(e), "cards": []}

@app.route("/api/re_implications")
def api_re_implications(): return jsonify(fetch_re_implications())

# ── P7.16 RE Regime Panel ─────────────────────────────────────────────────────
def fetch_re_regime():
    cached = get_cache("re_regime", 3600)
    if cached: return cached
    try:
        ie     = fetch_intel_engine()
        seci   = fetch_seci_tenders()
        pulse  = ie.get("industry_pulse", 50)
        stream = ie.get("intel_stream", [])
        cross  = ie.get("cross_signals", [])
        cat_counts = {}
        for item in stream:
            c = item.get("category","")
            cat_counts[c] = cat_counts.get(c, 0) + 1
        commission_n = cat_counts.get("COMMISSIONING", 0)
        tender_n     = cat_counts.get("TENDER_ISSUED", 0)
        policy_n     = cat_counts.get("POLICY_NOTIFICATION", 0)
        funding_n    = cat_counts.get("FUNDING", 0)
        risk_n       = cat_counts.get("REGULATORY_RISK", 0) + cat_counts.get("SUPPLY_CHAIN", 0)
        if pulse >= 70 and tender_n >= 2 and commission_n >= 1:
            regime, regime_color = "EXPANSION PHASE", "#6a9a3a"
        elif pulse >= 60 and (funding_n >= 2 or commission_n >= 2):
            regime, regime_color = "GROWTH SURGE", "#c4922a"
        elif policy_n >= 3:
            regime, regime_color = "POLICY TRANSITION", "#4a9eff"
        elif risk_n >= 3 or pulse < 35:
            regime, regime_color = "STRESS PHASE", "#c23535"
        elif pulse >= 50:
            regime, regime_color = "CONSOLIDATION", "#e8a030"
        else:
            regime, regime_color = "CONTRACTION", "#c23535"
        policy_score  = min(10, policy_n  * 2 + (1 if cross else 0))
        finance_score = min(10, funding_n * 3)
        grid_score    = max(0,  10 - risk_n * 3)
        actors = {}
        for item in stream:
            co = item.get("entities",{}).get("company","")
            if co and co in RE_STOCKS:
                actors[RE_STOCKS[co]] = actors.get(RE_STOCKS[co], 0) + 1
        top_actors = sorted(actors.items(), key=lambda x: x[1], reverse=True)[:4]
        watchpoints = []
        if seci.get("count", 0) > 0:
            watchpoints.append(f"SECI {seci['count']} tenders active · {int(seci.get('total_mw',0))}MW pipeline")
        if risk_n > 0:
            watchpoints.append("Regulatory/Supply chain risk items in feed")
        if policy_n >= 2:
            watchpoints.append(f"{policy_n} MNRE policy notifications in cycle")
        result = {
            "regime": regime, "regime_color": regime_color, "pulse": pulse,
            "policy_score": policy_score, "finance_score": finance_score, "grid_score": grid_score,
            "actors":  [{"name": n, "count": c} for n, c in top_actors],
            "watchpoints": watchpoints[:3],
            "drivers": ie.get("hot_topics", []),
            "fetched_at": datetime.now().strftime("%H:%M IST"),
        }
        set_cache("re_regime", result)
        return result
    except Exception as e:
        return {"error": str(e), "regime": "UNKNOWN", "regime_color": "#888"}

@app.route("/api/re_regime")
def api_re_regime(): return jsonify(fetch_re_regime())

# ── P7.11 RE Event Probability Tracker ───────────────────────────────────────
def fetch_re_forecast():
    cached = get_cache("re_forecast", 7200)
    if cached: return cached
    try:
        seci   = fetch_seci_tenders()
        mnre_n = fetch_mnre_notifications()
        ie     = fetch_intel_engine()
        pulse  = ie.get("industry_pulse", 50)
        stream = ie.get("intel_stream", [])
        tenders  = seci.get("tenders", [])
        policy_n = sum(1 for i in stream if i.get("category") == "POLICY_NOTIFICATION")
        fund_n   = sum(1 for i in stream if i.get("category") == "FUNDING")
        notifs   = mnre_n.get("items", [])
        def _parse_dl(s):
            if not s: return None
            for fmt in ("%d/%m/%Y","%d-%m-%Y","%Y-%m-%d","%d.%m.%Y"):
                try: return datetime.strptime(s.strip(), fmt)
                except: pass
            return None
        now = datetime.now()
        within_30 = sum(1 for t in tenders if _parse_dl(t.get("deadline","")) and
                        0 <= (_parse_dl(t["deadline"]) - now).days <= 30)
        within_90 = sum(1 for t in tenders if _parse_dl(t.get("deadline","")) and
                        0 <= (_parse_dl(t["deadline"]) - now).days <= 90)
        forecasts = [
            {"event":"SECI Tender L1 Discovery","p30":min(95,40+within_30*15),
             "p90":min(95,60+within_90*10),"p180":75,
             "signal":f"{within_30} tenders deadline ≤30d · {within_90} ≤90d","color":"#2aa198"},
            {"event":"MNRE Policy Notification","p30":min(90,30+policy_n*10),
             "p90":min(90,55+policy_n*8),"p180":80,
             "signal":f"{policy_n} notifications in cycle · {len(notifs)} total","color":"#4a9eff"},
            {"event":"New ALMM List Update","p30":30,"p90":55,"p180":80,
             "signal":"Typical cycle: quarterly — historical MNRE pattern","color":"#c4922a"},
            {"event":"Capital Raise / Green Bond","p30":min(85,25+fund_n*15),
             "p90":min(90,50+fund_n*10),"p180":75,
             "signal":f"{fund_n} funding events in current intel cycle","color":"#6a9a3a"},
            {"event":"Sector Pulse ≥ 70 (Bullish)","p30":min(80,20+max(0,pulse-50)*2),
             "p90":min(85,40+max(0,pulse-40)),"p180":60,
             "signal":f"Current pulse: {pulse} · Baseline: 50","color":"#6c71c4"},
        ]
        result = {"forecasts": forecasts, "fetched_at": datetime.now().strftime("%H:%M IST")}
        set_cache("re_forecast", result)
        return result
    except Exception as e:
        return {"error": str(e), "forecasts": []}

@app.route("/api/re_forecast")
def api_re_forecast(): return jsonify(fetch_re_forecast())

# ── P7.15 Sentiment Spread from pulse_history ─────────────────────────────────
@app.route("/api/sentiment_spread")
def api_sentiment_spread():
    try:
        con  = sqlite3.connect(DB_PATH)
        rows = con.execute(
            """SELECT date, pulse, label,
                      COALESCE(bullish_count,0), COALESCE(bearish_count,0), COALESCE(neutral_count,0)
               FROM pulse_history ORDER BY date DESC LIMIT 60"""
        ).fetchall()
        con.close()
        data = []
        for r in reversed(rows):
            total    = (r[3] or 0) + (r[4] or 0) + (r[5] or 0)
            bull_pct = round(r[3]/total*100, 1) if total else 0
            bear_pct = round(r[4]/total*100, 1) if total else 0
            spread   = round(bull_pct - bear_pct, 1)
            data.append({"date":r[0],"pulse":r[1],"label":r[2],
                          "bullish_count":r[3],"bearish_count":r[4],"neutral_count":r[5],
                          "bull_pct":bull_pct,"bear_pct":bear_pct,"spread":spread})
        return jsonify(data)
    except Exception as e:
        return jsonify({"error": str(e)})

# ── P7.18 RE Stock News Velocity ──────────────────────────────────────────────
@app.route("/api/re_velocity")
def api_re_velocity():
    cached = get_cache("re_velocity", 3600)
    if cached: return jsonify(cached)
    try:
        company_news = fetch_company_news()
        results = {}
        for sym, news_list in company_news.items():
            total = len(news_list)
            if not total: continue
            recent_count = 0
            for n in news_list:
                pub = n.get("published","")
                if pub:
                    try:
                        dt = datetime.strptime(pub, "%Y-%m-%d %H:%M")
                        if (datetime.now() - dt).total_seconds() <= 7200:
                            recent_count += 1
                    except Exception:
                        pass
            baseline = max(1.0, total / 5.0)
            ratio    = recent_count / baseline
            results[sym] = {
                "name":     RE_STOCKS.get(sym, sym),
                "recent":   recent_count,
                "total":    total,
                "baseline": round(baseline, 1),
                "ratio":    round(ratio, 2),
                "spike":    ratio >= 2.0,
                "alert":    ratio >= 5.0,
            }
        set_cache("re_velocity", results)
        return jsonify(results)
    except Exception as e:
        return jsonify({"error": str(e)})

# ── P7.17 ALMM Modules Full Search ───────────────────────────────────────────
@app.route("/api/almm/modules/search")
def api_almm_modules_search():
    q        = (request.args.get("q","") or "").strip().lower()
    sort_by  = request.args.get("sort","efficiency")
    sort_dir = request.args.get("dir","desc")
    mfr_filt = (request.args.get("mfr","") or "").strip().lower()
    page     = max(1, int(request.args.get("page","1") or "1"))
    per_page = 50
    try:
        con = sqlite3.connect(DB_PATH)
        where_clauses = []; params = []
        if q:
            where_clauses.append("(LOWER(model) LIKE ? OR LOWER(mfr) LIKE ? OR LOWER(parent_company) LIKE ?)")
            params += [f"%{q}%", f"%{q}%", f"%{q}%"]
        if mfr_filt:
            where_clauses.append("LOWER(parent_company) LIKE ?")
            params.append(f"%{mfr_filt}%")
        where_sql = ("WHERE " + " AND ".join(where_clauses)) if where_clauses else ""
        order_map = {"efficiency":"efficiency","capacity":"module_wp","validity":"validity_to","mfr":"parent_company"}
        order_col = order_map.get(sort_by, "efficiency")
        order_dir = "DESC" if sort_dir == "desc" else "ASC"
        offset    = (page - 1) * per_page
        total     = con.execute(f"SELECT COUNT(*) FROM almm_modules {where_sql}", params).fetchone()[0]
        rows      = con.execute(
            f"SELECT mfr,model,capacity_mw_yr,efficiency,module_type,module_wp,validity_from,validity_to,parent_company "
            f"FROM almm_modules {where_sql} ORDER BY {order_col} {order_dir} LIMIT ? OFFSET ?",
            params + [per_page, offset]
        ).fetchall()
        con.close()
        return jsonify({
            "total": total, "page": page, "per_page": per_page,
            "records": [{"mfr":r[0],"model":r[1],"capacity_mw_yr":r[2],"efficiency":r[3],
                          "module_type":r[4],"module_wp":r[5],"validity_from":r[6],
                          "validity_to":r[7],"parent_company":r[8]} for r in rows],
        })
    except Exception as e:
        return jsonify({"error": str(e)})

# ── v21 Tender Intelligence + Watch API ─────────────────────────────────────
try:
    import tender_intel as _ti
    _ti_ok = True
except Exception as _ti_err:
    _ti_ok = False
    print(f"  [..] tender_intel import failed: {_ti_err}")

@app.route("/api/tender_intel/stats")
def api_ti_stats():
    if not _ti_ok: return jsonify({"error":"tender_intel not loaded"})
    try: return jsonify(_fix_nan(_ti.get_stats()))
    except Exception as e: return jsonify({"error":str(e)})

@app.route("/api/tender_intel/tenders")
def api_ti_tenders():
    if not _ti_ok: return jsonify({"tenders":[],"count":0,"total_mw":0,"sectors":[]})
    try:
        return jsonify(_fix_nan(_ti.get_tenders(
            sector=request.args.get("sector"),
            entity_type=request.args.get("entity_type"),
            status=request.args.get("status"),
            limit=int(request.args.get("limit",100))
        )))
    except Exception as e: return jsonify({"error":str(e),"tenders":[]})

@app.route("/api/tender_intel/capacity_pipeline")
def api_ti_pipeline():
    if not _ti_ok: return jsonify({})
    try: return jsonify(_ti.get_capacity_pipeline())
    except Exception as e: return jsonify({"error":str(e)})

@app.route("/api/tender_intel/ingest", methods=["POST"])
def api_ti_ingest():
    if not _ti_ok: return jsonify({"ok":False,"error":"tender_intel not loaded"})
    try:
        arts = []
        # Pull recent articles from RSS cache / v11 sources
        try:
            con = sqlite3.connect(DB_PATH)
            rows = con.execute(
                "SELECT title,link,summary,source_id,published_at FROM v11_articles ORDER BY id DESC LIMIT 300"
            ).fetchall()
            con.close()
            arts = [{"title":r[0],"link":r[1],"summary":r[2],"source":r[3],"date":r[4]} for r in rows if r[0]]
        except Exception:
            pass
        result = _ti.scan_and_ingest(arts)
        return jsonify({"ok":True,**result})
    except Exception as e: return jsonify({"ok":False,"error":str(e)})

@app.route("/api/watch/companies")
def api_watch_companies():
    if not _ti_ok: return jsonify([])
    try: return jsonify(_ti.get_watch_companies())
    except Exception as e: return jsonify({"error":str(e)})

@app.route("/api/watch/causal_chains")
def api_watch_causal():
    if not _ti_ok: return jsonify([])
    try:
        return jsonify(_ti.get_causal_chains(
            limit=int(request.args.get("limit",20)),
            company=request.args.get("company")
        ))
    except Exception as e: return jsonify({"error":str(e)})

@app.route("/api/watch/anomalies")
def api_watch_anomalies():
    if not _ti_ok: return jsonify([])
    try: return jsonify(_ti.get_anomalies())
    except Exception as e: return jsonify({"error":str(e)})

@app.route("/v21")
@app.route("/index_v21")
def index_v21():
    # v21 template retired — redirect to main cockpit
    from flask import redirect
    return redirect("/", code=301)

def _daily_brief_worker():
    """Background thread: sends daily 9am brief to Telegram."""
    import time as _t
    while True:
        now = datetime.now()
        # Target 09:00 IST daily
        target = now.replace(hour=9, minute=0, second=0, microsecond=0)
        if now >= target:
            target += timedelta(days=1)   # safe across month/year boundaries
        sleep_secs = (target - now).total_seconds()
        _t.sleep(max(1, sleep_secs))
        try:
            # P15 A5 — the morning brief is the Synthesis Desk's output (always
            # available: LLM when a key is alive, heuristic otherwise), prefixed
            # with last night's consolidation memo. Prior bug: this called an
            # undefined fetch_brief() → silent NameError, so the brief never sent.
            data = v11_intel.synthesis_brief()
            text = data.get("brief", "") if isinstance(data, dict) else str(data)
            memo = v11_sources.kv_get("night_memo")
            if memo:
                text = f"📋 {memo}\n\n{text}"
            if text and len(text) > 20:
                send_telegram(f"🌅 <b>NEURON Morning Brief</b>\n\n{text[:3500]}")
        except Exception:
            pass

def _init_embedder():
    """P16.2 — upgrade MemoryOS to real semantics with a local ONNX embedder
    (fastembed bge-small-en-v1.5, 384-dim). Runs in the expression layer (I/O
    allowed); memory.py itself stays import-clean. Degrades to the numpy floor
    embedder if fastembed/model is unavailable — recall never breaks."""
    try:
        from fastembed import TextEmbedding
        _m = TextEmbedding(model_name="BAAI/bge-small-en-v1.5")
        def _one(t): return next(iter(_m.embed([t or ""])))
        def _many(ts): return list(_m.embed([t or "" for t in ts]))
        v16_mem.set_embedder(_one, "bge-small-en-v1.5", batch_fn=_many)
        print("  [OK] memory embedder: BAAI/bge-small-en-v1.5 (384-dim, local ONNX)")
        return True
    except Exception as e:
        print("  [..] memory embedder: floor-hash (fastembed unavailable: "
              + str(e)[:60] + ")")
        return False

def boot_diagnostics():
    """P15 A4/C3 — phased startup checks. Fast (DB + file reads, milliseconds),
    run synchronously before serving; results logged loudly and stored for
    /api/health. We deliberately do NOT block request serving on the worker
    heartbeat (that would delay the instant-boot UX) — heartbeat liveness is
    surfaced in /api/health instead, so a dead worker is visible, never silent."""
    report = {"phases": [], "ts": datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
    ok_all = True
    # Phase 1 — DB schema integrity
    try:
        con = sqlite3.connect(DB_PATH, timeout=15)
        jm = con.execute("PRAGMA journal_mode").fetchone()[0]
        have = {r[0] for r in con.execute("SELECT name FROM sqlite_master WHERE type='table'")}
        con.close()
        need = {"v11_articles", "v14_entity_ledger", "cea_national_snap", "kv_store", "v11_kv"}
        missing = sorted(need - have)
        p = {"phase": "db_schema", "ok": not missing, "journal_mode": jm,
             "missing_tables": missing, "table_count": len(have)}
    except Exception as e:
        p = {"phase": "db_schema", "ok": False, "error": str(e)[:120]}
    ok_all = ok_all and p["ok"]; report["phases"].append(p)
    # Phase 2 — user_data Excel integrity (don't silently fall through to Playwright)
    ud = os.path.join(os.path.dirname(__file__), "user_data")
    for fname, expected in (
            ("pm_kusum.xlsx", ["State_Name", "Total_Sanction_MW", "Total_Installed_MW"]),
            ("pm_surya_ghar.xlsx", ["State_Name", "Applications_Registered", "Sanctioned", "Installed"])):
        fp = os.path.join(ud, fname)
        try:
            import openpyxl
            wb = openpyxl.load_workbook(fp, read_only=True)
            ws = wb.active
            hdr = [c.value for c in next(ws.iter_rows(max_row=1))]
            wb.close()
            pf = {"phase": f"user_data/{fname}", "ok": all(h in hdr for h in expected),
                  "headers": hdr, "expected": expected}
        except Exception as e:
            pf = {"phase": f"user_data/{fname}", "ok": False, "error": str(e)[:120]}
        ok_all = ok_all and pf["ok"]; report["phases"].append(pf)
    report["ok"] = ok_all
    print("-" * 62)
    print("  BOOT DIAGNOSTICS (P15 A4/C3)")
    for p in report["phases"]:
        flag = "OK" if p.get("ok") else "!!"
        extra = p.get("error") or (("missing " + ",".join(p["missing_tables"]))
                                    if p.get("missing_tables") else "")
        print(f"  [{flag}] {p['phase']}" + (f"  -- {extra}" if extra else ""))
    if not ok_all:
        print("  WARNING: a boot check FAILED -- see /api/health. Affected fetchers")
        print("           may degrade (e.g. KUSUM/Surya Ghar fall back to live scrape).")
    print("-" * 62)
    try: v11_sources.kv_set("boot_diagnostics", json.dumps(report))
    except Exception: pass
    return report

def _consolidation_worker():
    """P15 B4 — the sleep cycle. Runs the daily consolidation pass at ~02:30
    local each day; also runs once shortly after boot if today's delta is
    missing, so the brain always has a current picture."""
    import time as _t
    _t.sleep(120)                       # let ingestion warm up first
    try:
        v11_cog.get_today_delta()       # compute today's delta if not present
    except Exception:
        pass
    try:
        v16_mem.ingest_recent()         # P16 — curate new high-signal facts
    except Exception:
        pass
    try:
        fetch_india_macro_plus()        # P20 — seed IMF macro/trade beliefs
    except Exception:
        pass
    try:
        # P17 — record today's decisions to the ledger + resolve matured ones.
        _ctx = _decision_context()
        _res = v17_dec.synthesize_decisions(_ctx, cite=False)
        v17_dec.record_decisions(_res["decisions"], _ctx.get("prices"))
        v17_dec.resolve_decisions(_ctx.get("prices"))
    except Exception:
        pass
    try:
        push_strong_decisions()         # P20 — proactive STRONG-call Telegram push
    except Exception:
        pass
    while True:
        now = datetime.now()
        target = now.replace(hour=2, minute=30, second=0, microsecond=0)
        if now >= target:
            target += timedelta(days=1)
        _t.sleep(max(60, (target - now).total_seconds()))
        try:
            v11_cog.run_consolidation(force=True)
        except Exception:
            pass
        try:
            swot_engine.generate_daily_swot()
        except Exception:
            pass
        try:
            # P16 sleep-phase: curate the day's new facts, then dedup/decay/promote.
            v16_mem.ingest_recent()
            v16_mem.consolidate()
        except Exception:
            pass
        try:
            fetch_india_macro_plus()    # P20 — refresh macro/trade beliefs nightly
        except Exception:
            pass
        try:
            # P17 sleep-phase: record the day's decisions + grade matured ones.
            _ctx = _decision_context()
            _res = v17_dec.synthesize_decisions(_ctx, cite=False)
            v17_dec.record_decisions(_res["decisions"], _ctx.get("prices"))
            v17_dec.resolve_decisions(_ctx.get("prices"))
        except Exception:
            pass
        try:
            push_strong_decisions()     # P20 — proactive STRONG-call Telegram push
        except Exception:
            pass
        try:
            # Curiosity Engine — nightly full thinking cycle
            _recent = v11_sources.recent_articles(hours=6, limit=300)
            _curiosity.think_cycle(_recent)
        except Exception:
            pass

def _heartbeat_confirm():
    """P15 C3 — confirm the ingestion worker reported a heartbeat within 90s of
    boot. Does not block serving; logs loudly if the worker failed to start."""
    import time as _t
    for _ in range(18):                 # 18 × 5s = 90s
        _t.sleep(5)
        if v11_sources.kv_get("worker_heartbeat"):
            print("  [OK] ingestion worker heartbeat confirmed.")
            return
    print("  [!!] ingestion worker produced NO heartbeat in 90s -- see /api/health.")

@app.route("/api/swot")
def api_swot():
    return jsonify(swot_engine.get_latest_swot())

@app.route("/api/swot/run", methods=["POST"])
def api_swot_run():
    return jsonify(swot_engine.generate_daily_swot())

# ── Brain Centers — Master Meta-Cognitive Space ───────────────────────────────
# All /api/brain routes require header: X-Neuron-Key: <value of NEURON_MASTER_KEY in .env>

_BRAIN_AUTH_FAILS: dict = {}   # ip -> [ts1, ts2, …] rolling window
_BRAIN_RATE_WINDOW = 60        # seconds
_BRAIN_RATE_LIMIT  = 10        # max failed attempts per window per IP

def _brain_auth():
    """
    Returns None if authorised, else a Flask error response.
    - Constant-time key comparison (hmac.compare_digest via metacog.verify_key)
    - Rate-limits failed attempts: 10 bad keys / 60s per IP → 429
    - Never reveals the expected key in error messages
    """
    import time as _t
    ip = request.remote_addr or "unknown"
    key = request.headers.get("X-Neuron-Key", "")

    # Rate-limit check on IP before expensive verify
    now = _t.time()
    fails = _BRAIN_AUTH_FAILS.get(ip, [])
    fails = [ts for ts in fails if now - ts < _BRAIN_RATE_WINDOW]
    if len(fails) >= _BRAIN_RATE_LIMIT:
        _security_center.log_access(request.path, False, ip, "rate-limited")
        return jsonify({"error": "Too many requests — try again later"}), 429

    granted = _metacog.verify_key(key)
    _security_center.log_access(request.path, granted, ip,
                                 "OK" if granted else "bad key")
    if not granted:
        fails.append(now)
        _BRAIN_AUTH_FAILS[ip] = fails
        return jsonify({"error": "Unauthorized"}), 401
    # Clear failure history on successful auth
    _BRAIN_AUTH_FAILS.pop(ip, None)
    return None

@app.route("/api/brain")
def api_brain():
    """Master view through the Meta-Cognitive Space — full brain state."""
    err = _brain_auth()
    if err:
        return err
    return jsonify(_metacog.full_brain_state())

@app.route("/api/brain/health")
def api_brain_health():
    """Lightweight center health — all 7 centers, no heavy computation."""
    err = _brain_auth()
    if err:
        return err
    health = _metacog.center_health()
    assessment = _metacog.metacognitive_assessment(health)
    return jsonify({"assessment": assessment, "centers": health})

@app.route("/api/brain/route")
def api_brain_route():
    """Route a natural-language query to the correct Brain Center."""
    err = _brain_auth()
    if err:
        return err
    q = request.args.get("q", "").strip()
    if not q:
        return jsonify({"error": "Provide ?q=<your question>"}), 400
    return jsonify(_metacog.route_query(q))

@app.route("/api/brain/center/<center_key>")
def api_brain_center(center_key):
    """Full detailed report from a specific Brain Center."""
    err = _brain_auth()
    if err:
        return err
    return jsonify(_metacog.brain_report(center_key))

@app.route("/api/brain/backup", methods=["POST"])
def api_brain_backup():
    """Trigger a DB backup from the Backup Center."""
    err = _brain_auth()
    if err:
        return err
    note = request.json.get("note", "api-trigger") if request.is_json else "api-trigger"
    return jsonify(_backup_center.backup(note=note))

@app.route("/api/brain/restore", methods=["POST"])
def api_brain_restore():
    """Restore DB from a backup (requires confirm_token in body)."""
    err = _brain_auth()
    if err:
        return err
    data = request.get_json(silent=True) or {}
    filename      = data.get("filename", "")
    confirm_token = data.get("confirm_token", "")
    if not filename:
        return jsonify({"error": "Provide filename in JSON body"}), 400
    return jsonify(_backup_center.restore(filename, confirm_token))

# ── Curiosity Engine — Public Endpoints (Neuron's open mind, no key needed) ──

@app.route("/api/thoughts")
def api_thoughts():
    """Neuron's recent autonomous observations — its stream of consciousness."""
    limit = min(int(request.args.get("limit", 20)), 50)
    return jsonify({"thoughts": _curiosity.get_today_thoughts(limit=limit)})

@app.route("/api/thoughts/today")
def api_thoughts_today():
    """Full curiosity cycle summary for today."""
    stats = _curiosity.curiosity_stats()
    thoughts = _curiosity.get_today_thoughts(limit=10)
    return jsonify({"stats": stats, "thoughts": thoughts})

@app.route("/api/curiosities")
def api_curiosities():
    """Open questions Neuron is currently investigating."""
    return jsonify({"curiosities": _curiosity.get_open_curiosities(limit=20)})

@app.route("/api/insights")
def api_insights():
    """Resolved insights — what Neuron figured out on its own."""
    limit = min(int(request.args.get("limit", 10)), 30)
    return jsonify({"insights": _curiosity.get_recent_insights(limit=limit)})

# ── Curiosity Engine — Secured Endpoints (X-Neuron-Key required) ─────────────

@app.route("/api/brain/curiosity")
def api_brain_curiosity():
    """Full curiosity center report."""
    err = _brain_auth()
    if err: return err
    return jsonify(_metacog.brain_report("curiosity_center"))

@app.route("/api/brain/agenda")
def api_brain_agenda():
    """What Neuron plans to investigate next cycle."""
    err = _brain_auth()
    if err: return err
    import sqlite3 as _sq
    con = _sq.connect(v11_sources.DB_PATH, timeout=10)
    rows = con.execute(
        "SELECT item, priority, cycle FROM v24_agenda ORDER BY priority DESC LIMIT 10"
    ).fetchall()
    con.close()
    return jsonify({"agenda": [{"item": r[0], "priority": r[1], "cycle": r[2]} for r in rows]})

@app.route("/api/brain/learning")
def api_brain_learning():
    """Neuron's learning journal — what it was right and wrong about."""
    err = _brain_auth()
    if err: return err
    import sqlite3 as _sq
    con = _sq.connect(v11_sources.DB_PATH, timeout=10)
    rows = con.execute(
        "SELECT ts, insight_id, was_right, note FROM v24_learning ORDER BY ts DESC LIMIT 20"
    ).fetchall()
    con.close()
    return jsonify({"learning": [{"ts":r[0],"insight_id":r[1],"was_right":bool(r[2]),"note":r[3]} for r in rows]})

@app.route("/api/brain/think-now", methods=["POST"])
def api_brain_think_now():
    """Trigger an immediate curiosity thinking cycle."""
    err = _brain_auth()
    if err: return err
    articles = v11_sources.recent_articles(hours=6, limit=200)
    result = _curiosity.think_cycle(articles)
    return jsonify(result)

# ── Father interaction hook — every /api/ask call teaches Neuron ──────────────
# NOTE: This wraps the existing /api/ask endpoint (defined earlier in neuron.py).
# We add a post-processing hook so Neuron learns from father's questions.
# The original endpoint is unchanged; this complementary route captures the intent.

@app.route("/api/father/chat", methods=["POST"])
def api_father_chat():
    """
    Father sends a message to Neuron in the hidden room.
    POST {"message": "..."} → returns father msg + Neuron's response.
    """
    data = request.get_json(silent=True) or {}
    msg = (data.get("message") or "").strip()[:2000]   # hard cap — no prompt stuffing
    if not msg:
        return jsonify({"error": "Provide message"}), 400
    try:
        result = _curiosity.chat_with_father(msg)
        return jsonify({"ok": True, **result})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/father/history")
def api_father_history():
    """Return the full father-child conversation history."""
    limit = min(int(request.args.get("limit", 80)), 200)
    try:
        history = _curiosity.get_father_conversation(limit=limit)
        return jsonify({"ok": True, "messages": history, "count": len(history)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/ask/learn", methods=["POST"])
def api_ask_learn():
    """
    When father asks Neuron something, Neuron records it as a high-priority curiosity.
    POST {"question": "...", "response": "..."} — called automatically after /api/ask.
    """
    data = request.get_json(silent=True) or {}
    q = (data.get("question") or "").strip()[:1500]
    r = (data.get("response") or "").strip()[:3000]
    if not q:
        return jsonify({"error": "Provide question"}), 400
    _curiosity.learn_from_father(q, r)
    return jsonify({"ok": True, "message": "Neuron has recorded your question as a priority curiosity."})

if __name__ == "__main__":
    print("=" * 62)
    print("  NEURON v7 - Indian RE Intelligence Monitor")
    print("  by Vipul Jakhar")
    print("  P15 Nervous System | Beliefs · Diff Engine · Watchdog")
    print("  -> http://localhost:5000")
    print("=" * 62)
    import threading
    # Phase 0 — ensure all schemas exist, then run synchronous boot checks.
    v11_sources.init_v11_tables()
    v11_cog.init_cognition_tables()
    v16_mem.init_memory_tables()
    v17_dec.init_decision_tables()
    _curiosity.init_curiosity_tables()
    threading.Thread(target=_curiosity.warm_up, daemon=True).start()
    # P16.2 — set the real embedder before any fact is embedded, then migrate
    # existing floor vectors to it in the background (never blocks serving).
    if _init_embedder():
        threading.Thread(target=lambda: v16_mem.reembed_all(), daemon=True).start()
    boot_diagnostics()
    # Daily brief + nightly consolidation background threads
    threading.Thread(target=_daily_brief_worker, daemon=True).start()
    threading.Thread(target=_consolidation_worker, daemon=True).start()
    # Pre-warm the heaviest caches at boot so the first visitor never pays
    # the 2-5 min ALMM PDF parse interactively
    def _prewarm():
        try: fetch_almm()
        except Exception: pass
        try: fetch_almm_modules()
        except Exception: pass
    threading.Thread(target=_prewarm, daemon=True).start()
    # v11 Observatory: 585-source tiered ingestion (daemon, never blocks requests)
    v11_sources.start_ingestion()
    threading.Thread(target=_heartbeat_confirm, daemon=True).start()
    app.run(debug=False, port=5000, threaded=True)
